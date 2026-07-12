from __future__ import annotations

import io
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_图表精简改稿版.docx")
OUT = SRC.with_name("毕业设计正文模板_图表黑白统一版.docx")
ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")


ABSTRACT_ADD = (
    "与单纯页面展示相比，本文更加关注监测指标与运维动作之间的对应关系。系统在诊断结果中保留风险因子、主要异常来源、健康评分变化和建议处置窗口，使运维人员能够沿着“发现异常、查看依据、生成建议、记录闭环”的路径理解系统输出，也便于后续接入真实风场数据后继续扩展。"
)


REMOVE_TABLE_CAPTIONS = {
    "表1.1 本文研究内容与系统产出",
    "表2.2 系统非功能需求说明",
    "表2.3 系统数据来源与用途",
    "表3.2 振动特征提取指标",
    "表4.2 数据库表关系说明",
    "表4.4 后端接口分组说明",
    "表4.5 诊断流程关键输入输出",
    "表5.1 数据采集模块输出字段",
    "表5.2 故障诊断模块实现要点",
    "表6.2 典型诊断场景测试结果",
    "表7.1 系统不足与改进方向",
}


TEXT_REPLACEMENTS = {
    "本文研究内容由故障机理、数据融合、诊断模型、Web 平台和系统测试五部分组成，见表1.1。": "本文研究内容包括齿轮箱故障机理分析、多源状态数据融合、风险评分与诊断规则设计、Web 平台实现以及系统功能测试五部分。",
    "非功能需求主要包括可用性、可扩展性、可解释性和轻量部署，见表2.2。": "非功能需求主要包括可用性、可扩展性、可解释性和轻量部署。界面需要便于快速识别异常，后端结构需要为真实 SCADA、CMS 和油液数据接入预留空间。",
    "系统数据需求见表2.3。两类数据共同支撑风场总览、单机详情、诊断判断和健康报告。": "系统数据来源包括模拟采集数据和风机数据摘要文件。前者用于实时监测与诊断演示，后者用于风场总览、单机详情、故障统计和健康报告。",
    "振动特征的含义和用途见表3.2。系统重点关注能量类、冲击类和包络类指标。": "系统重点关注能量类、冲击类和包络类振动指标，其中 RMS 反映整体振动能量，峭度和峰值因子反映冲击特征，包络峰值用于辅助判断局部损伤。",
    "诊断流程以多源状态数据为输入，经清洗、特征提取、风险评分和规则判断后，输出故障类型、健康评分、RUL、诊断依据和维护建议。关键输入输出见表4.5。": "诊断流程以多源状态数据为输入，经清洗、特征提取、风险评分和规则判断后，输出故障类型、健康评分、RUL、诊断依据和维护建议。",
    "诊断场景测试结果见表6.2。系统通过不同输入变化验证故障类型、风险等级和维护建议是否一致。": "系统通过正常运行、齿轮磨损、点蚀剥落、油温过高、油液污染和基础松动等场景验证诊断输出的一致性。",
    "系统不足和改进方向见表7.1。当前原型已完成核心演示流程，但真实数据验证、模型训练和生产级部署仍需继续完善。": "当前原型已完成核心演示流程，但真实数据验证、模型训练、RUL 参数校准和生产级部署仍需继续完善。",
}


CAPTION_RENAMES = {
    "表3.3 齿轮箱健康评估主要风险因子": "表3.2 齿轮箱健康评估主要风险因子",
    "表4.3 系统主要后端接口": "表4.2 系统主要后端接口",
    "表4.6 前端页面与交互功能对应": "表4.3 前端页面与交互功能对应",
}


EXTRA_FIGURES = [
    ("图5.12 系统登录入口页面", ROOT / "docs/system_screenshots_updated/01_main_login.png"),
    ("图5.13 登录后风场总览补充页面", ROOT / "docs/system_screenshots_updated/02_after_login_windfarm.png"),
    ("图5.14 实时监测运行概览页面", ROOT / "docs/system_screenshots_updated/03_realtime_dashboard.png"),
    ("图5.15 振动趋势监测页面", ROOT / "docs/system_screenshots_updated/04_dashboard_vibration_tab.png"),
    ("图5.16 HMI故障代码页面", ROOT / "docs/system_screenshots_updated/08_hmi_fault_code_page.png"),
    ("图5.17 AI问答助手补充页面", ROOT / "docs/system_screenshots_updated/09_ai_qa_assistant_updated.png"),
    ("图5.18 单机详情运行页面", ROOT / "docs/system_screenshots/07_turbine_detail.png"),
    ("图5.19 故障诊断运行页面", ROOT / "docs/system_screenshots/02_fault_diagnosis.png"),
    ("图5.20 故障记录管理页面", ROOT / "docs/system_screenshots/03_fault_records.png"),
    ("图5.21 健康报告运行页面", ROOT / "docs/system_screenshots/04_health_report.png"),
    ("图5.22 参数设置运行页面", ROOT / "docs/system_screenshots/06_settings.png"),
    ("图5.23 齿轮箱ER关系图补充", ROOT / "docs/system_screenshots/gearbox_er_diagram.png"),
    ("图5.24 系统功能架构补充图", ROOT / "docs/figures/system_function_architecture.png"),
    ("图5.25 健康管理流程补充图", ROOT / "docs/figures/gearbox_health_management_flowchart.png"),
]


def paragraph_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def set_paragraph_text(paragraph, text):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def remove_captioned_tables(doc):
    body = doc._body._element
    children = list(body)
    idx = 0
    while idx < len(children):
        child = children[idx]
        if child.tag == qn("w:p") and paragraph_text(child) in REMOVE_TABLE_CAPTIONS:
            body.remove(child)
            if idx + 1 < len(children) and children[idx + 1].tag == qn("w:tbl"):
                body.remove(children[idx + 1])
                idx += 2
            else:
                idx += 1
            continue
        idx += 1


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def normalize_tables(doc):
    for table in doc.tables:
        table.style = "Normal Table"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                clear_cell_shading(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(p.text.strip()) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "宋体"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(9)
                        run.font.color.rgb = None


def normalize_captions(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.match(r"^[图表]\d+(\.\d+)+", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(10)
                run.bold = False


def add_abstract_paragraph(doc):
    if any(ABSTRACT_ADD in p.text for p in doc.paragraphs):
        return
    keyword = next(p for p in doc.paragraphs if p.text.strip().startswith("关键词"))
    new_p = doc.add_paragraph()
    set_paragraph_text(new_p, ABSTRACT_ADD)
    keyword._p.addprevious(new_p._p)


def apply_text_replacements(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in TEXT_REPLACEMENTS:
            set_paragraph_text(p, TEXT_REPLACEMENTS[text])
        elif text in CAPTION_RENAMES:
            set_paragraph_text(p, CAPTION_RENAMES[text])


def add_picture_after(paragraph, caption, img_path):
    pic_p = paragraph._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(img_path), width=Inches(5.8))
    paragraph._p.addnext(pic_p._p)

    cap_p = paragraph._parent.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = "宋体"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap_run.font.size = Pt(10)
    pic_p._p.addnext(cap_p._p)
    return cap_p


def add_extra_figures(doc):
    if any(p.text.strip().startswith("图5.25") for p in doc.paragraphs):
        return
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "图5.11 系统参数设置页面")
    cursor = anchor
    for caption, path in EXTRA_FIGURES:
        if path.exists():
            cursor = add_picture_after(cursor, caption, path)


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def grayscale_media(docx_path):
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            lower = item.filename.lower()
            if lower.startswith("word/media/") and lower.endswith((".png", ".jpg", ".jpeg")):
                try:
                    img = Image.open(io.BytesIO(data))
                    gray = img.convert("L").convert("RGB")
                    out = io.BytesIO()
                    fmt = "JPEG" if lower.endswith((".jpg", ".jpeg")) else "PNG"
                    gray.save(out, format=fmt)
                    data = out.getvalue()
                except Exception:
                    pass
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    add_abstract_paragraph(doc)
    apply_text_replacements(doc)
    remove_captioned_tables(doc)
    add_extra_figures(doc)
    normalize_tables(doc)
    normalize_captions(doc)
    set_update_fields(doc)
    doc.save(OUT)
    grayscale_media(OUT)

    checked = Document(OUT)
    captions = [p.text.strip() for p in checked.paragraphs if re.match(r"^[图表]\d+(\.\d+)+", p.text.strip())]
    fills = []
    for table in checked.tables:
        for row in table.rows:
            for cell in row.cells:
                shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if shd is not None:
                    fills.append(shd.get(qn("w:fill")))
    print(f"saved={OUT}")
    print(f"paragraphs={sum(1 for p in checked.paragraphs if p.text.strip())}")
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")
    print(f"captions={len(captions)}")
    print(f"table_shading_count={len(fills)}")


if __name__ == "__main__":
    main()
