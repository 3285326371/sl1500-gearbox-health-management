from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_结构图美化版.docx")
OUT = SRC.with_name("毕业设计正文模板_无边框摘要对应版.docx")


EN_ABSTRACT = [
    "This thesis takes the gearbox of the Sinovel SL1500 doubly-fed wind turbine as the research object and implements a health management system for graduation-design demonstration and prototype verification. During development, simulated SCADA/CMS data, wind turbine summary files, oil temperature, vibration and NAS oil-quality indicators are integrated into a unified set of pages and interfaces. The backend is built with Flask and SQLite, while the frontend uses HTML, CSS and JavaScript, together with ECharts, Three.js and SSE, to implement wind farm overview, turbine detail monitoring, real-time monitoring, fault diagnosis, health reports and operation assistance.",
    "The diagnosis module focuses on vibration RMS, peak value, kurtosis, crest factor, impulse factor, envelope peak, oil temperature and oil quality grade. A rule-based fusion method is used to identify typical scenarios such as gear wear, pitting and spalling, broken teeth, bearing damage, high oil temperature, oil contamination, shaft misalignment, foundation looseness and cooling abnormalities. The system outputs health score, RUL, diagnostic evidence and maintenance suggestions. Test results show that the prototype has completed the basic workflow of data acquisition, risk calculation, record storage, report generation and simulated work order processing, and can serve as an initial form of an intelligent gearbox operation and maintenance platform.",
    "In terms of functional organization, this thesis does not stop at page display, but connects monitoring indicators with subsequent maintenance actions. The diagnostic result keeps risk factors, abnormal sources, health-score changes and suggested handling windows, so that operators can understand the system output along the path of discovering abnormalities, checking evidence, generating suggestions and recording the closed loop. If real wind farm data are connected later, the current interfaces and data structures can continue to be extended.",
]


def set_paragraph_text(paragraph, text):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
        paragraph._p = paragraph._element = None


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    en_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "ABSTRACT")
    key_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Key words"))
    body = [p for p in doc.paragraphs[en_idx + 1 : key_idx] if p.text.strip()]

    while len(body) < len(EN_ABSTRACT):
        new_p = doc.add_paragraph()
        doc.paragraphs[key_idx]._p.addprevious(new_p._p)
        body.append(new_p)
        key_idx += 1

    for p, text in zip(body, EN_ABSTRACT):
        set_paragraph_text(p, text)
    for p in body[len(EN_ABSTRACT) :]:
        remove_paragraph(p)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUT)
    checked = Document(OUT)
    en_idx = next(i for i, p in enumerate(checked.paragraphs) if p.text.strip() == "ABSTRACT")
    key_idx = next(i for i, p in enumerate(checked.paragraphs) if p.text.strip().startswith("Key words"))
    en_count = sum(1 for p in checked.paragraphs[en_idx + 1 : key_idx] if p.text.strip())
    cn_idx = next(i for i, p in enumerate(checked.paragraphs) if p.text.strip() == "摘  要")
    kw_idx = next(i for i, p in enumerate(checked.paragraphs) if p.text.strip().startswith("关键词"))
    cn_count = sum(1 for p in checked.paragraphs[cn_idx + 1 : kw_idx] if p.text.strip())
    print(f"saved={OUT}")
    print(f"cn_abstract_paragraphs={cn_count}")
    print(f"en_abstract_paragraphs={en_count}")
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")


if __name__ == "__main__":
    main()
