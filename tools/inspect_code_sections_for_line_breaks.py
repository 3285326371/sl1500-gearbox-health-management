from pathlib import Path

from docx import Document


DOCX = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板.docx")

doc = Document(DOCX)
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if "代码清单" in text or "核心代码如下" in text:
        print("\nMARK", i, paragraph.style.name, repr(text))
        for j in range(i + 1, min(i + 28, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.rstrip()
            if not t:
                continue
            print(j, doc.paragraphs[j].style.name, repr(t[:180]))
            if j > i + 3 and ("代码清单" in t or t.startswith("图") or t.startswith("表") or t.startswith("第") or t.startswith("当用户") or t.startswith("系统") or t.startswith("健康")):
                break
