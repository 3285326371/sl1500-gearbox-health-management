from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_改题目.docx"
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_改题目前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)

OLD_TITLES = [
    "华锐SL1500型双馈式风机齿轮箱智能健康管理系统设计与实现",
    "华锐 SL1500 型双馈式风机齿轮箱智能健康管理系统设计与实现",
    "华锐SL1500型双馈式风机齿轮箱\n智能健康管理系统设计与实现",
    "华锐 SL1500 型双馈式风机齿轮箱智能健康管理系统",
    "华锐SL1500型双馈式风机齿轮箱智能健康管理系统",
]
NEW_TITLE = "华锐SL1500型双馈式风机齿轮箱智能健康管理体系"


def set_font(run, east_asia="黑体", ascii_font="Arial", size=22, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_cover_title(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.25
    r1 = paragraph.add_run("华锐SL1500型双馈式风机齿轮箱")
    set_font(r1)
    r1.add_break()
    r2 = paragraph.add_run("智能健康管理体系")
    set_font(r2)


def replace_in_paragraph(paragraph):
    text = paragraph.text
    new = text
    for old in OLD_TITLES:
        new = new.replace(old, NEW_TITLE)
    if new != text:
        # Preserve simple paragraph-level formatting well enough for title/declaration cases.
        style = paragraph.style
        alignment = paragraph.alignment
        paragraph.clear()
        paragraph.style = style
        paragraph.alignment = alignment
        run = paragraph.add_run(new)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)

    # Cover title is paragraph 6 in the school template after prior in-place edits.
    set_cover_title(doc.paragraphs[6])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"docx={TEMPLATE}")
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
