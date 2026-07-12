from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")
DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
SHOT_DIR = ROOT / "docs" / "system_screenshots_updated"


def font_run(run, size=12, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def style_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for r in p.runs:
        font_run(r, 12)


def style_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = False
    for r in p.runs:
        font_run(r, 10.5)


def style_image_paragraph(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"Paragraph not found: {text}")


def paragraph_after(anchor, text="", style=None):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    if style:
        p.style = style
    if text:
        style_body(p)
    return p


def image_after(anchor, image_name: str, width_cm=14.8):
    p = anchor._parent.add_paragraph()
    anchor._p.addnext(p._p)
    style_image_paragraph(p)
    run = p.add_run()
    run.add_picture(str(SHOT_DIR / image_name), width=Cm(width_cm))
    return p


def replace_paragraph_image(p, image_name: str, width_cm=14.8):
    p.clear()
    style_image_paragraph(p)
    p.add_run().add_picture(str(SHOT_DIR / image_name), width=Cm(width_cm))


def update_text(doc: Document):
    replacements = {
        "前端首页包含登录界面和主应用界面。登录后默认进入健康总览，左侧导航按照健康监测区、智能诊断区和系统配置区分组，顶部显示当前模块说明和通知。风场总览页面展示多台机组卡片，单机详情页面展示齿轮箱温度、寿命、部件状态、告警和集群对标。实时监测页面展示油温、振动、功率、RUL、采集链路、振动波形和数字孪生热图。":
        "前端首页包含登录、注册和主应用界面。用户通过登录页进入系统，系统根据角色显示高级工程师或运维人员身份；注册页可选择运维人员或管理员角色。登录后默认进入健康总览，左侧导航按照健康监测区、智能诊断区和系统配置区分组，顶部显示当前模块说明、通知状态和用户身份。风场总览页面展示多台机组卡片，单机详情页面展示齿轮箱温度、寿命、部件状态、告警和集群对标。实时监测页面进一步改为多标签结构，包含运行概览、振动趋势、孪生与采集、故障统计等页面。",
        "故障诊断页面提供场景选择和诊断按钮，运行后展示疑似故障类型、风险等级、关联部件、置信度、健康评分、RUL、特征值、风险因子、诊断依据、建议措施和工单信息。故障数据页面展示历史故障记录，报告页面展示健康评估结果，问答页面支持输入自然语言问题并获得排查建议。HMI 页面则以单机齿轮箱为对象，提供运行总览、状态寿命、透明模型、数据采集、测点信号、阈值设置、故障代码和临界警报等功能。":
        "故障诊断页面提供场景选择和诊断按钮，运行后展示疑似故障类型、风险等级、关联部件、置信度、健康评分、RUL、特征值、风险因子、诊断依据、建议措施和工单信息。故障数据页面展示历史故障记录，报告页面展示健康评估结果，问答页面支持输入自然语言问题并获得排查建议。更新后的 HMI 页面增加二次权限登录，登录后进入独立的单机齿轮箱控制台，提供运行总览、状态寿命、透明模型、数据采集、测点信号、阈值设置、故障代码和临界警报等功能。",
        "HMI 页面更强调现场设备交互。该页面围绕单台齿轮箱展开，将齿轮箱油温、高速轴承温度、振动RMS、润滑油等级、健康评分和报警状态集中到同一视图，并预留透明模型、测点信号、阈值设置和临界报警入口。相比普通管理页面，HMI 界面能够体现设备对象、运行状态和操作反馈之间的对应关系，是系统面向现场监控应用的重要补充。":
        "HMI 页面更强调现场设备交互。该页面围绕单台齿轮箱展开，用户从单机详情页点击 HMI 运维控制台后，首先进入权限登录页面，确认账号、密码和角色级别，再进入控制台。登录后的 HMI 页面将齿轮箱油温、振动RMS、健康评分、RUL、运行状态和故障风险集中到同一视图，并提供状态寿命、透明模型、数据采集、测点信号、阈值设置、故障代码和临界警报等菜单。相比普通管理页面，HMI 界面能够体现设备对象、运行状态和操作反馈之间的对应关系，是系统面向现场监控应用的重要补充。",
        "HMI页面主要用于展示齿轮箱关键运行信息和人机交互状态。与普通Web仪表盘相比，HMI页面更强调现场监控感和设备对象感，适合在答辩演示时说明系统具备工业监控界面的扩展能力。页面可根据机组编号加载不同设备状态，并与后端接口保持一致的数据来源。":
        "HMI页面主要用于展示齿轮箱关键运行信息和人机交互状态。与普通Web仪表盘相比，HMI页面更强调现场监控感、权限确认和设备对象感，适合在答辩演示时说明系统具备工业监控界面的扩展能力。页面可根据机组编号加载不同设备状态，并与后端接口保持一致的数据来源；阈值设置页面还能根据操作员、工程师或管理员权限决定是否允许修改关键阈值。",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            p.text = replacements[text]
            style_body(p)

    # Update the HMI figure lead-in and caption.
    for p in doc.paragraphs:
        if p.text.strip() == "齿轮箱HMI页面如图5.8所示。该页面面向现场监控场景，用更加接近工业界面的方式展示单机齿轮箱状态和关键指标。":
            p.text = "齿轮箱HMI权限登录页面如图5.8所示。该页面面向现场监控场景，在进入控制台前要求用户确认账号、密码和权限级别。"
            style_body(p)
        if p.text.strip() == "图5.8 齿轮箱HMI人机交互页面":
            p.text = "图5.8 齿轮箱HMI权限登录页面"
            style_caption(p)


def update_table(doc: Document):
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) >= 4 and table.cell(0, 0).text.strip() == "前端页面":
            for row in table.rows:
                if row.cells[0].text.strip() == "HMI界面":
                    row.cells[1].text = "单机详情-HMI控制台，二次权限登录"
                    row.cells[2].text = "运行总览、阈值设置、故障代码、测点和临界报警"
                    row.cells[3].text = "hmi.html、settings_route.py、data_route.py"
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.first_line_indent = None
                            for r in p.runs:
                                font_run(r, 10.5)


def update_images(doc: Document):
    # Replace old windfarm overview screenshot with the updated logged-in view.
    fig51 = find_para(doc, "图5.1 风场健康总览页面")
    if fig51._p.getprevious() is not None:
        prev_para = None
        for p in doc.paragraphs:
            if p._p is fig51._p.getprevious():
                prev_para = p
                break
        if prev_para is not None:
            replace_paragraph_image(prev_para, "02_after_login_windfarm.png", 15.2)

    # Replace old HMI image with the updated HMI login page.
    fig58 = find_para(doc, "图5.8 齿轮箱HMI权限登录页面")
    if fig58._p.getprevious() is not None:
        prev_para = None
        for p in doc.paragraphs:
            if p._p is fig58._p.getprevious():
                prev_para = p
                break
        if prev_para is not None:
            replace_paragraph_image(prev_para, "05_hmi_login.png", 13.0)

    after_fig58 = paragraph_after(fig58, "用户完成 HMI 权限登录后进入运行总览页面，如图5.9所示。该页面左侧为 HMI 功能菜单，顶部显示当前机组和运行状态，中部展示油温、振动、健康评分、RUL 等关键指标，下方提供运行总览数据。")
    img = image_after(after_fig58, "06_hmi_after_login_overview.png", 15.2)
    cap = paragraph_after(img, "图5.9 HMI登录后运行总览页面")
    style_caption(cap)
    p = paragraph_after(cap, "HMI 登录后页面将单机关键指标和控制入口集中到一个界面中，便于运维人员在确认权限后快速查看齿轮箱状态。阈值设置和故障代码页面进一步提供配置与排查入口，其中阈值设置与后端 settings_route.py 参数接口关联，故障代码页面与 data_route.py 的故障码和故障状态数据关联。")
    style_body(p)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_前后端更新论文修改前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)
    update_text(doc)
    update_table(doc)
    update_images(doc)
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
