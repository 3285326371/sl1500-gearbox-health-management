from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_附录B前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_附录B.docx"


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_p(doc, text="", first=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24) if first else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=16 if level == 1 else 12, bold=True)
    return p


def shade_cell(cell, fill="F4F4F4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code(doc, title, code):
    p = doc.add_paragraph(title)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=10.5, bold=True)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell)
    cell.text = ""
    cp = cell.paragraphs[0]
    cp.paragraph_format.line_spacing = Pt(10)
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(3)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            cp.add_run("\n")
        r = cp.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(30, 30, 30)


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    if any("附录B 核心代码片段与接口说明" in p.text for p in doc.paragraphs):
        print("already exists")
    else:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_heading(doc, "附录B 核心代码片段与接口说明")
        add_p(doc, "本附录列出系统实现中具有代表性的核心代码片段，主要用于说明后端应用初始化、诊断接口组织、前端页面切换和实时数据刷新等关键实现。正文第5章已经给出部分算法代码，本附录进一步补充系统工程层面的代码结构，使论文内容与实际项目代码保持一致。")
        add_code(doc, "代码清单B.1 Flask应用初始化与蓝图注册", """
def create_app():
    app = Flask(__name__, static_folder='../frontend', static_url_path='')
    CORS(app)
    project_root = Path(app.root_path).resolve().parent
    database_path = project_root / 'instance' / 'gearbox_system.db'
    app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{database_path.as_posix()}"
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

    from models.database import init_db
    init_db(app)

    @app.route('/')
    def index():
        return app.send_static_file('index.html')

    app.register_blueprint(data_bp, url_prefix='/api/data')
    app.register_blueprint(windfarm_bp, url_prefix='/api/windfarm')
    app.register_blueprint(report_bp, url_prefix='/api/report')
    app.register_blueprint(qa_bp, url_prefix='/api/qa')
    app.register_blueprint(settings_bp, url_prefix='/api/settings')
    return app
""")
        add_p(doc, "代码清单B.1展示了系统后端入口的组织方式。Flask应用创建后首先配置静态前端目录、数据库路径和跨域访问，然后初始化数据库模型，并按业务域注册多个蓝图。该结构使认证、数据、风场、报告、问答和设置接口相互独立，便于后续维护。")
        add_code(doc, "代码清单B.2 故障诊断接口组织方式", """
@data_bp.route('/diagnosis', methods=['POST'])
def diagnose():
    payload = request.get_json(silent=True) or {}
    unit_id = payload.get('unit_id', 'WTG-001')
    scenario = payload.get('scenario', 'normal')
    raw_signal = generate_vibration_data(scenario=scenario)
    status = build_status_from_scenario(unit_id, scenario)
    result = run_diagnosis(unit_id=unit_id, raw_signal=raw_signal, status=status)
    record = FaultRecord(
        unit_id=unit_id,
        fault_type=result['fault_type'],
        severity=result['severity'],
        probability=result['probability'],
        health_score=result['health_score'],
        rul_days=result['rul_days'],
        advice=result['advice'],
        status='pending'
    )
    db.session.add(record)
    db.session.commit()
    return jsonify(result)
""")
        add_p(doc, "代码清单B.2体现了故障诊断接口的基本流程。接口接收前端传入的机组编号和诊断场景，生成或读取振动数据与状态数据，调用run_diagnosis完成诊断，并将结果写入故障记录表。这样可以保证诊断结果既能在页面展示，也能在故障数据管理模块中追踪。")
        add_code(doc, "代码清单B.3 前端模块切换与视图控制", """
function showView(targetId) {
    const navItems = document.querySelectorAll('.nav-item');
    const sections = document.querySelectorAll('.view-section');
    navItems.forEach(item => {
        item.classList.toggle('active', item.dataset.target === targetId);
    });
    sections.forEach(section => {
        section.style.display = section.id === targetId ? 'block' : 'none';
        section.classList.toggle('active', section.id === targetId);
    });
    updateModuleIndicator(targetId);
}

document.querySelectorAll('.nav-item').forEach(item => {
    item.addEventListener('click', event => {
        event.preventDefault();
        showView(item.dataset.target);
    });
});
""")
        add_p(doc, "代码清单B.3展示了前端单页应用的视图切换方式。系统通过导航项的data-target属性确定目标模块，然后隐藏其他section并显示当前模块。该方式实现简单、依赖少，适合本系统这种本地演示型健康管理平台。")
        add_code(doc, "代码清单B.4 健康报告关键指标渲染", """
async function fetchReport() {
    const unitId = document.getElementById('report-unit-select').value;
    const days = document.getElementById('report-range-select').value;
    const response = await fetch(`/api/report/generate?unit_id=${unitId}&days=${days}`);
    const report = await response.json();
    document.getElementById('report-health-score').textContent = report.health_score;
    document.getElementById('report-oil-temp').textContent = report.oil_temp + ' ℃';
    document.getElementById('report-vibration').textContent = report.vibration_rms + ' mm/s';
    document.getElementById('report-rul').textContent = report.rul_days + ' 天';
    document.getElementById('report-conclusion').textContent = report.conclusion;
    renderReportTrend(report.trend);
}
""")
        add_p(doc, "代码清单B.4说明健康报告页面如何通过接口获取数据并渲染关键指标。报告页面不是静态文本，而是由后端根据当前机组状态、历史趋势和诊断记录动态生成。该实现使报告内容能够随系统数据变化而更新。")
        add_p(doc, "综上，附录B中的代码片段覆盖了系统启动、接口注册、诊断接口、前端模块切换和报告渲染等关键环节，与正文中的系统需求、总体设计、详细实现和测试分析相互对应，能够证明本文系统并非停留在概念设计层面，而是具有可运行的工程实现。")

    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
