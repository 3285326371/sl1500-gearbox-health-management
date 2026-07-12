from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "风电场数据对接与项目阶段成果说明.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_style(table):
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(15, 76, 129) if level == 1 else RGBColor(30, 64, 175)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_style(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], "DCEBFA")
        set_cell_text(hdr[idx], header, bold=True, color="0F4C81")
        if widths:
            hdr[idx].width = Cm(widths[idx])
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)
            if widths:
                row[idx].width = Cm(widths[idx])
    doc.add_paragraph()
    return table


def add_info_box(doc, title, lines, fill="EFF6FF"):
    table = doc.add_table(rows=1, cols=1)
    set_table_style(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(15, 76, 129)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("• " + line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9.8)
    doc.add_paragraph()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    r = title.add_run("风电场数据对接与项目阶段成果说明")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(15, 76, 129)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("华锐 SL1500 型双馈式风机齿轮箱智能健康管理系统")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(71, 85, 105)

    meta = doc.add_table(rows=4, cols=2)
    set_table_style(meta)
    meta_rows = [
        ("文档用途", "风电场数据对接沟通、毕设阶段成果展示、研究方向说明"),
        ("系统对象", "SL1500 型双馈式风电机组齿轮箱及关键传动链部件"),
        ("当前阶段", "工程原型已完成，正在准备真实 SCADA/PLC/CMS 数据接入与现场验证"),
        ("生成日期", "2026 年 5 月 7 日"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        set_cell_shading(meta.cell(i, 0), "E2E8F0")
        set_cell_text(meta.cell(i, 0), k, bold=True)
        set_cell_text(meta.cell(i, 1), v)
    doc.add_page_break()

    add_heading(doc, "一、项目背景与对接目标", 1)
    add_paragraph(
        doc,
        "本项目面向风电场实际运维场景，围绕 SL1500 型双馈式风机齿轮箱健康管理展开。"
        "齿轮箱承担风轮低速大扭矩到发电机高速轴的传动转换，长期受交变载荷、风速波动、润滑状态变化和环境温度影响，"
        "容易出现齿面磨损、点蚀剥落、轴承损伤、油温异常和润滑油污染等问题。"
    )
    add_paragraph(
        doc,
        "当前系统已完成工程原型开发，下一步需要与风电场 SCADA、PLC、CMS 或历史运行数据库对接，"
        "获取真实机组运行数据、故障记录和检修记录，用于验证模型效果、优化健康评分和形成可用于现场运维的诊断闭环。"
    )
    add_info_box(
        doc,
        "对接目标",
        [
            "建立稳定的数据采集链路，实现机组运行状态、齿轮箱油温、振动、功率、风速、转速、报警等数据同步。",
            "将真实数据接入故障诊断、健康评估、RUL 估计和健康报告模块，替换当前演示性模拟数据。",
            "形成从异常识别、诊断建议、复检周期、检修工单到复测验收的闭环运维流程。",
        ],
    )

    add_heading(doc, "二、风电场数据对接方案", 1)
    add_paragraph(
        doc,
        "建议采用“只读接入、分级授权、边缘缓存、统一清洗”的方式对接风电场数据。"
        "系统优先通过风场现有 SCADA 网关或历史数据库获取数据，不直接改写 PLC 控制逻辑；若需下发控制命令，"
        "应单独经过权限审批、操作确认和安全联锁校验。"
    )
    add_table(
        doc,
        ["数据源", "建议接口", "主要数据内容", "用途"],
        [
            ("SCADA", "OPC UA / Modbus TCP / 历史库 API", "风速、功率、转速、变桨角、偏航角、发电量、机组状态", "实时监测、趋势分析、工况归一化"),
            ("PLC", "只读寄存器映射 / OPC UA", "运行链路、控制状态、停机原因、保护动作、关键开关量", "状态码解析、停机原因判断"),
            ("CMS", "文件导入 / API / 数据库同步", "高速轴承振动、包络谱、RMS、峭度、峰值因子", "齿轮箱与轴承故障识别"),
            ("油液监测", "CSV / API / 在线油液传感器", "NAS 等级、水分、铁磁颗粒、黏度、油温", "润滑状态评估与故障辅助判断"),
            ("检修台账", "Excel / 工单系统 API", "故障时间、处理过程、备件、更换记录、复测结果", "模型标签、闭环验证、论文案例支撑"),
        ],
        widths=[2.4, 3.4, 5.2, 4.0],
    )

    add_heading(doc, "三、建议采集字段与数据粒度", 1)
    add_table(
        doc,
        ["类别", "关键字段", "建议频率", "说明"],
        [
            ("机组运行", "机组编号、运行状态、有功功率、风速、发电机转速、变桨角、偏航角、转矩", "1 s - 10 s", "用于实时监测、功率曲线与工况修正"),
            ("齿轮箱温度", "齿轮箱油温、高速轴承温度、发电机温度、柜内温度、环境温度", "1 s - 10 s", "用于 M-IALO-SVR 油温残差预测和过温预警"),
            ("振动特征", "RMS、峭度、峰值、峰值因子、包络谱特征频率", "CMS 原始高频；特征 1 min - 10 min", "用于齿轮、轴承、轴系不对中等故障诊断"),
            ("报警事件", "报警码、报警名称、严重程度、触发时间、恢复时间、停机原因", "事件触发", "用于故障历史趋势和严重程度统计"),
            ("检修记录", "工单号、故障部位、处理措施、备件、验收结论、复检结果", "人工录入 / 工单同步", "用于诊断结果闭环和模型验证标签"),
        ],
        widths=[2.2, 6.0, 2.6, 4.2],
    )

    add_heading(doc, "四、当前项目成果", 1)
    add_paragraph(
        doc,
        "目前项目已形成一套可演示、可扩展的风机齿轮箱智能健康管理系统原型，后端采用 Flask，前端采用 HTML、CSS、JavaScript、ECharts 和 Three.js/SVG 可视化，"
        "数据库采用 SQLite，并预留真实数据接入和模型替换位置。"
    )
    add_table(
        doc,
        ["模块", "已实现成果", "工程价值"],
        [
            ("登录与权限", "管理员登录、用户管理、权限配置、系统通知已读处理", "满足基本系统使用和安全管理需求"),
            ("风场总览", "多台风机卡片、健康 99 显示、小风机叶片旋转、环境/列表视图切换、添加机组", "展示场站运行态势和机组分布状态"),
            ("单机详情", "单机切换、实时指标卡、温度寿命、控制按钮、机舱剖视模型、集群对标", "支撑单台机组精细化诊断和运维查看"),
            ("实时监测", "机组选择、SCADA/CMS 指标刷新、振动波形、数据质量与传感器状态", "模拟真实在线监测中心工作流"),
            ("故障诊断", "基于特征的异常识别、故障类型分布、严重程度统计、诊断机组选择", "支撑齿轮箱异常检测和检修建议生成"),
            ("故障数据管理", "历史故障记录、机组编号、严重程度、处理状态、导出能力", "形成故障样本积累和论文验证数据基础"),
            ("健康评估报告", "关键指标、健康评分、RUL、复检周期、导出报告", "服务运维决策与阶段性汇报"),
            ("AI 专家问答", "内置风电运维知识、可接入大模型 API、回答结合当前工况", "辅助非专家人员快速理解诊断结论"),
        ],
        widths=[2.5, 7.0, 5.0],
    )

    add_heading(doc, "五、具体研究方向", 1)
    add_table(
        doc,
        ["研究方向", "研究内容", "预期成果"],
        [
            ("多源数据融合", "融合 SCADA、CMS、油液、报警和检修台账，统一时间戳、机组编号和字段口径", "形成面向齿轮箱健康评估的数据底座"),
            ("油温残差预测", "利用 M-IALO-SVR 建立油温预测模型，以实际油温与预测油温残差识别异常热状态", "获得过温早期预警指标和残差阈值策略"),
            ("齿轮箱故障诊断", "提取振动时域、频域、包络谱和工况参数特征，识别齿轮磨损、点蚀、断齿、轴承过温等故障", "输出故障类型、置信度、严重程度和维修建议"),
            ("健康评分与 RUL", "结合油温、振动、油液、报警和历史趋势计算健康评分，按 20 年设计寿命和 RUL 推算复检周期", "形成可解释的健康评估报告"),
            ("运维闭环", "将诊断结果转为检修建议、工单优先级、复检周期和验收记录", "实现从诊断到处置的闭环管理"),
        ],
        widths=[3.0, 7.0, 4.5],
    )

    add_heading(doc, "六、前期准备工作", 1)
    add_bullets(
        doc,
        [
            "完成风机齿轮箱常见故障机理梳理，重点关注齿面磨损、点蚀剥落、断齿、轴承磨损、轴系不对中、冷却系统故障和润滑油污染。",
            "完成系统原型搭建，形成后端接口、前端页面、图表展示、机组详情和健康报告的基础框架。",
            "完成数据字段设计，已覆盖油温、振动 RMS、油液 NAS、功率、风速、转速、变桨角、报警状态和 RUL 等关键指标。",
            "完成模拟数据流和接口联调，能够支持实时刷新、机组切换、故障趋势展示和不同机组差异化分析。",
            "完成大模型运维问答接入准备，能够将知识条目、当前工况和运维建议结合起来生成回答。",
        ],
    )

    add_heading(doc, "七、对研究方向的深入理解", 1)
    add_paragraph(
        doc,
        "风机齿轮箱故障诊断不能只依赖单一阈值。油温升高可能来自负载升高、环境温度变化、冷却系统效率下降、润滑油劣化或轴承摩擦异常；"
        "振动增大也可能与风速波动、转速变化、传动链不对中、齿轮啮合冲击或轴承局部损伤有关。因此，研究重点应放在工况修正、多源特征融合和趋势变化识别上。"
    )
    add_paragraph(
        doc,
        "M-IALO-SVR 在本课题中的价值不在于简单预测一个温度值，而是建立“正常工况下油温应当是多少”的参考模型。"
        "当实测油温持续高于模型预测值，且振动、油液或报警信号同步出现异常时，系统可以更早识别潜在故障，减少单纯阈值判断造成的漏报和误报。"
    )
    add_paragraph(
        doc,
        "真实风场对接后，研究工作应重点解决三类问题：第一，数据质量问题，包括缺失、延迟、异常跳变和不同系统时间不一致；第二，标签稀缺问题，即真实故障样本少、故障边界不清；"
        "第三，工程可解释问题，即诊断结果必须能说明依据、风险等级、建议复检周期和现场处理动作。"
    )

    add_heading(doc, "八、风电场现场对接实施步骤", 1)
    add_table(
        doc,
        ["阶段", "主要工作", "交付物"],
        [
            ("1. 对接确认", "确认风场系统结构、可开放接口、机组清单、数据字段、网络边界和权限要求", "数据对接清单、字段字典、接口授权说明"),
            ("2. 数据试采", "选择 1-3 台机组进行只读试采，验证时间戳、采样频率、单位和数据质量", "试采样本、数据质量报告"),
            ("3. 映射清洗", "建立字段映射、单位换算、异常值处理、缺失补齐和缓存机制", "标准化数据表、清洗规则"),
            ("4. 模型验证", "基于真实数据验证油温残差、故障类型识别、健康评分和 RUL 估计", "模型评估结果、阈值建议"),
            ("5. 系统联调", "将真实数据接入现有页面，联调总览、单机、实时监测、诊断、报告和工单流程", "现场演示版本、联调记录"),
            ("6. 运维闭环", "把诊断结论转为检修建议和复检周期，并记录处理反馈", "闭环案例、论文验证材料"),
        ],
        widths=[2.2, 8.0, 4.2],
    )

    add_heading(doc, "九、需要风电场配合提供的资料", 1)
    add_bullets(
        doc,
        [
            "机组清单：机组编号、型号、投运日期、容量、所属线路或区域、当前运行状态。",
            "数据接口与信息安全资料：包括 SCADA、PLC、CMS 等系统的数据访问方式、通信协议、点表信息或字段字典等；涉及 IP 地址、端口、账号权限、访问密钥等敏感信息时，应按风电场网络安全要求脱敏或仅在现场受控环境下查看，系统原则上以只读方式接入运行数据，不直接修改 PLC 控制逻辑。",
            "历史数据：至少 3-12 个月运行数据，包含油温、风速、功率、转速、振动、报警和停机事件。",
            "故障与检修记录：齿轮箱、轴承、冷却系统、润滑系统相关故障案例及处理结果。",
            "现场规则：网络安全要求、数据脱敏要求、控制命令权限边界和现场验收流程。",
        ],
    )

    add_heading(doc, "十、阶段计划与预期成果", 1)
    add_table(
        doc,
        ["时间安排", "工作重点", "成果形式"],
        [
            ("第 1 阶段", "完成风场数据字段确认、接口沟通和试采方案设计", "字段表、接口方案、试采计划"),
            ("第 2 阶段", "完成真实数据导入、清洗、可视化页面适配", "真实数据演示页面、数据质量报告"),
            ("第 3 阶段", "完成 M-IALO-SVR 油温残差模型训练和故障样本验证", "模型结果、误差指标、残差阈值"),
            ("第 4 阶段", "完成健康评分、RUL、复检周期和故障诊断报告优化", "健康评估报告、诊断案例"),
            ("第 5 阶段", "完成论文撰写、系统测试和答辩展示材料整理", "毕业论文、演示系统、答辩 PPT"),
        ],
        widths=[2.4, 8.0, 4.0],
    )

    add_heading(doc, "十一、结论", 1)
    add_paragraph(
        doc,
        "本项目已经具备风电场智能健康管理系统的基础形态，能够展示风场总览、单机详情、实时监测、故障诊断、故障数据管理、健康报告和智能问答等核心功能。"
        "后续工作的关键是接入真实风场数据，通过真实运行工况和检修记录验证模型效果，使系统从毕设演示原型进一步提升为具有工程应用价值的智能运维辅助平台。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
