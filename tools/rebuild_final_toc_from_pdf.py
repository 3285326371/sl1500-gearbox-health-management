from __future__ import annotations

import importlib.util
import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path.cwd()
DOCX = next(p for p in (Path.home() / "Desktop").glob("*正文模板.docx") if not p.name.startswith("~$"))
PDF = max((ROOT / "docs" / "qa_template_inplace").glob("*重构*检.pdf"), key=lambda p: p.stat().st_mtime)

spec = importlib.util.spec_from_file_location("restructure", ROOT / "tools" / "restructure_thesis_40pages.py")
mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(mod)
TOC_LINES = mod.TOC_LINES


def norm(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def text(p):
    return p.text.strip().replace("\n", "")


def set_font(run, size, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def delete_between(doc, start_p, end_p):
    body = doc._body._element
    children = list(body)
    start = children.index(start_p._element) + 1
    end = children.index(end_p._element)
    for el in children[start:end]:
        body.remove(el)


def main():
    doc = Document(DOCX)
    toc_title = next(p for p in doc.paragraphs if norm(p.text) == "目录")
    seen = False
    first_chapter = None
    for p in doc.paragraphs:
        if p._element is toc_title._element:
            seen = True
            continue
        if seen and text(p).startswith("第1章") and "\t" not in p.text:
            first_chapter = p
            break
    if first_chapter is None:
        raise RuntimeError("Cannot locate first chapter after TOC")

    pdf = fitz.open(PDF)
    page_texts = [norm(page.get_text("text")) for page in pdf]
    chapter_phys = next(
        i for i, t in enumerate(page_texts, start=1)
        if norm("第1章 绪论") in t and norm("风力发电是新能源") in t
    )
    page_map = {}
    for title, _, _ in TOC_LINES:
        nt = norm(title)
        found = None
        for i in range(chapter_phys, len(page_texts) + 1):
            if nt in page_texts[i - 1]:
                found = i
                break
        page_map[title] = str(found - chapter_phys + 1) if found else ""

    delete_between(doc, toc_title, first_chapter)
    for title, _, level in TOC_LINES:
        p = first_chapter.insert_paragraph_before("")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24 * level)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        set_font(r, 12 if level == 0 else 10.5, bold=(level == 0))
        p.add_run("\t")
        r2 = p.add_run(page_map.get(title, ""))
        set_font(r2, 12 if level == 0 else 10.5, bold=(level == 0))
    toc_title.paragraph_format.alignment = 1
    toc_title.paragraph_format.first_line_indent = Pt(0)
    toc_title.paragraph_format.space_after = Pt(12)
    for r in toc_title.runs:
        r.font.name = "Times New Roman"
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(16)
        r.bold = True

    doc.save(DOCX)
    print(f"updated={DOCX}")
    print(f"chapter_phys={chapter_phys}")
    for title, _, _ in TOC_LINES:
        print(f"{title}\t{page_map.get(title, '')}")


if __name__ == "__main__":
    main()
