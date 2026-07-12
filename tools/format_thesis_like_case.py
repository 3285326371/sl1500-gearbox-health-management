from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DESKTOP = Path.home() / "Desktop"
TARGET = next(DESKTOP.glob("*正文模板.docx"))
OUT = TARGET


def set_run_font(run, size=12, east="宋体", west="Times New Roman", bold=None):
    run.font.name = west
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_para_font(p, size=12, east="宋体", west="Times New Roman", bold=None):
    for r in p.runs:
        set_run_font(r, size=size, east=east, west=west, bold=bold)


def normalize_para(p, first_indent=True, size=12):
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Pt(24) if first_indent else Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_para_font(p, size=size)


def format_title(p, level=1):
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(21)
    if level == 1:
        fmt.space_before = Pt(40)
        fmt.space_after = Pt(20)
        set_para_font(p, size=15, east="黑体", west="Times New Roman", bold=False)
    elif level == 2:
        fmt.space_before = Pt(24)
        fmt.space_after = Pt(6)
        set_para_font(p, size=14, east="黑体", west="Times New Roman", bold=False)
    else:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
        set_para_font(p, size=12, east="黑体", west="Times New Roman", bold=False)


def shade_para(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def clear_shade_para(p):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is not None:
        pPr.remove(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill="D9EAF7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def is_code_table(table):
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return any(x in text for x in ("def ", "async function", "function ", "return ", "const ", "@data_bp.route", "create_app"))


def format_table(table):
    code = is_code_table(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.autofit = True
    except Exception:
        pass
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if ri == 0 and not code:
                shade_cell(cell, "D9EAF7")
            if code:
                shade_cell(cell, "F2F2F2")
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE if code else WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    if code:
                        set_run_font(r, size=9, east="Consolas", west="Consolas", bold=False)
                    else:
                        set_run_font(r, size=10.5, east="宋体", west="Times New Roman", bold=(True if ri == 0 else False))


def para_has_drawing(p):
    return bool(p._element.xpath(".//w:drawing"))


def update_caption_number(text, kind_counts):
    m = re.match(r"^(图|表)(\d+)\.(\d+)\s*(.*)$", text.strip())
    if not m:
        return text
    kind, chap, _, rest = m.groups()
    key = (kind, chap)
    kind_counts[key] = kind_counts.get(key, 0) + 1
    return f"{kind}{chap}.{kind_counts[key]} {rest}".strip()


def main():
    backup = TARGET.with_name(f"{TARGET.stem}_格式修改前备份_{datetime.now():%Y%m%d_%H%M%S}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)

    doc = Document(TARGET)

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.2)

    toc_mode = False
    in_body = False
    kind_counts = {}
    caption_re = re.compile(r"^(图|表)\d+\.\d+")
    h1_re = re.compile(r"^第\d+章\s*.+")
    h2_re = re.compile(r"^\d+\.\d+\s+.+")
    h3_re = re.compile(r"^\d+\.\d+\.\d+\s+.+")

    for p in doc.paragraphs:
        text = p.text.strip()
        clear_shade_para(p)
        if not text:
            continue

        if text == "目  录":
            toc_mode = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            set_para_font(p, size=16, east="黑体", west="Times New Roman", bold=True)
            continue

        if toc_mode and "\t" in text:
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.35 if re.match(r"^\d+\.", text) else 0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.tab_stops.clear_all()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), 2, 1)
            set_para_font(p, size=10.5, east="宋体", west="Times New Roman", bold=False)
            continue

        if h1_re.match(text):
            toc_mode = False
            in_body = True
            p.style = doc.styles["Heading 1"]
            format_title(p, level=1)
            continue

        if text in ("摘  要", "ABSTRACT", "参考文献", "致谢") or text.startswith("附录"):
            p.style = doc.styles["Heading 1"]
            format_title(p, level=1)
            continue

        if caption_re.match(text):
            new_text = update_caption_number(text, kind_counts)
            if new_text != text:
                p.clear()
                p.add_run(new_text)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(18)
            set_para_font(p, size=10.5, east="宋体", west="Times New Roman", bold=False)
            continue

        if h3_re.match(text):
            p.style = doc.styles["Heading 3"]
            format_title(p, level=3)
            continue

        if h2_re.match(text):
            p.style = doc.styles["Heading 2"]
            format_title(p, level=2)
            continue

        if para_has_drawing(p):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            continue

        if text.startswith(("关键词：", "Key words:")):
            normalize_para(p, first_indent=False, size=12)
            continue

        if text.startswith("{") or text.startswith("["):
            normalize_para(p, first_indent=False, size=9)
            for r in p.runs:
                set_run_font(r, size=9, east="Consolas", west="Consolas", bold=False)
            shade_para(p)
            continue

        if in_body or p.paragraph_format.first_line_indent:
            normalize_para(p, first_indent=True, size=12)

    for shape in doc.inline_shapes:
        max_width = Inches(5.9)
        if shape.width and shape.width > max_width:
            ratio = shape.height / shape.width
            shape.width = max_width
            shape.height = int(max_width * ratio)

    for table in doc.tables:
        format_table(table)

    doc.save(OUT)
    print(f"saved={OUT}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
