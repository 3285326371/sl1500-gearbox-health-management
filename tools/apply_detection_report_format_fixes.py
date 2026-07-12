from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCX = (
    Path.home()
    / "Desktop"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f_\u68c0\u6d4b\u62a5\u544a"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
)


def set_run_font(run, size=12, east="宋体", ascii_font="Times New Roman", bold=None, color=None):
    run.font.name = ascii_font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_char_indent(p):
    ind = p._p.get_or_add_pPr().find(qn("w:ind"))
    if ind is not None:
        for attr in ["firstLineChars", "hangingChars"]:
            ind.attrib.pop(qn(f"w:{attr}"), None)


def set_first_line_chars(p, chars=200):
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = None
    p.paragraph_format.right_indent = None
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), str(chars))


def set_left_chars(p, chars=0):
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.attrib.pop(qn("w:left"), None)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:firstLineChars"), None)
    if chars:
        ind.set(qn("w:leftChars"), str(chars))
    else:
        ind.attrib.pop(qn("w:leftChars"), None)


def set_exact_line(p, points):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(points)


def fmt_para(p, size=12, east="宋体", ascii_font="Times New Roman", bold=None):
    for run in p.runs:
        set_run_font(run, size=size, east=east, ascii_font=ascii_font, bold=bold)


def is_body_text(p):
    text = p.text.strip()
    if not text:
        return False
    if p.style.name.startswith("Heading"):
        return False
    if text in {"摘  要", "ABSTRACT", "目  录", "参考文献", "致谢"}:
        return False
    if re.match(r"^(图|表)\s*\d+(\.\d+)?", text):
        return False
    if re.match(r"^代码清单\d+\.\d+", text) or text == "核心代码如下：":
        return False
    if re.match(r"^\[\d+\]", text):
        return False
    if text.startswith(("关键词", "Key words", "Key words:", "Key words：")):
        return False
    if "\t" in text and (text.startswith("第") or re.match(r"^\d+(\.\d+)+", text) or text in {"参考文献", "致谢"}):
        return False
    # Common code lines are intentionally kept unindented.
    code_marks = ["def ", "return ", "const ", "async ", "await ", "if ", "for ", "});", "}", "{", "@", "document.", "localStorage", "fetch("]
    if any(text.startswith(mark) for mark in code_marks) or text.startswith(("\"", "'")):
        return False
    return True


def format_front_matter(doc):
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if re.match(r"^20\d{2}\s*年", text):
            fmt_para(p, 16, "宋体", "Times New Roman")
        if text == "摘  要":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(20)
            set_exact_line(p, 20)
            fmt_para(p, 15, "黑体", "Arial", bold=False)
        elif text == "ABSTRACT":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(20)
            set_exact_line(p, 20)
            fmt_para(p, 15, "黑体", "Arial", bold=False)
        elif text.startswith("关键词"):
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            clear_char_indent(p)
            set_exact_line(p, 20)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Keep only five keywords to satisfy the template reminder.
            if "；" in text:
                head, kws = text.split("：", 1) if "：" in text else ("关键词", text.replace("关键词", "", 1))
                p.text = "关键词：" + "；".join([k.strip() for k in kws.split("；")[:5]])
            for run in p.runs:
                set_run_font(run, 12, "黑体", "Times New Roman", bold=True)
        elif text.startswith("Key words"):
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            clear_char_indent(p)
            set_exact_line(p, 20)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if ":" in text:
                head, kws = text.split(":", 1)
                p.text = "Key words: " + "; ".join([k.strip() for k in kws.split(";")[:5]])
            for run in p.runs:
                set_run_font(run, 12, "黑体", "Times New Roman", bold=True)

    # Abstract body paragraphs.
    abstract = False
    english = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "摘  要":
            abstract = True
            english = False
            continue
        if text == "ABSTRACT":
            abstract = True
            english = True
            continue
        if text == "目  录":
            abstract = False
        if abstract and text and not text.startswith(("关键词", "Key words")):
            set_first_line_chars(p, 200)
            set_exact_line(p, 20)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            fmt_para(p, 12, "宋体", "Times New Roman")


def format_toc(doc):
    in_toc = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "目  录":
            in_toc = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(20)
            set_exact_line(p, 20)
            fmt_para(p, 15, "黑体", "Arial", bold=False)
            continue
        if in_toc and text == "第1章 绪论":
            in_toc = False
        if in_toc and text:
            p.paragraph_format.space_after = Pt(0)
            set_exact_line(p, 20)
            if text.startswith("第") or text.startswith("参考文献") or text.startswith("致谢"):
                p.paragraph_format.space_before = Pt(6)
                set_left_chars(p, 0)
                fmt_para(p, 12, "黑体", "Arial", bold=False)
            elif re.match(r"^\d+\.\d+\.\d+", text):
                p.paragraph_format.space_before = Pt(0)
                set_left_chars(p, 400)
                fmt_para(p, 12, "宋体", "Times New Roman")
            elif re.match(r"^\d+\.\d+", text):
                p.paragraph_format.space_before = Pt(0)
                set_left_chars(p, 200)
                fmt_para(p, 12, "宋体", "Times New Roman")


def format_body(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name
        if style == "Heading 1" and text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(20)
            p.paragraph_format.keep_with_next = True
            set_exact_line(p, 20)
            fmt_para(p, 15, "黑体", "Arial", bold=True)
        elif style == "Heading 2" and text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            set_exact_line(p, 20)
            fmt_para(p, 14, "黑体", "Arial", bold=True)
        elif style == "Heading 3" and text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            set_exact_line(p, 20)
            fmt_para(p, 12, "黑体", "Arial", bold=True)
        elif re.match(r"^(图|表)\s*\d+(\.\d+)?", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            clear_char_indent(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt_para(p, 10.5, "黑体", "Arial", bold=False)
        elif re.match(r"^代码清单\d+\.\d+", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            clear_char_indent(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt_para(p, 10.5, "黑体", "Arial", bold=False)
        elif re.match(r"^\[\d+\]", text):
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            clear_char_indent(p)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            set_exact_line(p, 16)
            fmt_para(p, 10.5, "宋体", "Times New Roman")
        elif is_body_text(p):
            set_first_line_chars(p, 200)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_exact_line(p, 20)
            fmt_para(p, 12, "宋体", "Times New Roman")


def format_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = None
                    clear_char_indent(p)
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    fmt_para(p, 10.5, "宋体", "Times New Roman")


def fix_headers_and_margins(doc):
    for section in doc.sections:
        section.footer_distance = Cm(1.75)
    if len(doc.sections) > 1:
        for header in (doc.sections[1].header, doc.sections[1].first_page_header, doc.sections[1].even_page_header):
            for p in header.paragraphs:
                if p.text.strip():
                    p.text = "河北建筑工程学院学士学位论文"
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fmt_para(p, 9, "宋体", "Times New Roman")


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_检测批注批量修复前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)
    fix_headers_and_margins(doc)
    format_front_matter(doc)
    format_toc(doc)
    format_body(doc)
    format_tables(doc)
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
