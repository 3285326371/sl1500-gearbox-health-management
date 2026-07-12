from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCX = next((Path.home() / "Desktop").glob("*正文模板.docx"))
PDF = max((Path.cwd() / "docs" / "qa_template_inplace").glob("*分节正确初检.pdf"), key=lambda p: p.stat().st_mtime)


def norm(text: str) -> str:
    return re.sub(r"\s+", "", text or "").replace("\u3000", "")


def set_font(run, size=10.5, east="宋体", west="Times New Roman", bold=False):
    run.font.name = west
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def main():
    doc = Document(DOCX)
    toc_title_idx = next(i for i, p in enumerate(doc.paragraphs) if norm(p.text) == norm("目  录"))
    toc_entries = []
    for i in range(toc_title_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if "\t" not in text:
            break
        title = text.split("\t", 1)[0].strip()
        toc_entries.append((i, title))

    pdf = fitz.open(PDF)
    page_texts = [norm(page.get_text("text")) for page in pdf]

    first_chapter = next(title for _, title in toc_entries if norm(title).startswith("第1章"))
    first_section = next(title for _, title in toc_entries if norm(title).startswith("1.1"))
    chapter_phys = None
    for pno, text in enumerate(page_texts, start=1):
        if norm(first_chapter) in text and norm(first_section) in text and text.count(".") < 80:
            chapter_phys = pno
            break
    if chapter_phys is None:
        raise RuntimeError("Cannot locate first chapter page in rendered PDF")

    page_map = {}
    for _, title in toc_entries:
        nt = norm(title)
        if title == "摘要":
            page_map[title] = "I"
            continue
        if title == "ABSTRACT":
            page_map[title] = "II"
            continue
        found = None
        for pno in range(chapter_phys, len(page_texts) + 1):
            if nt in page_texts[pno - 1]:
                found = pno
                break
        if found is not None:
            page_map[title] = str(found - chapter_phys + 1)
        else:
            page_map[title] = ""

    for idx, title in toc_entries:
        p = doc.paragraphs[idx]
        level = 1 if re.match(r"^\d+\.", title) else 0
        p.clear()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Inches(0.35 if level else 0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        set_font(r)
        p.add_run("\t")
        r2 = p.add_run(page_map.get(title, ""))
        set_font(r2)

    doc.save(DOCX)
    print(f"updated={DOCX}")
    print(f"chapter_phys={chapter_phys}")
    for title in page_map:
        print(f"{title}\t{page_map[title]}")


if __name__ == "__main__":
    main()
