from __future__ import annotations

from pathlib import Path
import re

import fitz
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path(r"C:\Users\隐涅\Desktop\毕业设计正文模板_检测报告\毕业设计正文模板.docx")
PDF = ROOT / "docs" / "qa_detection_report" / "final_cover_ai_toc_checked.pdf"


def compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def set_run_font(run, name: str, size: int = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def format_toc_paragraph(paragraph, level: int):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Cm({1: 0.0, 2: 0.74, 3: 1.48}.get(level, 0.0))
    fmt.right_indent = Pt(0)
    fmt.space_before = Pt(6 if level == 1 else 0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(20)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def get_heading_level(paragraph) -> int | None:
    name = paragraph.style.name
    if name in ("Heading 1", "标题 1"):
        return 1
    if name in ("Heading 2", "标题 2"):
        return 2
    if name in ("Heading 3", "标题 3"):
        return 3
    return None


def collect_doc_headings(doc: Document):
    headings = []
    started = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("第1章"):
            started = True
        if not started:
            continue
        level = get_heading_level(paragraph)
        if level is None:
            continue
        if text in {"摘要", "ABSTRACT"}:
            continue
        headings.append((level, text))
    return headings


def find_body_start_pdf_page(pdf: fitz.Document) -> int:
    for i, page in enumerate(pdf):
        if "第1章绪论" in compact(page.get_text("text")):
            # skip the table of contents hit
            if i >= 5:
                return i
    raise RuntimeError("Cannot locate body start page in exported PDF")


def page_for_heading(pdf: fitz.Document, heading: str, body_start: int) -> int:
    target = compact(heading)
    for i in range(body_start, len(pdf)):
        text = compact(pdf[i].get_text("text"))
        if target in text:
            return i - body_start + 1
    # Fallback for headings whose extracted text has odd spacing.
    short = target.replace(" ", "")
    for i in range(body_start, len(pdf)):
        text = compact(pdf[i].get_text("text"))
        if short[: min(8, len(short))] in text:
            return i - body_start + 1
    raise RuntimeError(f"Cannot locate heading in PDF: {heading}")


def remove_between(start_para, end_para):
    node = start_para._p.getnext()
    while node is not None and node is not end_para._p:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt


def main():
    doc = Document(DOCX)
    pdf = fitz.open(PDF)
    body_start = find_body_start_pdf_page(pdf)
    headings = collect_doc_headings(doc)
    toc_rows = []
    for level, heading in headings:
        toc_rows.append((level, heading, page_for_heading(pdf, heading, body_start)))

    toc_title = next(p for p in doc.paragraphs if compact(p.text) == "目录")
    first_body = next(
        p
        for p in doc.paragraphs
        if p.text.strip().startswith("第1章") and get_heading_level(p) == 1
    )
    remove_between(toc_title, first_body)

    anchor = toc_title
    for level, heading, page in reversed(toc_rows):
        paragraph = insert_paragraph_after(anchor)
        paragraph.style = doc.styles["Normal"]
        format_toc_paragraph(paragraph, level)
        run = paragraph.add_run(heading)
        set_run_font(run, "黑体" if level == 1 else "宋体", 12, level == 1)
        run = paragraph.add_run(f"\t{page}")
        set_run_font(run, "Times New Roman", 12, level == 1)

    doc.save(DOCX)
    print(f"Rebuilt TOC with {len(toc_rows)} entries; body PDF page={body_start + 1}")


if __name__ == "__main__":
    main()
