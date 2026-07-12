from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_附录C前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_附录C.docx"


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def para(doc, text, first=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24) if first else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        set_font(r)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    for line_no, line in enumerate(text.strip().splitlines()):
        if line_no:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
    return p


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    if any("附录C 系统部署与运行说明" in p.text for p in doc.paragraphs):
        print("already exists")
    else:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = doc.add_paragraph("附录C 系统部署与运行说明")
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            set_font(r, east_asia="黑体", ascii_font="Arial", size=16, bold=True)
        para(doc, "本系统采用本地Flask服务和静态前端页面运行，适合毕业设计演示、功能测试和答辩展示。系统启动前应确认Python虚拟环境、依赖包、数据库目录和前端文件均存在。项目根目录下的backend目录保存后端程序，frontend目录保存页面文件，instance目录保存SQLite数据库。")
        para(doc, "启动系统时，进入项目根目录并执行后端入口文件。系统默认监听127.0.0.1:5000端口，启动后可通过浏览器访问首页。默认管理员账号用于演示登录，登录后可以查看风场总览、故障诊断、故障数据、健康报告、AI问答和参数设置等模块。")
        code(doc, r"""
cd "D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1"
.\.venv\Scripts\python.exe backend\app.py
http://127.0.0.1:5000
默认账号：admin
默认密码：admin
""")
        para(doc, "系统健康检查接口为/health，用于确认后端服务和数据库路径是否正常。若返回status为ok，说明Flask服务已启动并能够访问数据库目录。系统页面中使用的实时数据、诊断接口、报告接口和问答接口均通过/api前缀访问。")
        code(doc, r"""
GET  /health
GET  /api/data/status
POST /api/data/diagnosis
GET  /api/windfarm/overview
GET  /api/report/generate
POST /api/qa/ask
""")
        para(doc, "答辩演示时建议先展示风场总览页面，说明系统能够对多台机组进行健康状态监测；随后进入单机详情页面，说明油温、振动、油液和RUL等指标；再进入故障诊断页面执行一次典型故障诊断；最后展示健康报告、AI问答和参数设置页面，说明系统已经形成从监测、诊断、解释到维护建议的闭环。")
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
