from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")
DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
PDF = ROOT / "docs" / "qa_template_inplace" / "frontend_hmi_table_check.pdf"


def norm(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def paragraph_after(paragraph, text="", style=None):
    p = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(p._p)
    return p


def set_run_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def format_toc_paragraph(p, level: int, is_chapter: bool):
    p.style = "Normal"
    p.paragraph_format.left_indent = Cm({1: 0, 2: 0.74, 3: 1.48}.get(level, 0))
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.alignment = 0
    for run in p.runs:
        set_run_font(run, 12, bold=is_chapter)


def remove_duplicate_empty_heading(doc: Document):
    seen = set()
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if p.style.name.startswith("Heading") and text:
            key = (p.style.name, text)
            if key in seen and text == "6.4 结果评价":
                p._element.getparent().remove(p._element)
                continue
            seen.add(key)


def collect_body_headings(doc: Document):
    started = False
    result = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "第1章 绪论" and p.style.name == "Heading 1":
            started = True
        if not started or not text:
            continue
        if p.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            if text.startswith("摘") or text == "ABSTRACT":
                continue
            level = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}[p.style.name]
            result.append((text, level))
    return result


def page_map(headings):
    pdf = fitz.open(PDF)
    page_texts = [norm(page.get_text("text")) for page in pdf]
    chapter_start = None
    # Front matter also contains the static TOC text. The real body starts after
    # cover/declaration/abstract/catalog pages, so skip the first several pages.
    for idx in range(6, len(page_texts)):
        text = page_texts[idx]
        if norm("第1章 绪论") in text:
            chapter_start = idx
            break
    if chapter_start is None:
        raise RuntimeError("Cannot locate first chapter in PDF")

    mapping = {}
    for heading, _level in headings:
        n = norm(heading)
        found = None
        for idx in range(chapter_start, len(page_texts)):
            if n in page_texts[idx]:
                found = idx - chapter_start + 1
                break
        if found is None:
            found = ""
        mapping[heading] = found
    return mapping


def rebuild_toc(doc: Document, headings, mapping):
    toc_title = None
    for p in doc.paragraphs:
        if p.text.strip() == "目  录":
            toc_title = p
            break
    if toc_title is None:
        raise RuntimeError("TOC title not found")

    body_start = None
    for p in doc.paragraphs:
        if p.text.strip() == "第1章 绪论" and p.style.name == "Heading 1":
            body_start = p
            break
    if body_start is None:
        raise RuntimeError("Body start not found")

    body = doc.element.body
    children = list(body)
    start = children.index(toc_title._p) + 1
    end = children.index(body_start._p)
    for el in children[start:end]:
        body.remove(el)

    anchor = toc_title
    for heading, level in headings:
        page = mapping.get(heading, "")
        p = paragraph_after(anchor, f"{heading}\t{page}")
        format_toc_paragraph(p, level, level == 1)
        anchor = p


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_目录重建前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)
    remove_duplicate_empty_heading(doc)
    headings = collect_body_headings(doc)
    mapping = page_map(headings)
    rebuild_toc(doc, headings, mapping)
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")
    print(f"toc_items={len(headings)}")


if __name__ == "__main__":
    main()
