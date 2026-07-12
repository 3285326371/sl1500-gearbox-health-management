from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def style_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(text)


def add_after(anchor, text):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    style_body(p)
    return p


def insert_once(doc, anchor_text, marker, text):
    if any(marker in p.text for p in doc.paragraphs):
        return
    add_after(find_para(doc, anchor_text), text)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_章节末留白微调前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)
    insert_once(
        doc,
        "Web系统开发技术的作用还体现在系统集成能力上。齿轮箱健康管理涉及实时数据、故障规则、历史记录、报告输出和用户交互等多个环节，如果缺少统一的平台载体，诊断结果难以被运维人员持续使用。通过 Web 系统将各类功能组织在同一界面中，可以提高系统演示的完整性，也便于后续扩展移动端访问、权限控制和远程运维功能。",
        "综上，本文后续系统设计将围绕数据流、业务流和界面流展开",
        "综上，本文后续系统设计将围绕数据流、业务流和界面流展开。数据流解决齿轮箱状态如何采集、处理和保存的问题，业务流解决故障如何诊断、记录和闭环的问题，界面流解决运维人员如何查看、理解和处理结果的问题。三类流程相互配合，构成系统总体设计和详细实现的基础。",
    )
    insert_once(
        doc,
        "因此，本文前端设计并非单纯追求页面美观，而是围绕运维人员的判断效率展开。总览页面强调快速发现问题，详情页面强调指标解释，诊断页面强调结论可信度，HMI 页面强调现场设备对象感，报告页面强调结果沉淀。不同页面承担不同任务，能够减少信息堆叠造成的阅读负担。",
        "第4章的总体设计为第5章实现提供了结构依据",
        "第4章的总体设计为第5章实现提供了结构依据。后续实现章节将按照数据采集、故障诊断、风场与单机页面、数据管理、问答辅助和健康报告等模块展开，分别说明关键代码、系统截图和运行效果，使系统设计与实际开发结果能够相互对应。",
    )
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
