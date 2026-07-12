from pathlib import Path
import re

from docx import Document


DOCX = Path(r"C:\Users\隐涅\Desktop\毕业设计正文模板_检测报告\毕业设计正文模板.docx")

doc = Document(DOCX)
patterns = [
    re.compile(r"\[[0-9,，\-—\s]+\]"),
    re.compile(r"文献\s*\[?\d+\]?"),
    re.compile(r"参考文献"),
]
generic = ["具有重要", "本文", "系统", "随着", "因此", "能够", "实现了", "提高"]

in_refs = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    if text == "参考文献":
        in_refs = True
    if any(pat.search(text) for pat in patterns):
        print("CIT", i, p.style.name, "refs" if in_refs else "body", text[:180])

print("--- sample generic paragraphs ---")
count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) < 60 or text == "参考文献":
        continue
    if any(g in text for g in generic):
        print("GEN", i, p.style.name, text[:220])
        count += 1
        if count >= 30:
            break
