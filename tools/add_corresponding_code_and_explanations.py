from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def set_run_font(run, size=12, name="宋体", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def style_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.left_indent = None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        set_run_font(run, 12)


def style_code_caption(p):
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, 12)


def style_code_line(p):
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.95)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def paragraph_after(anchor, text=""):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    return p


def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    raise RuntimeError("target paragraph not found")


def add_code_block(anchor, title, lines, explanation):
    p = paragraph_after(anchor, title)
    style_code_caption(p)
    anchor = p
    p = paragraph_after(anchor, "核心代码如下：")
    style_body(p)
    anchor = p
    for line in lines:
        p = paragraph_after(anchor, line)
        style_code_line(p)
        anchor = p
    p = paragraph_after(anchor, explanation)
    style_body(p)
    return p


def main():
    backup = DOCX.with_name(
        f"{DOCX.stem}_补充对应代码前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}"
    )
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)

    if any("代码清单4.1 主系统登录请求与角色状态保存" in p.text for p in doc.paragraphs):
        print("code blocks already inserted")
        return

    front_anchor = find_para(doc, lambda t: t.startswith("故障诊断页面提供场景选择和诊断按钮"))
    front_anchor = add_code_block(
        front_anchor,
        "代码清单4.1 主系统登录请求与角色状态保存",
        [
            'loginForm.addEventListener("submit", async (e) => {',
            "    e.preventDefault();",
            '    const response = await fetch("/api/auth/login", {',
            '        method: "POST",',
            '        headers: { "Content-Type": "application/json" },',
            "        body: JSON.stringify({",
            "            username: usernameInput.value,",
            "            password: passwordInput.value",
            "        })",
            "    });",
            "    const data = await response.json();",
            "    if (response.ok) {",
            '        localStorage.setItem("currentUser", JSON.stringify(data.user));',
            "        applyAuthority(data.user.role);",
            '        document.getElementById("login-overlay").style.display = "none";',
            '        document.getElementById("main-app").style.display = "flex";',
            "        showPostLoginDefault();",
            "    }",
            "});",
        ],
        "代码清单4.1位于frontend/js/main.js。登录表单提交后，前端向/api/auth/login发送账号和密码；认证成功后把用户信息写入localStorage，并根据角色调用applyAuthority控制页面权限，最后隐藏登录遮罩并进入默认业务页面。该逻辑对应论文中“登录后进入健康总览、按照角色显示身份和权限”的交互设计。",
    )
    add_code_block(
        front_anchor,
        "代码清单4.2 后端用户认证接口",
        [
            '@auth_bp.route("/login", methods=["POST"])',
            "def login():",
            "    data = request.json or {}",
            '    username = (data.get("username") or "").strip()',
            '    password = data.get("password") or ""',
            "    user = User.query.filter_by(username=username).first()",
            "    password_ok = False",
            "    if user:",
            "        try:",
            "            password_ok = check_password_hash(user.password, password)",
            "        except ValueError:",
            "            password_ok = False",
            "        if not password_ok and user.password == password:",
            "            password_ok = True",
            "            user.password = generate_password_hash(password)",
            "            db.session.commit()",
            "    if user and password_ok:",
            '        return jsonify({"status": "success",',
            '            "user": {"username": user.username, "role": user.role}})',
            '    return jsonify({"status": "error", "message": "用户名或密码错误。"}), 401',
        ],
        "代码清单4.2位于backend/routes/auth_route.py。接口先查询用户，再校验密码哈希；若历史演示数据仍为明文密码，则在首次正确登录时自动迁移为哈希值。接口返回username和role供前端完成角色显示与权限控制，使登录功能既满足演示便利性，也保留基本的账号安全处理。",
    )

    hmi_anchor = find_para(doc, lambda t: t.startswith("通过风场总览、单机详情和HMI页面的组合"))
    hmi_anchor = add_code_block(
        hmi_anchor,
        "代码清单5.3 HMI权限登录与菜单切换逻辑",
        [
            "document.getElementById('hmi-login-form').addEventListener('submit', async (event) => {",
            "    event.preventDefault();",
            "    hmiRole = document.getElementById('hmi-role')?.value || 'operator';",
            "    document.getElementById('login-shell').style.display = 'none';",
            "    document.getElementById('hmi-app').style.display = 'grid';",
            "    await loadDetail();",
            "});",
            "document.getElementById('hmi-menu').addEventListener('click', (event) => {",
            "    const button = event.target.closest('button[data-panel]');",
            "    if (!button) return;",
            "    document.querySelectorAll('#hmi-menu button').forEach(item =>",
            "        item.classList.toggle('active', item === button));",
            "    renderDetail(button.dataset.panel);",
            "});",
        ],
        "代码清单5.3位于frontend/hmi.html。HMI页面先显示权限登录表单，提交后隐藏login-shell并显示hmi-app，然后加载单机详情、故障代码和阈值参数。左侧菜单通过data-panel区分运行总览、状态寿命、阈值设置、故障代码等面板，点击菜单后调用renderDetail刷新当前区域，使登录后的各页面能够在同一个HMI控制台中切换。",
    )
    add_code_block(
        hmi_anchor,
        "代码清单5.4 HMI阈值修改与系统参数同步逻辑",
        [
            "async function saveHmiThreshold(key, rawValue) {",
            "    const meta = hmiThresholdMeta[key];",
            "    const value = Number(rawValue);",
            "    if (!Number.isFinite(value) || value < meta.min || value > meta.max) {",
            "        alert(`${meta.label}需在 ${meta.min} - ${meta.max} 范围内`);",
            "        renderDetail('params');",
            "        return;",
            "    }",
            "    const next = { ...hmiThresholds, [key]: value };",
            "    if (next.temp_warning_threshold >= next.temp_threshold ||",
            "        next.vibration_threshold >= next.vibration_critical_threshold ||",
            "        next.oil_warning_threshold >= next.oil_quality_threshold) return;",
            "    hmiThresholds = next;",
            "    await fetch('/api/settings/configs', {",
            "        method: 'POST',",
            "        headers: { 'Content-Type': 'application/json' },",
            "        body: JSON.stringify({ [key]: value })",
            "    });",
            "    renderDetail('params');",
            "}",
        ],
        "代码清单5.4对应HMI阈值设置页面。系统先根据hmiThresholdMeta校验输入范围，再检查关注阈值与临界阈值的大小关系；通过校验后，将当前阈值提交到/api/settings/configs。这样HMI页面中的阈值调整不是单纯的前端显示，而是会同步影响系统参数、告警判断和后续健康报告。",
    )

    settings_anchor = find_para(doc, lambda t: t == "图5.10 系统参数设置页面")
    settings_anchor = add_code_block(
        settings_anchor,
        "代码清单5.5 系统参数阈值校验与保存接口",
        [
            '@settings_bp.route("/configs", methods=["POST"])',
            "def update_configs():",
            "    data = request.json or {}",
            "    existing = {c.config_key: c.config_value for c in SystemConfig.query.all()}",
            "    candidate = {**defaults, **existing,",
            "                 **{k: v for k, v in data.items() if k in ALLOWED_CONFIGS}}",
            "    for warning_key, critical_key in threshold_pairs:",
            "        warning_value = float(candidate.get(warning_key, 0))",
            "        critical_value = float(candidate.get(critical_key, 0))",
            "        if warning_value >= critical_value:",
            '            return jsonify({"status": "error",',
            '                "message": "关注阈值必须小于告警阈值。"}), 400',
            "    for key, value in data.items():",
            "        if key not in ALLOWED_CONFIGS:",
            "            continue",
            "        numeric_value = float(value)",
            "        min_value, max_value = CONFIG_RANGES[key]",
            "        if not min_value <= numeric_value <= max_value:",
            "            return jsonify({'status': 'error'}), 400",
            "        config = SystemConfig.query.filter_by(config_key=key).first()",
            "        config.config_value = str(value)",
            "    db.session.commit()",
            '    return jsonify({"status": "success", "configs": updated})',
        ],
        "代码清单5.5位于backend/routes/settings_route.py。后端接口只允许ALLOWED_CONFIGS中的参数被修改，并通过CONFIG_RANGES限制每个阈值的取值范围；同时用threshold_pairs保证关注阈值小于告警阈值。该接口是系统参数页面和HMI阈值设置页面的共同后端入口，保证不同前端入口写入的是同一套系统配置。",
    )

    for p in doc.paragraphs:
        if p.text.strip() == "代码清单5.3 前端实时状态刷新逻辑":
            p.text = "代码清单5.6 前端实时状态刷新逻辑"
            style_code_caption(p)

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
