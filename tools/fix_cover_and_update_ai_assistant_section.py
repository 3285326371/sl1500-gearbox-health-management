from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")
DOCX = (
    Path.home()
    / "Desktop"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f_\u68c0\u6d4b\u62a5\u544a"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
)
SHOT = ROOT / "docs" / "system_screenshots_updated" / "09_ai_qa_assistant_updated.png"


def set_run_font(run, size=12, east="宋体", ascii_font="Times New Roman", bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(p):
    for child in list(p._p):
        p._p.remove(child)


def style_body(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, 12, "宋体", "Times New Roman")


def style_code_title(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, 10.5, "黑体", "Arial", bold=False)


def style_code_line(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.95)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, 10.5, "宋体", "Consolas")


def style_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        set_run_font(run, 10.5, "黑体", "Arial", bold=False)


def paragraph_after(anchor, text=""):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    return p


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"not found: {text}")


def para_index(doc, target):
    for i, p in enumerate(doc.paragraphs):
        if p._p is target._p:
            return i
    raise RuntimeError("paragraph index not found")


def replace_ai_code(doc):
    start = find_para(doc, "代码清单5.6 前端实时状态刷新逻辑")
    idx = para_index(doc, start)
    # Remove old title, "核心代码如下：" and old code lines through refreshStatus().
    for p in list(doc.paragraphs[idx:idx + 14]):
        remove_paragraph(p)

    anchor = doc.paragraphs[idx - 1]
    lines = [
        "代码清单5.6 智能运维诊断助手流式问答逻辑",
        "核心代码如下：",
        "async function sendQuestion(customQuestion = null) {",
        "    const question = customQuestion || qaInput.value.trim();",
        "    addMessage(question, true);",
        "    const deepThinking = qaDeepThinking ? qaDeepThinking.checked : true;",
        "    const streaming = createStreamingAssistantMessage();",
        "    let streamedAnswer = '';",
        "    let streamMeta = {};",
        "    const response = await fetch('/api/qa/ask_stream', {",
        "        method: 'POST',",
        "        headers: { 'Content-Type': 'application/json' },",
        "        body: JSON.stringify({",
        "            question: question,",
        "            current_status: qaCurrentStatusPayload(),",
        "            deep_thinking: deepThinking,",
        "            answer_mode: deepThinking ? 'model_first' : 'local_only'",
        "        })",
        "    });",
        "    await readSseResponse(response, {",
        "        delta(payload) {",
        "            streamedAnswer += payload.text || '';",
        "            appendStreamingText(streaming.streamBox, payload.text || '');",
        "        },",
        "        done(payload) { streamMeta = { ...streamMeta, ...payload }; }",
        "    });",
        "    finalizeStreamingMessage(streaming, streamedAnswer, {",
        "        sources: streamMeta.sources || [],",
        "        confidence: streamMeta.confidence,",
        "        riskLevel: streamMeta.risk_level,",
        "        engine: streamMeta.engine",
        "    });",
        "}",
    ]
    for pos, line in enumerate(lines):
        p = paragraph_after(anchor, line)
        if pos == 0:
            style_code_title(p)
        elif pos == 1:
            style_body(p)
        else:
            style_code_line(p)
        anchor = p


def update_text(doc):
    replacements = {
        "AI 运维问答模块位于 services/rag_service.py 和 routes/qa_route.py。系统内置本地知识库，知识项包括油温过高排查、高速轴承冲击诊断、齿面磨损与点蚀、齿轮断齿与啮合异常、轴承磨损与高速端过温、润滑油污染、联轴器不对中、箱体或基础松动、齿轮箱故障清单、剩余寿命评估、M-IALO-SVR 方法和检修策略。":
            "智能运维诊断助手位于 frontend/index.html、frontend/js/main.js、routes/qa_route.py 和 services/rag_service.py。更新后的页面由当前机组上下文、诊断对话、输出要求和智能程度提示四部分组成，能够把油温、振动 RMS、油液 NAS、RUL、健康评分和功率等实时工况随问题一并提交给后端问答接口。",
        "当用户提出问题时，系统先识别问题意图，如故障分类、算法原理、排查建议、原因分析、寿命预测或状态评估。然后根据关键词和内容匹配知识库条目，并结合当前油温、振动、油液、功率、健康评分和 RUL 生成结论、实时工况、判断依据、建议步骤和可继续追问。如果开启大模型辅助分析且配置了 API Key，系统会调用兼容 OpenAI 格式的聊天接口生成更完整回答；若调用失败，则回退到本地专家引擎。":
            "当用户提出问题时，前端先将问题显示为用户消息，再创建流式助手消息，并把当前机组状态、是否启用“大模型主答”和回答模式一并提交到 /api/qa/ask_stream。后端返回 SSE 数据流，前端按 delta 事件逐步追加文本，按 done 事件补充置信度、风险等级、证据来源和后续追问，使回答过程更接近真实智能助手交互。",
        "AI运维问答页面如图5.9所示。用户可以围绕油温、振动、轴承、齿轮、润滑和检修策略提出问题，系统结合知识库和当前状态生成回答。":
            "智能运维诊断助手页面如图5.9所示。用户可以围绕当前机组状态、停机判断、振动诊断、油温排查、工单生成和算法解释提出问题，系统结合实时工况、本地知识库和大模型接口生成结构化处置建议。",
        "AI运维问答页面面向运维人员的日常咨询场景。传统系统通常只能提供固定报警，而问答模块可以让用户以自然语言提出问题，例如“齿轮箱油温过高如何排查”“振动RMS升高是否需要停机”“油液NAS等级偏高如何处理”等。系统根据问题意图匹配知识库条目，并结合当前状态给出排查顺序。":
            "智能运维诊断助手面向运维人员的日常咨询和答辩演示场景。左侧上下文面板用于确认当前机组和关键指标，中间诊断对话区用于显示用户问题、助手流式回答和证据标签，右侧输出要求区明确回答必须包含结论、依据、排查步骤和工单判断。该设计使问答结果不只是文本回复，而是与齿轮箱健康管理业务流程相连接。",
        "在实现上，问答模块先使用本地知识库保证回答稳定性，再根据配置决定是否调用外部大模型接口。即使没有网络或API Key，系统仍然能够返回专家规则答案。这样既满足毕业设计演示稳定性，也体现了智能运维系统可扩展到大模型辅助分析的方向。":
            "在实现上，系统默认使用大模型主答、本地知识库作为证据补充；当大模型接口不可用时，后端会回退到本地专家规则，保证演示稳定性。回答渲染函数会把“结论、风险等级、关键依据、建议步骤、后续追问”等标题转换为卡片化结构，并提供查看实时趋势、故障记录和健康报告的跳转按钮，增强运维闭环能力。",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            p.text = replacements[text]
            style_body(p)


def replace_ai_image(doc):
    caption = find_para(doc, "图5.9 AI运维问答页面")
    # Previous paragraph contains the old image.
    prev = None
    for p in doc.paragraphs:
        if p._p is caption._p.getprevious():
            prev = p
            break
    if prev is not None:
        clear_paragraph(prev)
        prev.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prev.paragraph_format.first_line_indent = None
        prev.add_run().add_picture(str(SHOT), width=Cm(15.2))
    caption.text = "图5.9 智能运维诊断助手页面"
    style_caption(caption)


def fix_cover_page_break(doc):
    # The originality declaration must start on page 2; the batch format pass
    # removed this separation and caused the cover to spill into the next page.
    target = find_para(doc, "毕业设计原创性声明")
    target.paragraph_format.page_break_before = True


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_封面与助手页更新前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)
    fix_cover_page_break(doc)
    update_text(doc)
    replace_ai_code(doc)
    replace_ai_image(doc)
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
