from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def style_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(text)


def add_after(anchor, text):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    style_body(p)
    return p


def insert_once(doc, anchor_text, marker, text):
    if any(marker in p.text for p in doc.paragraphs):
        return
    add_after(find_para(doc, anchor_text), text)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_二次减少留白前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)

    insert_once(
        doc,
        "本文技术路线可以概括为：需求分析与故障机理研究、数据接入与预处理、特征提取与风险评分、故障诊断与健康评估、后端接口与数据库设计、前端可视化实现、系统测试与结果分析。系统首先获取齿轮箱油温、振动、油液、功率、风速、转速、故障码等数据；其次进行滤波、异常值处理和字段归一化；然后提取振动特征并计算风险因子；最后输出故障诊断、健康评分、RUL、复检周期和工单建议，并通过前端页面展示给运维人员。",
        "该技术路线的重点在于把机械故障机理与软件系统实现结合起来",
        "该技术路线的重点在于把机械故障机理与软件系统实现结合起来。一方面，论文需要说明齿轮箱故障产生的物理原因和监测指标意义；另一方面，系统需要把这些指标转化为可操作的页面、接口和数据记录。通过这种方式，本文不是单独讨论算法，也不是单独开发管理页面，而是围绕齿轮箱健康管理目标形成较完整的技术闭环。",
    )

    insert_once(
        doc,
        "在毕业设计实现中，技术选型更强调可实现性和可演示性。Flask 便于快速组织接口，SQLite 适合轻量化部署，ECharts 能够直观展示趋势和分布，SSE 可以降低实时刷新实现复杂度。该组合虽然不等同于生产级风电场平台，但能够覆盖原型系统的核心功能，并为后续接入真实 SCADA、CMS 和油液监测数据预留扩展空间。",
        "Web系统开发技术的作用还体现在系统集成能力上",
        "Web系统开发技术的作用还体现在系统集成能力上。齿轮箱健康管理涉及实时数据、故障规则、历史记录、报告输出和用户交互等多个环节，如果缺少统一的平台载体，诊断结果难以被运维人员持续使用。通过 Web 系统将各类功能组织在同一界面中，可以提高系统演示的完整性，也便于后续扩展移动端访问、权限控制和远程运维功能。",
    )

    insert_once(
        doc,
        "系统数据库围绕用户、风机资产、故障记录、诊断结果、健康报告和系统参数展开。风机资产表是业务核心，故障记录、诊断结果和健康报告均通过机组编号与其关联，系统参数表为阈值配置和诊断规则提供支撑。",
        "ER图设计强调数据之间的业务关系",
        "ER图设计强调数据之间的业务关系。用户表对应系统访问主体，风机资产表对应被监测对象，故障记录表对应诊断事件，诊断结果和健康报告对应分析输出，系统参数表对应规则配置。通过这些实体关系，系统能够把一次故障诊断从页面操作转化为可保存、可查询、可统计的业务数据。",
    )

    insert_once(
        doc,
        "在人机交互反馈方面，系统将报警状态、健康评分、RUL、故障类型和建议措施放在同一业务链路中呈现。用户既能看到实时指标，也能看到系统给出判断的依据和后续处理建议。对于齿轮箱这类高价值设备，仅显示单一报警数值并不足够，必须同时提供故障解释、处理优先级和闭环入口，才能体现智能健康管理系统的实际意义。",
        "因此，本文前端设计并非单纯追求页面美观",
        "因此，本文前端设计并非单纯追求页面美观，而是围绕运维人员的判断效率展开。总览页面强调快速发现问题，详情页面强调指标解释，诊断页面强调结论可信度，HMI 页面强调现场设备对象感，报告页面强调结果沉淀。不同页面承担不同任务，能够减少信息堆叠造成的阅读负担。",
    )

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
