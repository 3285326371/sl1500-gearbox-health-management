from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX = Path(r"C:\Users\隐涅\Desktop\毕业设计正文模板_检测报告\毕业设计正文模板.docx")


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_children(parent, tag):
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    remove_children(p_pr, "w:keepNext")
    if value:
        p_pr.append(OxmlElement("w:keepNext"))


def set_keep_together(paragraph, value=False):
    p_pr = paragraph._p.get_or_add_pPr()
    remove_children(p_pr, "w:keepLines")
    if value:
        p_pr.append(OxmlElement("w:keepLines"))


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    remove_children(tr_pr, "w:cantSplit")
    tr_pr.append(OxmlElement("w:cantSplit"))


def remove_table_shading(table):
    # Keep the three-line borders already present, but remove any cell background fill.
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            remove_children(tc_pr, "w:shd")


def main():
    backup = DOCX.with_name(
        DOCX.stem + "_表格不跨页修改前备份_" + datetime.now().strftime("%Y%m%d_%H%M%S") + DOCX.suffix
    )
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    for table in doc.tables:
        for row in table.rows:
            set_row_cant_split(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Avoid row-internal paragraph rules that can create large page gaps.
                    set_keep_together(paragraph, False)
        remove_table_shading(table)

    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs[:-1]):
        text = paragraph.text.strip()
        next_el = paragraph._p.getnext()
        if text.startswith("表") and next_el is not None and next_el.tag == qn("w:tbl"):
            set_keep_with_next(paragraph, True)
            set_keep_together(paragraph, False)

    doc.save(DOCX)
    print("tables", len(doc.tables))
    print("backup", backup)


if __name__ == "__main__":
    main()
