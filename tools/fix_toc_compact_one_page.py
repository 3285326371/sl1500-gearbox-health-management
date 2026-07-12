from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_目录修改前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_目录修改.docx"

TOC_LINES = [
    ("摘要", "I", 0),
    ("ABSTRACT", "II", 0),
    ("第1章 绪论", "1", 0),
    ("1.1 研究背景与意义", "1", 1),
    ("1.2 国内外研究现状", "1", 1),
    ("1.3 本文研究内容", "2", 1),
    ("1.4 技术路线", "2", 1),
    ("第2章 系统需求分析", "3", 0),
    ("2.1 业务需求分析", "3", 1),
    ("2.2 功能需求分析", "3", 1),
    ("2.3 非功能需求分析", "4", 1),
    ("2.4 数据需求分析", "4", 1),
    ("第3章 相关技术与理论基础", "5", 0),
    ("3.1 齿轮箱典型故障机理", "5", 1),
    ("3.2 多源状态监测技术", "5", 1),
    ("3.3 振动特征提取方法", "5", 1),
    ("3.4 风险评分与健康评估方法", "5", 1),
    ("3.5 Web系统开发技术", "6", 1),
    ("第4章 系统总体设计", "7", 0),
    ("4.1 系统架构设计", "7", 1),
    ("4.2 数据库设计", "7", 1),
    ("4.3 后端接口设计", "7", 1),
    ("4.4 诊断流程设计", "8", 1),
    ("4.5 前端页面设计", "9", 1),
    ("第5章 系统详细实现", "10", 0),
    ("5.1 数据采集与融合模块实现", "10", 1),
    ("5.2 故障诊断算法模块实现", "10", 1),
    ("5.3 风场总览与单机详情实现", "10", 1),
    ("5.4 故障代码库与数据管理实现", "11", 1),
    ("5.5 AI运维问答模块实现", "11", 1),
    ("5.6 健康报告与工单闭环实现", "11", 1),
    ("第6章 系统测试与结果分析", "12", 0),
    ("6.1 测试环境", "12", 1),
    ("6.2 功能测试", "12", 1),
    ("6.3 诊断场景测试分析", "12", 1),
    ("6.4 结果评价", "13", 1),
    ("第7章 总结与展望", "14", 0),
    ("参考文献", "15", 0),
    ("致谢", "16", 0),
    ("附录A 系统主要接口", "17", 0),
]


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def format_toc_paragraph(p, level):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.38 if level else 0)
    fmt.first_line_indent = Inches(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(15)
    fmt.tab_stops.clear_all()
    fmt.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() in {"目  录", "目□□录"}) + 1
    end = next(i for i in range(start, len(doc.paragraphs)) if doc.paragraphs[i].text.strip() == "第1章 绪论")

    for offset, (title, page, level) in enumerate(TOC_LINES):
        p = doc.paragraphs[start + offset]
        p.clear()
        format_toc_paragraph(p, level)
        r1 = p.add_run(title)
        set_run_font(r1, bold=False)
        p.add_run("\t")
        r2 = p.add_run(page)
        set_run_font(r2, bold=False)

    for i in range(start + len(TOC_LINES), end):
        doc.paragraphs[i].clear()
        doc.paragraphs[i].paragraph_format.space_before = Pt(0)
        doc.paragraphs[i].paragraph_format.space_after = Pt(0)
        doc.paragraphs[i].paragraph_format.line_spacing = Pt(1)

    # Tighten TOC title spacing slightly.
    title_p = doc.paragraphs[start - 1]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    for r in title_p.runs:
        set_run_font(r, size=16, bold=True)

    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
