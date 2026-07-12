from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
HEADER_TEXT = "河北建筑工程学院学士学位毕业设计"


CHAPTER_TITLES = {
    "摘  要", "ABSTRACT", "目  录", "关于毕业设计（论文）的著作权声明",
    "毕业设计原创性声明", "毕业设计版权使用授权书", "参考文献", "致谢", "附录",
}


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold


def ensure_runs(paragraph):
    if not paragraph.runs:
        paragraph.add_run("")
    return paragraph.runs


def is_chapter(text):
    stripped = text.strip()
    return stripped in CHAPTER_TITLES or stripped.startswith("第") and "章" in stripped


def is_second_level(text):
    stripped = text.strip()
    if len(stripped) < 4:
        return False
    parts = stripped.split(maxsplit=1)
    head = parts[0]
    nums = head.split(".")
    return len(nums) == 2 and all(n.isdigit() for n in nums)


def is_third_level(text):
    stripped = text.strip()
    parts = stripped.split(maxsplit=1)
    head = parts[0] if parts else ""
    nums = head.split(".")
    return len(nums) == 3 and all(n.isdigit() for n in nums)


def normalize_heading_space(text):
    stripped = text.strip()
    if stripped.startswith("第") and "章" in stripped:
        # 第1章 -> 第 1 章
        import re
        m = re.match(r"第\s*(\d+)\s*章\s*(.*)", stripped)
        if m:
            return f"第 {m.group(1)} 章 {m.group(2).strip()}".strip()
    import re
    m = re.match(r"^(\d+(?:\.\d+){1,2})\s*(.+)$", stripped)
    if m:
        return f"{m.group(1)} {m.group(2).strip()}"
    return stripped


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def format_chapter(paragraph):
    text = normalize_heading_space(paragraph.text)
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(40)
    fmt.space_after = Pt(20)
    fmt.line_spacing = Pt(20)
    fmt.first_line_indent = None
    for run in ensure_runs(paragraph):
        set_run_font(run, "黑体", "Arial", 15, True)


def format_second(paragraph):
    text = normalize_heading_space(paragraph.text)
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(24)
    fmt.space_after = Pt(6)
    fmt.line_spacing = Pt(20)
    fmt.first_line_indent = None
    for run in ensure_runs(paragraph):
        set_run_font(run, "黑体", "Arial", 14, True)


def format_third(paragraph):
    text = normalize_heading_space(paragraph.text)
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = Pt(20)
    fmt.first_line_indent = None
    for run in ensure_runs(paragraph):
        set_run_font(run, "黑体", "Arial", 12, True)


def format_body(paragraph):
    text = paragraph.text.strip()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(20)
    fmt.first_line_indent = Cm(0.74) if text else None
    for run in ensure_runs(paragraph):
        set_run_font(run, "宋体", "Times New Roman", 12, False)


def format_cover(paragraph):
    # Keep cover centered but make it consistent.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(20)
    for run in ensure_runs(paragraph):
        text = paragraph.text.strip()
        if text == "学士学位毕业设计" or "系统设计与实现" in text or "齿轮箱" in text:
            set_run_font(run, "黑体", "Arial", 16, True)
        else:
            set_run_font(run, "宋体", "Times New Roman", 12, False)


def set_header(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = HEADER_TEXT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, "宋体", "Times New Roman", 10.5, False)


def format_table_text(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = Pt(20)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for run in ensure_runs(p):
                        set_run_font(run, "宋体", "Times New Roman", 12, run.bold)


def main():
    candidates = [
        p for p in DOCS.glob("*毕业论文*.docx")
        if not p.name.startswith("~$") and "格式规范版" not in p.stem
    ]
    source = max(candidates, key=lambda p: p.stat().st_mtime)
    out = source.with_name(source.stem + "_格式规范版.docx")
    doc = Document(source)
    set_header(doc)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if idx < 12:
            format_cover(paragraph)
        elif is_chapter(text):
            format_chapter(paragraph)
        elif is_third_level(text):
            format_third(paragraph)
        elif is_second_level(text):
            format_second(paragraph)
        else:
            format_body(paragraph)

    format_table_text(doc)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
