from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_正文逐段润色版.docx")
OUT = SRC


FIXES = {
    "前端页面按“总览、下钻、诊断、报告、配置”组织，页面职责见表4.7。": "前端页面按“总览、下钻、诊断、报告、配置”组织，页面职责见表4.3。",
    "登录注册测试中，系统可以完成管理员账号登录和新用户注册，并在前端显示用户角色。风场总览测试中，系统可以展示多机组运行状态、风场统计指标和状态图例。单机详情测试中，选择不同机组后，油温、功率、风速、健康评分、RUL、告警和集群对标可以随之变化。": "登录注册测试中，管理员账号登录、新用户注册和角色显示均能正常完成。风场总览页面可以展示多机组运行状态、风场统计指标和状态图例；切换到单机详情后，油温、功率、风速、健康评分、RUL、告警和集群对标会随所选机组同步变化。",
    "实时监测测试中，系统通过 SSE 持续推送状态数据和振动波形，前端可以刷新油温、振动、功率、RUL、采集链路和图表。故障诊断测试中，选择不同场景后，系统可以生成相应故障类型、风险等级、置信度、健康评分、RUL、风险因子和维护建议。故障数据测试中，系统可以显示诊断记录和历史故障记录，并支持导出 CSV。": "实时监测测试中，后端通过 SSE 持续推送状态数据和振动波形，前端可以刷新油温、振动、功率、RUL、采集链路和图表。故障诊断测试中，选择不同场景后，页面会返回对应的故障类型、风险等级、置信度、健康评分、RUL、风险因子和维护建议。故障数据页面可以显示诊断记录和历史故障记录，并支持导出 CSV。",
    "在算法流程方面，系统实现了振动特征提取、风险因子计算、故障分类、健康评分和 RUL 估计。系统可以识别齿轮齿面磨损、齿面点蚀/剥落、齿轮断齿、轴承磨损、轴承点蚀/剥落、高速轴承过温、齿轮箱油温过高、润滑油污染/油液劣化、轴系不对中、齿轮啮合异常、箱体或基础松动和冷却系统故障等常见异常，并给出诊断依据和维护建议。": "算法流程方面，振动特征提取、风险因子计算、故障分类、健康评分和 RUL 估计已经串联起来。当前规则库覆盖齿轮齿面磨损、齿面点蚀/剥落、齿轮断齿、轴承磨损、轴承点蚀/剥落、高速轴承过温、齿轮箱油温过高、润滑油污染/油液劣化、轴系不对中、齿轮啮合异常、箱体或基础松动和冷却系统故障等常见异常，并输出诊断依据和维护建议。",
}


def set_text(paragraph, text):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main():
    tmp = SRC.with_name(SRC.stem + "_tmp.docx")
    shutil.copy2(SRC, tmp)
    doc = Document(tmp)
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in FIXES:
            set_text(paragraph, FIXES[text])
            changed += 1
    doc.save(tmp)
    shutil.move(tmp, OUT)
    print(f"saved={OUT}")
    print(f"changed={changed}")


if __name__ == "__main__":
    main()
