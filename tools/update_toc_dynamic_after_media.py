from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_图片代码目录同步.docx"

HEADINGS = [
    ("摘要", "I", 0),
    ("ABSTRACT", "II", 0),
    ("第1章 绪论", None, 0),
    ("1.1 研究背景与意义", None, 1),
    ("1.2 国内外研究现状", None, 1),
    ("1.3 本文研究内容", None, 1),
    ("1.4 技术路线", None, 1),
    ("第2章 系统需求分析", None, 0),
    ("2.1 业务需求分析", None, 1),
    ("2.2 功能需求分析", None, 1),
    ("2.3 非功能需求分析", None, 1),
    ("2.4 数据需求分析", None, 1),
    ("第3章 相关技术与理论基础", None, 0),
    ("3.1 齿轮箱典型故障机理", None, 1),
    ("3.2 多源状态监测技术", None, 1),
    ("3.3 振动特征提取方法", None, 1),
    ("3.4 风险评分与健康评估方法", None, 1),
    ("3.5 Web系统开发技术", None, 1),
    ("第4章 系统总体设计", None, 0),
    ("4.1 系统架构设计", None, 1),
    ("4.2 数据库设计", None, 1),
    ("4.3 后端接口设计", None, 1),
    ("4.4 诊断流程设计", None, 1),
    ("4.5 前端页面设计", None, 1),
    ("第5章 系统详细实现", None, 0),
    ("5.1 数据采集与融合模块实现", None, 1),
    ("5.2 故障诊断算法模块实现", None, 1),
    ("5.3 风场总览与单机详情实现", None, 1),
    ("5.4 故障代码库与数据管理实现", None, 1),
    ("5.5 AI运维问答模块实现", None, 1),
    ("5.6 健康报告与工单闭环实现", None, 1),
    ("第6章 系统测试与结果分析", None, 0),
    ("6.1 测试环境", None, 1),
    ("6.2 功能测试", None, 1),
    ("6.3 诊断场景测试分析", None, 1),
    ("6.4 结果评价", None, 1),
    ("第7章 总结与展望", None, 0),
    ("参考文献", None, 0),
    ("致谢", None, 0),
    ("附录A 系统主要接口", None, 0),
]


def word_pages():
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(TEMPLATE), False, True)
    try:
        wanted = {h[0] for h in HEADINGS if h[1] is None}
        found = {}
        for i in range(1, doc.Paragraphs.Count + 1):
            text = doc.Paragraphs.Item(i).Range.Text.strip().replace("\r", "").replace("\a", "")
            normalized = text.replace(" ", "")
            for target in wanted:
                if normalized == target.replace(" ", "") and target not in found:
                    found[target] = doc.Paragraphs.Item(i).Range.Information(3)
        first = found["第1章 绪论"]
        offset = first - 1
        return {k: str(v - offset) for k, v in found.items()}
    finally:
        doc.Close(False)
        word.Quit()


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def format_toc_paragraph(p, level):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.38 if level else 0)
    fmt.first_line_indent = Inches(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(15)
    fmt.tab_stops.clear_all()
    fmt.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def main():
    pages = word_pages()
    doc = Document(TEMPLATE)
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() in {"目  录", "目□□录"}) + 1
    end = next(i for i in range(start, len(doc.paragraphs)) if doc.paragraphs[i].text.strip().replace(" ", "") == "第1章绪论")
    for offset, (title, fixed, level) in enumerate(HEADINGS):
        p = doc.paragraphs[start + offset]
        p.clear()
        format_toc_paragraph(p, level)
        r1 = p.add_run(title)
        set_run_font(r1)
        p.add_run("\t")
        r2 = p.add_run(fixed or pages.get(title, ""))
        set_run_font(r2)
    for i in range(start + len(HEADINGS), end):
        doc.paragraphs[i].clear()
        doc.paragraphs[i].paragraph_format.line_spacing = Pt(1)
        doc.paragraphs[i].paragraph_format.space_after = Pt(0)
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(TMP_OUT)
    print(pages)


if __name__ == "__main__":
    main()
