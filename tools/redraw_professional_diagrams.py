from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_人工化润色版.docx")
OUT = SRC.with_name("毕业设计正文模板_结构图美化版.docx")
ASSET_DIR = SRC.with_name("结构图美化版_assets")


DIAGRAMS = {
    "图1.1 本文技术路线图": {
        "type": "pipeline",
        "title": "本文技术路线",
        "steps": [
            ("需求分析", "齿轮箱故障机理\n业务场景梳理"),
            ("数据接入", "油温/振动/油液\n风机摘要文件"),
            ("特征计算", "RMS/峭度/包络\n风险因子归一"),
            ("诊断评估", "故障分类\n健康评分/RUL"),
            ("系统实现", "前后端接口\n报告与闭环"),
            ("测试评价", "功能测试\n典型场景验证"),
        ],
    },
    "图3.1 健康评分与 RUL 估计示意图": {
        "type": "model",
        "title": "健康评分与 RUL 估计模型",
        "left": ["振动 RMS", "峭度", "峰值因子", "包络峰值", "油温", "油液 NAS"],
        "middle": ["归一化处理", "权重融合", "严重度修正"],
        "right": ["综合风险", "健康评分", "RUL", "复检周期"],
    },
    "图4.2 系统数据流与模块关系图": {
        "type": "layer",
        "title": "系统数据流与模块关系",
        "layers": [
            ("数据来源", ["模拟采集", "风机摘要", "故障统计", "系统配置"]),
            ("后端处理", ["数据融合", "特征提取", "风险评分", "规则诊断", "报告生成"]),
            ("前端展示", ["风场总览", "单机详情", "实时监测", "诊断结果", "健康报告"]),
        ],
    },
    "图5.6 HMI运行总览页面结构示意图": {
        "type": "layer",
        "title": "HMI 运行总览页面结构",
        "layers": [
            ("权限入口", ["二次登录", "角色级别", "机组选择"]),
            ("运行状态", ["油温", "振动 RMS", "健康评分", "RUL", "透明模型"]),
            ("现场操作", ["阈值设置", "测点信号", "故障代码", "临界报警"]),
        ],
    },
    "图5.10 故障诊断与运维闭环图": {
        "type": "closed_loop",
        "title": "故障诊断与运维闭环",
        "steps": [
            ("发现", "实时监测\n阈值告警"),
            ("判断", "风险评分\n故障分类"),
            ("处理", "维护建议\n工单生成"),
            ("记录", "故障记录\n报告归档"),
            ("复盘", "复测结果\n参数优化"),
        ],
        "sub": ["总览定位\n单机下钻", "特征依据\n置信度", "责任班组\n处理窗口", "状态更新\nCSV导出", "知识库补充\n模型迭代"],
    },
    "图5.12 系统整体分层架构图": {
        "type": "layer",
        "title": "系统整体分层架构",
        "layers": [
            ("表示层", ["风场总览", "单机详情", "实时监测", "故障诊断", "健康报告", "AI问答", "HMI"]),
            ("业务逻辑层", ["认证路由", "风场接口", "数据接口", "诊断服务", "报告服务", "问答服务"]),
            ("数据层", ["SQLite", "风机摘要", "模拟采集", "故障代码库", "系统配置"]),
        ],
    },
    "图5.13 风场总览与单机下钻流程图": {
        "type": "pipeline",
        "title": "风场总览与单机下钻流程",
        "steps": [
            ("登录", "用户验证\n角色识别"),
            ("总览", "机组矩阵\n统计指标"),
            ("定位", "异常状态\n告警数量"),
            ("下钻", "单机详情\n部件状态"),
            ("处置", "诊断/报告\nHMI入口"),
        ],
    },
    "图5.14 多源数据采集与融合流程图": {
        "type": "pipeline",
        "title": "多源数据采集与融合流程",
        "steps": [
            ("采集", "油温/振动\n油液/功率"),
            ("读取", "风机摘要\n故障统计"),
            ("清洗", "字段统一\n异常处理"),
            ("融合", "机组快照\n趋势数据"),
            ("输出", "页面展示\n诊断输入"),
        ],
    },
    "图5.15 齿轮箱故障诊断流程图": {
        "type": "pipeline",
        "title": "齿轮箱故障诊断流程",
        "steps": [
            ("选择机组", "读取状态\n选择场景"),
            ("生成信号", "振动波形\n工况修正"),
            ("提取特征", "RMS/峭度\n包络峰值"),
            ("风险评分", "油温/油液\n综合风险"),
            ("诊断输出", "故障类型\n建议措施"),
        ],
    },
    "图5.16 风险评分与健康评估模型图": {
        "type": "model",
        "title": "风险评分与健康评估模型",
        "left": ["RMS", "峭度", "峰值因子", "包络峰值", "油温", "油液 NAS"],
        "middle": ["分项风险", "权重融合", "故障严重度"],
        "right": ["综合风险", "健康评分", "RUL", "维护建议"],
    },
    "图5.17 数据库实体关系图": {
        "type": "er",
        "title": "数据库实体关系",
        "entities": [
            ("User", ["id", "username", "password_hash", "role"]),
            ("TurbineAsset", ["unit_id", "model", "location", "design_life"]),
            ("FaultRecord", ["unit_id", "fault_type", "severity", "status"]),
            ("FaultClosure", ["record_id", "handler", "result", "closed_at"]),
            ("SystemConfig", ["config_key", "config_value", "description"]),
            ("HealthReport", ["unit_id", "score", "rul", "summary"]),
        ],
        "links": [(0, 2), (1, 2), (2, 3), (1, 5), (4, 2), (4, 5)],
    },
    "图5.18 后端接口与蓝图结构图": {
        "type": "layer",
        "title": "后端接口与蓝图结构",
        "layers": [
            ("应用入口", ["app.py", "Flask应用", "蓝图注册"]),
            ("接口层", ["auth_bp", "data_bp", "windfarm_bp", "report_bp", "qa_bp", "settings_bp"]),
            ("服务层", ["DataAcquisition", "WindRepository", "DiagnosisModel", "RagService", "ReportService"]),
        ],
    },
    "图5.19 前端页面模块结构图": {
        "type": "layer",
        "title": "前端页面模块结构",
        "layers": [
            ("导航分组", ["健康监测区", "智能诊断区", "系统配置区"]),
            ("功能页面", ["总览", "详情", "实时监测", "故障数据", "代码库", "报告", "问答"]),
            ("交互输出", ["指标卡片", "趋势曲线", "诊断卡片", "记录表格", "跳转按钮"]),
        ],
    },
    "图5.20 HMI运维控制台结构图": {
        "type": "layer",
        "title": "HMI 运维控制台结构",
        "layers": [
            ("入口", ["权限登录", "角色确认", "机组选择"]),
            ("监控", ["油温", "振动", "健康评分", "RUL", "透明模型"]),
            ("操作", ["阈值设置", "故障代码", "临界报警", "测点信号"]),
        ],
    },
    "图5.21 故障代码库与记录管理关系图": {
        "type": "model",
        "title": "故障代码库与记录管理关系",
        "left": ["GBX-001", "GBX-002", "GBX-003", "SCADA动态码"],
        "middle": ["诊断匹配", "状态判断", "记录入库"],
        "right": ["筛选查询", "详情查看", "CSV导出", "闭环处理"],
    },
    "图5.22 健康报告生成流程图": {
        "type": "pipeline",
        "title": "健康报告生成流程",
        "steps": [
            ("读取状态", "当前机组\n实时指标"),
            ("汇总历史", "故障记录\n趋势数据"),
            ("计算结果", "健康评分\nRUL"),
            ("形成结论", "风险说明\n维护计划"),
            ("输出报告", "页面展示\n归档保存"),
        ],
    },
    "图5.23 AI运维问答处理流程图": {
        "type": "pipeline",
        "title": "AI 运维问答处理流程",
        "steps": [
            ("输入问题", "用户提问\n当前机组"),
            ("补充上下文", "油温/振动\n健康评分"),
            ("检索知识", "本地知识库\n实时状态"),
            ("生成回答", "模型主答\n规则回退"),
            ("输出建议", "依据/步骤\n工单判断"),
        ],
    },
    "图5.24 参数设置影响链路图": {
        "type": "model",
        "title": "参数设置影响链路",
        "left": ["油温阈值", "振动阈值", "油液阈值", "报警策略"],
        "middle": ["实时颜色", "风险因子", "故障等级"],
        "right": ["报告建议", "工单优先级", "复检周期", "权限审计"],
    },
    "图5.25 运维工单闭环流程图": {
        "type": "closed_loop",
        "title": "运维工单闭环流程",
        "steps": [
            ("触发", "监测告警\n诊断异常"),
            ("派单", "生成工单\n确定优先级"),
            ("处理", "现场复核\n检修操作"),
            ("验收", "复测确认\n状态更新"),
            ("沉淀", "记录归档\n知识补充"),
        ],
        "sub": ["异常来源", "责任班组", "处理说明", "复测结果", "报告归档"],
    },
}


def get_font(size, bold=False):
    paths = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in paths:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


FONT_TITLE = get_font(34, True)
FONT_HEAD = get_font(25, True)
FONT_BODY = get_font(21)
FONT_SMALL = get_font(17)


def draw_text_center(d, box, text, font, gap=5):
    x1, y1, x2, y2 = box
    lines = str(text).split("\n")
    metrics = [d.textbbox((0, 0), line, font=font) for line in lines]
    heights = [m[3] - m[1] for m in metrics]
    total_h = sum(heights) + gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, bbox, h in zip(lines, metrics, heights):
        w = bbox[2] - bbox[0]
        d.text((x1 + (x2 - x1 - w) / 2, y), line, fill="black", font=font)
        y += h + gap


def arrow(d, start, end, width=3):
    x1, y1 = start
    x2, y2 = end
    d.line([x1, y1, x2, y2], fill="black", width=width)
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) >= abs(dy):
        if dx >= 0:
            pts = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
        else:
            pts = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    else:
        if dy >= 0:
            pts = [(x2, y2), (x2 - 8, y2 - 14), (x2 + 8, y2 - 14)]
        else:
            pts = [(x2, y2), (x2 - 8, y2 + 14), (x2 + 8, y2 + 14)]
    d.polygon(pts, fill="black")


def base_canvas(title, size=(1600, 700)):
    img = Image.new("RGB", size, "white")
    d = ImageDraw.Draw(img)
    w, h = size
    d.text((70, 52), title, fill="black", font=FONT_TITLE)
    d.line([70, 105, w - 70, 105], fill="black", width=3)
    return img, d


def draw_pipeline(path, spec):
    img, d = base_canvas(spec["title"], (1600, 610))
    steps = spec["steps"]
    n = len(steps)
    margin = 80
    gap = 34
    box_w = (1600 - margin * 2 - gap * (n - 1)) / n
    y1, y2 = 210, 355
    centers = []
    for i, (head, body) in enumerate(steps):
        x1 = int(margin + i * (box_w + gap))
        x2 = int(x1 + box_w)
        d.rounded_rectangle([x1, y1, x2, y2], radius=18, outline="black", width=3, fill="white")
        d.rectangle([x1, y1, x2, y1 + 48], outline="black", width=2, fill="white")
        draw_text_center(d, (x1, y1, x2, y1 + 48), head, FONT_HEAD)
        draw_text_center(d, (x1 + 8, y1 + 54, x2 - 8, y2 - 10), body, FONT_BODY)
        centers.append(((x1 + x2) // 2, (y1 + y2) // 2, x1, x2))
        if i > 0:
            arrow(d, (centers[i - 1][3] + 5, (y1 + y2) // 2), (x1 - 10, (y1 + y2) // 2))
    d.line([margin, 455, 1600 - margin, 455], fill="black", width=2)
    for cx, _, _, _ in centers:
        d.line([cx, 448, cx, 462], fill="black", width=2)
    labels = ["输入", "处理", "融合", "判断", "输出", "评价"][:n]
    for i, (cx, _, _, _) in enumerate(centers):
        draw_text_center(d, (cx - 60, 475, cx + 60, 520), labels[i] if i < len(labels) else "", FONT_SMALL)
    img.save(path)


def draw_closed_loop(path, spec):
    img, d = base_canvas(spec["title"], (1600, 660))
    steps = spec["steps"]
    subs = spec["sub"]
    margin = 95
    gap = 42
    n = len(steps)
    box_w = (1600 - margin * 2 - gap * (n - 1)) / n
    y1, y2 = 175, 320
    centers = []
    for i, ((head, body), sub) in enumerate(zip(steps, subs)):
        x1 = int(margin + i * (box_w + gap))
        x2 = int(x1 + box_w)
        d.rounded_rectangle([x1, y1, x2, y2], radius=16, outline="black", width=3, fill="white")
        draw_text_center(d, (x1, y1 + 12, x2, y1 + 52), head, FONT_HEAD)
        d.line([x1 + 18, y1 + 62, x2 - 18, y1 + 62], fill="black", width=2)
        draw_text_center(d, (x1 + 10, y1 + 72, x2 - 10, y2 - 12), body, FONT_BODY)
        sx1, sx2 = x1 + 8, x2 - 8
        sy1, sy2 = 390, 475
        d.rectangle([sx1, sy1, sx2, sy2], outline="black", width=2)
        draw_text_center(d, (sx1 + 5, sy1 + 5, sx2 - 5, sy2 - 5), sub, FONT_BODY)
        d.line([(x1 + x2) // 2, y2, (x1 + x2) // 2, sy1], fill="black", width=2)
        centers.append((x1, x2))
        if i > 0:
            arrow(d, (centers[i - 1][1] + 6, (y1 + y2) // 2), (x1 - 10, (y1 + y2) // 2))
    # Bottom feedback line keeps the loop visible without crossing content boxes.
    feedback_y = 565
    first_cx = (centers[0][0] + centers[0][1]) // 2
    last_cx = (centers[-1][0] + centers[-1][1]) // 2
    d.line([last_cx, 475, last_cx, feedback_y], fill="black", width=3)
    d.line([last_cx, feedback_y, first_cx, feedback_y], fill="black", width=3)
    arrow(d, (first_cx, feedback_y), (first_cx, 320), width=3)
    draw_text_center(d, (650, feedback_y - 42, 950, feedback_y - 6), "闭环反馈", FONT_SMALL)
    img.save(path)


def draw_model(path, spec):
    img, d = base_canvas(spec["title"], (1600, 680))
    columns = [("输入指标", spec["left"]), ("计算过程", spec["middle"]), ("输出结果", spec["right"])]
    xs = [(100, 470), (615, 985), (1130, 1500)]
    for col_idx, ((title, items), (x1, x2)) in enumerate(zip(columns, xs)):
        d.rectangle([x1, 155, x2, 205], outline="black", width=3)
        draw_text_center(d, (x1, 155, x2, 205), title, FONT_HEAD)
        item_h = 54
        start_y = 245
        for i, item in enumerate(items):
            y = start_y + i * (item_h + 16)
            d.rounded_rectangle([x1 + 20, y, x2 - 20, y + item_h], radius=12, outline="black", width=2, fill="white")
            draw_text_center(d, (x1 + 24, y, x2 - 24, y + item_h), item, FONT_BODY)
    for source_y in [272, 342, 412, 482]:
        arrow(d, (470, source_y), (615, 360), width=2)
    for target_y in [272, 342, 412, 482]:
        arrow(d, (985, 360), (1130, target_y), width=2)
    img.save(path)


def draw_layer(path, spec):
    layers = spec["layers"]
    img, d = base_canvas(spec["title"], (1600, 760))
    y = 150
    layer_h = 160
    left_label_w = 145
    for layer_name, items in layers:
        d.rectangle([80, y, 1520, y + layer_h], outline="black", width=2)
        d.rectangle([80, y, 80 + left_label_w, y + layer_h], outline="black", width=2)
        draw_text_center(d, (80, y, 80 + left_label_w, y + layer_h), layer_name, FONT_HEAD)
        n = len(items)
        gap = 22
        x_start = 250
        available = 1520 - x_start - 24
        box_w = min(170, (available - gap * (n - 1)) / n)
        total = box_w * n + gap * (n - 1)
        x = x_start + (available - total) / 2
        for item in items:
            d.rounded_rectangle([int(x), y + 55, int(x + box_w), y + 110], radius=10, outline="black", width=2, fill="white")
            draw_text_center(d, (int(x), y + 55, int(x + box_w), y + 110), item, FONT_BODY)
            x += box_w + gap
        y += layer_h + 45
    for i in range(len(layers) - 1):
        y1 = 150 + i * (layer_h + 45) + layer_h
        y2 = y1 + 45
        arrow(d, (800, y1 + 4), (800, y2 - 4), width=3)
    img.save(path)


def draw_er(path, spec):
    img, d = base_canvas(spec["title"], (1600, 800))
    positions = [
        (100, 170, 430, 330),
        (635, 170, 965, 330),
        (1170, 170, 1500, 330),
        (100, 500, 430, 660),
        (635, 500, 965, 660),
        (1170, 500, 1500, 660),
    ]
    centers = []
    for (name, fields), box in zip(spec["entities"], positions):
        x1, y1, x2, y2 = box
        d.rectangle(box, outline="black", width=3)
        d.rectangle([x1, y1, x2, y1 + 44], outline="black", width=2)
        draw_text_center(d, (x1, y1, x2, y1 + 44), name, FONT_HEAD)
        for i, f in enumerate(fields):
            d.text((x1 + 24, y1 + 58 + i * 24), f, fill="black", font=FONT_SMALL)
        centers.append(((x1 + x2) // 2, (y1 + y2) // 2))
    for a, b in spec["links"]:
        arrow(d, centers[a], centers[b], width=2)
    img.save(path)


def make_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {}
    for caption, spec in DIAGRAMS.items():
        path = ASSET_DIR / (caption.split()[0].replace(".", "_") + ".png")
        if spec["type"] == "pipeline":
            draw_pipeline(path, spec)
        elif spec["type"] == "closed_loop":
            draw_closed_loop(path, spec)
        elif spec["type"] == "model":
            draw_model(path, spec)
        elif spec["type"] == "layer":
            draw_layer(path, spec)
        elif spec["type"] == "er":
            draw_er(path, spec)
        paths[caption] = path
    return paths


def paragraph_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def replace_picture_before_caption(doc, caption, image_path):
    body = doc._body._element
    paragraphs = list(doc.paragraphs)
    target = next((p for p in paragraphs if p.text.strip() == caption), None)
    if target is None:
        return False
    children = list(body)
    pos = children.index(target._p)
    if pos > 0 and children[pos - 1].tag == qn("w:p") and children[pos - 1].xpath(".//w:drawing"):
        body.remove(children[pos - 1])
    pic_p = target._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Inches(5.85))
    target._p.addprevious(pic_p._p)
    for run in target.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10)
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def main():
    paths = make_diagrams()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    replaced = 0
    for caption, path in paths.items():
        if replace_picture_before_caption(doc, caption, path):
            replaced += 1
    doc.save(OUT)
    checked = Document(OUT)
    captions = [p.text.strip() for p in checked.paragraphs if re.match(r"^[图表]\d+(\.\d+)+", p.text.strip())]
    print(f"saved={OUT}")
    print(f"replaced={replaced}")
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")
    print(f"captions={len(captions)}")


if __name__ == "__main__":
    main()
