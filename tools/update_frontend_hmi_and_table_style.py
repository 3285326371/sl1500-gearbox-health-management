from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def set_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_para(paragraph, first_line=True, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.keep_together = True
    fmt.keep_with_next = keep_with_next
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    else:
        fmt.first_line_indent = None
    for run in paragraph.runs:
        set_font(run, 12)


def paragraph_after(paragraph, text="", style=None):
    p = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(p._p)
    return p


def paragraph_after_element(doc, element, text="", style=None):
    p = doc.add_paragraph(text, style=style)
    element.addnext(p._p)
    return p


def table_after(doc, paragraph, rows, cols):
    tbl = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(tbl._tbl)
    return tbl


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"Paragraph not found: {text}")


def move_block_before(doc, start_text, end_before_text, before_text):
    body = doc.element.body
    start = find_para(doc, start_text)._p
    end_before = find_para(doc, end_before_text)._p
    before = find_para(doc, before_text)._p

    children = list(body)
    si = children.index(start)
    ei = children.index(end_before)
    block = children[si:ei]
    for el in block:
        body.remove(el)

    children = list(body)
    bi = children.index(before)
    for offset, el in enumerate(block):
        body.insert(bi + offset, el)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in kwargs.items():
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    specs = {
        "top": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "bottom": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "left": {"val": "nil"},
        "right": {"val": "nil"},
        "insideH": {"val": "dotted", "sz": "4", "space": "0", "color": "BFBFBF"},
        "insideV": {"val": "single", "sz": "4", "space": "0", "color": "BFBFBF"},
    }
    for edge, attrs in specs.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def remove_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def style_table_like_sample(table):
    set_table_borders(table)
    table.autofit = True
    for ri, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            remove_shading(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = None
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, 10.5, bold=(ri == 0))
            if ri == 0:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
                )


def insert_frontend_design_details(doc):
    heading = find_para(doc, "4.5 前端页面设计")
    first = heading._p.getnext()
    if first is not None and first.xpath(".//w:t[contains(., '4.5.1 前端页面结构设计')]"):
        return

    p451 = paragraph_after(heading, "4.5.1 前端页面结构设计", style="Heading 3")
    set_para(p451, first_line=False, keep_with_next=True)
    p451._p.addnext(first)

    intro2 = find_para(doc, "故障诊断页面提供场景选择和诊断按钮，运行后展示疑似故障类型、风险等级、关联部件、置信度、健康评分、RUL、特征值、风险因子、诊断依据、建议措施和工单信息。故障数据页面展示历史故障记录，报告页面展示健康评估结果，问答页面支持输入自然语言问题并获得排查建议。HMI 页面则以单机齿轮箱为对象，提供运行总览、状态寿命、透明模型、数据采集、测点信号、阈值设置、故障代码和临界警报等功能。")
    caption = paragraph_after(intro2, "表4.5 前端页面与交互功能对应")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(caption, first_line=False, keep_with_next=True)
    for r in caption.runs:
        set_font(r, 10.5, bold=True)

    table = table_after(doc, caption, 8, 4)
    data = [
        ["前端页面", "交互入口", "主要展示内容", "前后端对应"],
        ["风场总览", "左侧导航-健康总览", "风场统计、多机组卡片、健康矩阵", "windfarm_route.py"],
        ["单机详情", "机组卡片下钻", "油温、振动、RUL、部件状态和历史故障", "windfarm_route.py"],
        ["故障诊断", "智能诊断-故障诊断", "场景选择、诊断结果、依据和建议措施", "data_route.py、ml_models.py"],
        ["故障数据", "智能诊断-故障数据", "故障记录、筛选、处理状态和CSV导出", "data_route.py、database.py"],
        ["AI问答", "智能诊断-运维辅助问答", "知识库检索、实时工况分析和排查建议", "qa_route.py、rag_service.py"],
        ["健康报告", "智能诊断-健康报告", "健康评分、关键指标、维护计划和报告历史", "report_route.py"],
        ["HMI界面", "单机详情-HMI控制台", "设备对象、透明模型、阈值、测点和临界报警", "hmi_page.html、data_route.py"],
    ]
    for r, row in enumerate(data):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    style_table_like_sample(table)

    p452 = paragraph_after_element(doc, table._element, "4.5.2 人机交互界面设计", style="Heading 3")
    set_para(p452, first_line=False, keep_with_next=True)
    body = paragraph_after(
        p452,
        "人机交互界面设计遵循“总览发现异常、详情定位原因、诊断形成结论、报告完成闭环”的操作路径。页面左侧采用稳定导航分组，减少用户在不同功能间切换时的认知负担；顶部区域显示当前页面名称、登录用户和通知状态，使运维人员能够确认当前操作对象。关键指标采用卡片和颜色状态表达，异常指标使用醒目色提示，但不改变页面整体结构，便于在长时间监控场景下保持可读性。",
    )
    set_para(body)
    body2 = paragraph_after(
        body,
        "HMI 页面更强调现场设备交互。该页面围绕单台齿轮箱展开，将齿轮箱油温、高速轴承温度、振动RMS、润滑油等级、健康评分和报警状态集中到同一视图，并预留透明模型、测点信号、阈值设置和临界报警入口。相比普通管理页面，HMI 界面能够体现设备对象、运行状态和操作反馈之间的对应关系，是系统面向现场监控应用的重要补充。",
    )
    set_para(body2)


def fix_section_order_and_hmi_heading(doc):
    try:
        move_block_before(doc, "单机齿轮箱详情页面如图5.7所示。页面从风场总览下钻到单台机组，集中展示齿轮箱油温、高速轴承温度、振动RMS、油液NAS等级、健康评分、RUL、复检周期和历史故障记录。", "第6章 系统测试与结果分析", "5.4 故障代码库与数据管理实现")
    except RuntimeError:
        pass

    detail_heading = find_para(doc, "5.3.2 单机详情页面")
    if detail_heading._p.getprevious() is not None:
        prev_text = "".join(t.text or "" for t in detail_heading._p.getprevious().xpath(".//w:t"))
        if prev_text.startswith("单机齿轮箱详情页面"):
            prev = detail_heading._p.getprevious()
            detail_heading._p.addnext(prev)

    hmi_intro = find_para(doc, "齿轮箱HMI页面如图5.8所示。该页面面向现场监控场景，用更加接近工业界面的方式展示单机齿轮箱状态和关键指标。")
    prev_text = "".join(t.text or "" for t in hmi_intro._p.getprevious().xpath(".//w:t")) if hmi_intro._p.getprevious() is not None else ""
    if "5.3.3 HMI人机交互页面" not in prev_text:
        h = paragraph_after(find_para(doc, "从健康管理角度看，单机详情页面能够把不同来源的数据组织到同一个设备视图中。油温和轴承温度用于判断热状态，振动RMS用于判断机械冲击，油液NAS用于判断润滑污染，RUL和复检周期用于形成维护计划。这样的页面结构有助于运维人员快速理解设备状态。"), "5.3.3 HMI人机交互页面", style="Heading 3")
        set_para(h, first_line=False, keep_with_next=True)
        h._p.addnext(hmi_intro._p)

    # Keep numbering/captions consistent after adding the HMI subsection.
    replacements = {
        "图4.2 齿轮箱智能健康管理流程图": "图4.3 齿轮箱智能健康管理流程图",
        "系统诊断流程如图4.1所示。": "系统诊断流程如图4.3所示。",
    }
    for p in doc.paragraphs:
        text = p.text
        for old, new in replacements.items():
            if old in text:
                p.text = text.replace(old, new)
                for run in p.runs:
                    set_font(run, 12)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_前端交互修改前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    insert_frontend_design_details(doc)
    fix_section_order_and_hmi_heading(doc)
    for table in doc.tables:
        style_table_like_sample(table)
    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
