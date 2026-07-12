from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))


def set_para(p, text):
    p.clear()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(15)
    r.bold = True


def main():
    doc = Document(TEMPLATE)
    set_para(doc.paragraphs[9], "姓    名  ____________________")
    set_para(doc.paragraphs[10], "学    号  ____________________")
    set_para(doc.paragraphs[13], "班    级  ____________________")
    set_para(doc.paragraphs[14], "指导教师  ____________________")
    set_para(doc.paragraphs[15], "辅导教师  ____________________")
    doc.save(TEMPLATE)
    print(TEMPLATE)


if __name__ == "__main__":
    main()
