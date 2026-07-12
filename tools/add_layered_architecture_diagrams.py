from __future__ import annotations

import io
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_图表黑白统一版.docx")
OUT = SRC.with_name("毕业设计正文模板_结构图增强版.docx")
ASSET_DIR = SRC.with_name("结构图增强版_assets")


DIAGRAMS = [
    ("图5.12 系统整体分层架构图", "整体分层架构", [
        ("表示层", ["风场总览", "单机详情", "实时监测", "故障诊断", "健康报告", "AI问答", "HMI控制台"]),
        ("业务逻辑层", ["认证路由", "风场接口", "数据接口", "诊断服务", "报告服务", "问答服务", "设置接口"]),
        ("数据层", ["SQLite数据库", "风机数据摘要", "模拟采集数据", "故障代码库", "系统配置"]),
    ]),
    ("图5.13 风场总览与单机下钻流程图", "风场总览与单机下钻", [
        ("入口层", ["登录验证", "角色识别", "进入健康总览"]),
        ("总览层", ["统计指标", "机组矩阵", "状态图例", "告警数量"]),
        ("单机层", ["关键指标", "部件状态", "历史故障", "HMI入口", "报告入口"]),
    ]),
    ("图5.14 多源数据采集与融合流程图", "多源数据采集与融合", [
        ("数据来源", ["油温测点", "振动测点", "油液NAS", "功率风速", "故障统计"]),
        ("处理过程", ["字段清洗", "异常值处理", "均值滤波", "机组编号统一", "状态快照生成"]),
        ("输出对象", ["风场总览", "单机详情", "趋势图表", "诊断输入", "健康报告"]),
    ]),
    ("图5.15 齿轮箱故障诊断流程图", "齿轮箱故障诊断流程", [
        ("输入", ["选择机组", "诊断场景", "实时状态", "振动波形"]),
        ("算法处理", ["特征提取", "风险评分", "规则匹配", "严重度判断"]),
        ("结果", ["故障类型", "置信度", "健康评分", "RUL", "维护建议"]),
    ]),
    ("图5.16 风险评分与健康评估模型图", "风险评分与健康评估模型", [
        ("指标层", ["RMS", "峭度", "峰值因子", "包络峰值", "油温", "油液NAS"]),
        ("风险层", ["振动风险", "冲击风险", "温度风险", "润滑风险", "历史风险"]),
        ("决策层", ["综合风险", "健康评分", "RUL估计", "复检周期", "工单建议"]),
    ]),
    ("图5.17 数据库实体关系图", "数据库实体关系", [
        ("用户与权限", ["User", "username", "password_hash", "role"]),
        ("设备与诊断", ["TurbineAsset", "FaultRecord", "DiagnosisResult", "HealthReport"]),
        ("配置与闭环", ["SystemConfig", "FaultClosure", "threshold", "handler"]),
    ]),
    ("图5.18 后端接口与蓝图结构图", "后端接口与蓝图结构", [
        ("入口", ["app.py", "Flask应用", "蓝图注册"]),
        ("接口层", ["auth_bp", "data_bp", "windfarm_bp", "report_bp", "qa_bp", "settings_bp", "ops_bp"]),
        ("服务层", ["DataAcquisition", "WindRepository", "DiagnosisModel", "RagService", "ReportService"]),
    ]),
    ("图5.19 前端页面模块结构图", "前端页面模块结构", [
        ("导航分组", ["健康监测区", "智能诊断区", "系统配置区"]),
        ("页面模块", ["总览", "详情", "实时监测", "故障数据", "代码库", "报告", "问答"]),
        ("交互输出", ["指标卡片", "趋势曲线", "诊断卡片", "记录表格", "跳转按钮"]),
    ]),
    ("图5.20 HMI运维控制台结构图", "HMI运维控制台结构", [
        ("权限入口", ["二次登录", "角色级别", "机组选择"]),
        ("运行监控", ["油温", "振动RMS", "健康评分", "RUL", "透明模型"]),
        ("现场操作", ["阈值设置", "故障代码", "临界报警", "测点信号", "数据采集"]),
    ]),
    ("图5.21 故障代码库与记录管理关系图", "故障代码库与记录管理", [
        ("故障代码库", ["GBX-001", "GBX-002", "GBX-003", "SCADA动态码"]),
        ("诊断记录", ["机组编号", "故障类型", "严重程度", "置信度", "处理状态"]),
        ("数据管理", ["筛选查询", "CSV导出", "详情查看", "闭环处理"]),
    ]),
    ("图5.22 健康报告生成流程图", "健康报告生成流程", [
        ("数据准备", ["当前状态", "历史故障", "趋势指标", "配置阈值"]),
        ("报告生成", ["健康评分", "关键指标", "诊断结论", "维护计划"]),
        ("输出使用", ["页面展示", "报告归档", "复检建议", "工单依据"]),
    ]),
    ("图5.23 AI运维问答处理流程图", "AI运维问答处理流程", [
        ("提问输入", ["用户问题", "当前机组", "实时状态", "输出要求"]),
        ("知识处理", ["本地知识库", "实时上下文", "模型主答", "规则回退"]),
        ("回答输出", ["结论", "依据", "排查步骤", "工单判断", "追问建议"]),
    ]),
    ("图5.24 参数设置影响链路图", "参数设置影响链路", [
        ("参数配置", ["油温阈值", "振动阈值", "油液阈值", "报警策略"]),
        ("业务影响", ["实时颜色", "风险因子", "故障等级", "报告建议"]),
        ("权限控制", ["管理员", "工程师", "普通运维", "操作日志"]),
    ]),
    ("图5.25 运维工单闭环流程图", "运维工单闭环流程", [
        ("发现异常", ["监测告警", "诊断触发", "风险升高"]),
        ("处理过程", ["生成工单", "现场复核", "检修处理", "复测验收"]),
        ("闭环沉淀", ["状态更新", "记录归档", "报告输出", "知识补充"]),
    ]),
]


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, fnt, line_gap=6):
    x1, y1, x2, y2 = box
    max_chars = max(4, int((x2 - x1) // (fnt.size * 0.95)))
    lines = []
    for part in str(text).split("\n"):
        line = ""
        for ch in part:
            line += ch
            if len(line) >= max_chars:
                lines.append(line)
                line = ""
        if line:
            lines.append(line)
    heights = [draw.textbbox((0, 0), ln, font=fnt)[3] for ln in lines]
    total = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for ln, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), ln, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), ln, fill="black", font=fnt)
        y += h + line_gap


def arrow(draw, start, end):
    x1, y1 = start
    x2, y2 = end
    draw.line([x1, y1, x2, y2], fill="black", width=2)
    if y2 >= y1:
        pts = [(x2, y2), (x2 - 8, y2 - 14), (x2 + 8, y2 - 14)]
    else:
        pts = [(x2, y2), (x2 - 8, y2 + 14), (x2 + 8, y2 + 14)]
    draw.polygon(pts, fill="black")


def draw_layered_diagram(path: Path, title: str, layers):
    img = Image.new("RGB", (1600, 1050), "white")
    d = ImageDraw.Draw(img)
    title_font = font(34, True)
    layer_font = font(26, True)
    box_font = font(22)
    small_font = font(18)

    d.rectangle([30, 30, 1570, 1020], outline="black", width=2)
    draw_centered(d, (120, 45, 1480, 90), title, title_font)

    top = 130
    layer_h = 230
    gap = 70
    side_x1, side_x2 = 70, 130
    content_x1, content_x2 = 170, 1500
    prev_centers = []
    for idx, (layer_name, boxes) in enumerate(layers):
        y1 = top + idx * (layer_h + gap)
        y2 = y1 + layer_h
        d.rectangle([side_x1, y1, side_x2, y2], outline="black", width=1)
        # vertical label
        chars = list(layer_name)
        cy = y1 + 18
        for ch in chars:
            bbox = d.textbbox((0, 0), ch, font=layer_font)
            d.text((side_x1 + (side_x2 - side_x1 - (bbox[2] - bbox[0])) / 2, cy), ch, fill="black", font=layer_font)
            cy += 31

        d.rectangle([content_x1, y1, content_x2, y2], outline="black", width=1)
        main_box = (content_x1 + 25, y1 + 18, content_x2 - 25, y1 + 68)
        d.rectangle(main_box, outline="black", width=2)
        draw_centered(d, main_box, layer_name, layer_font)

        count = len(boxes)
        box_w = min(175, (content_x2 - content_x1 - 80) // max(count, 1))
        spacing = (content_x2 - content_x1 - count * box_w) / (count + 1)
        centers = []
        for j, label in enumerate(boxes):
            x1 = content_x1 + spacing + j * (box_w + spacing)
            b = (int(x1), y1 + 115, int(x1 + box_w), y1 + 178)
            d.rectangle(b, outline="black", width=2)
            draw_centered(d, b, label, box_font)
            centers.append(((b[0] + b[2]) // 2, b[1]))
            arrow(d, ((main_box[0] + main_box[2]) // 2, main_box[3]), ((b[0] + b[2]) // 2, b[1]))
        if prev_centers:
            target_y = main_box[1]
            source_y = top + (idx - 1) * (layer_h + gap) + layer_h - 52
            for pc in prev_centers:
                arrow(d, (pc[0], source_y), ((main_box[0] + main_box[2]) // 2, target_y))
        prev_centers = [(c[0], y2) for c in centers]

    d.text((120, 985), "说明：图中箭头表示数据调用、接口访问或业务流转关系。", fill="black", font=small_font)
    img.save(path)


def paragraph_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def remove_old_extra_figures(doc: Document):
    body = doc._body._element
    children = list(body)
    targets = {f"图5.{i}" for i in range(12, 26)}
    for idx, child in list(enumerate(children)):
        if child.tag != qn("w:p"):
            continue
        text = paragraph_text(child)
        if any(text.startswith(prefix + " ") for prefix in targets):
            # Remove image paragraph immediately before this caption when present.
            current_children = list(body)
            pos = current_children.index(child)
            if pos > 0 and current_children[pos - 1].tag == qn("w:p") and current_children[pos - 1].xpath(".//w:drawing"):
                body.remove(current_children[pos - 1])
            if child.getparent() is not None:
                body.remove(child)


def add_picture_after(paragraph, caption, img_path):
    pic_p = paragraph._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(img_path), width=Inches(5.8))
    paragraph._p.addnext(pic_p._p)

    cap_p = paragraph._parent.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = "宋体"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap_run.font.size = Pt(10)
    pic_p._p.addnext(cap_p._p)
    return cap_p


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = docx_element("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def docx_element(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for caption, title, layers in DIAGRAMS:
        filename = caption.split()[0].replace(".", "_") + ".png"
        draw_layered_diagram(ASSET_DIR / filename, title, layers)

    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    remove_old_extra_figures(doc)
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "图5.11 系统参数设置页面")
    cursor = anchor
    for caption, _title, _layers in DIAGRAMS:
        filename = caption.split()[0].replace(".", "_") + ".png"
        cursor = add_picture_after(cursor, caption, ASSET_DIR / filename)
    set_update_fields(doc)
    doc.save(OUT)

    checked = Document(OUT)
    captions = [p.text.strip() for p in checked.paragraphs if re.match(r"^[图表]\d+(\.\d+)+", p.text.strip())]
    print(f"saved={OUT}")
    print(f"paragraphs={sum(1 for p in checked.paragraphs if p.text.strip())}")
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")
    print(f"captions={len(captions)}")


if __name__ == "__main__":
    main()
