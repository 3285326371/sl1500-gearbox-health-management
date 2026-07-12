from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_40页目录同步.docx"

LINES = [
    ("摘要", "I", 0), ("ABSTRACT", "II", 0),
    ("第1章 绪论", "1", 0), ("1.1 研究背景与意义", "1", 1), ("1.2 国内外研究现状", "1", 1), ("1.3 本文研究内容", "2", 1), ("1.4 技术路线", "2", 1),
    ("第2章 系统需求分析", "3", 0), ("2.1 业务需求分析", "3", 1), ("2.2 功能需求分析", "3", 1), ("2.3 非功能需求分析", "4", 1), ("2.4 数据需求分析", "4", 1), ("2.5 需求约束与边界分析", "4", 1), ("2.6 关键数据字段设计", "5", 1),
    ("第3章 相关技术与理论基础", "7", 0), ("3.1 齿轮箱典型故障机理", "7", 1), ("3.2 多源状态监测技术", "7", 1), ("3.3 振动特征提取方法", "7", 1), ("3.4 风险评分与健康评估方法", "7", 1), ("3.5 Web系统开发技术", "8", 1), ("3.6 健康评分与RUL估计原理", "8", 1), ("3.7 智能问答与知识库技术", "9", 1),
    ("第4章 系统总体设计", "10", 0), ("4.1 系统架构设计", "10", 1), ("4.2 数据库设计", "10", 1), ("4.3 后端接口设计", "11", 1), ("4.4 诊断流程设计", "12", 1), ("4.5 前端页面设计", "13", 1),
    ("第5章 系统详细实现", "14", 0), ("5.1 数据采集与融合模块实现", "14", 1), ("5.2 故障诊断算法模块实现", "15", 1), ("5.3 风场总览与单机详情实现", "17", 1), ("5.4 故障代码库与数据管理实现", "18", 1), ("5.5 AI运维问答模块实现", "19", 1), ("5.6 健康报告与工单闭环实现", "20", 1),
    ("第6章 系统测试与结果分析", "24", 0), ("6.1 测试环境", "24", 1), ("6.2 功能测试", "24", 1), ("6.3 诊断场景测试分析", "24", 1), ("6.4 结果评价", "25", 1), ("6.5 测试用例设计", "25", 1), ("6.6 典型诊断场景结果", "26", 1), ("6.7 系统运行效果分析", "26", 1), ("6.8 单机详情与HMI验证", "27", 1), ("6.9 非功能测试与可用性分析", "27", 1), ("6.10 存在问题与改进方向", "28", 1),
    ("第7章 总结与展望", "29", 0), ("参考文献", "30", 0), ("致谢", "31", 0), ("附录A 系统主要接口", "32", 0), ("附录B 核心代码片段与接口说明", "33", 0), ("附录C 系统部署与运行说明", "35", 0),
]


def set_font(run, size=8.7, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def fmt(p, level):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.34 if level else 0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(10.6)
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def main():
    doc = Document(TEMPLATE)
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() in {"目  录", "目□□录"}) + 1
    end = next(i for i in range(start, len(doc.paragraphs)) if doc.paragraphs[i].text.strip().replace(" ", "") == "第1章绪论")
    title = doc.paragraphs[start - 1]
    title.paragraph_format.space_after = Pt(5)
    for r in title.runs:
        set_font(r, size=16, bold=True)
    for idx, (text, page, level) in enumerate(LINES):
        p = doc.paragraphs[start + idx]
        p.clear()
        fmt(p, level)
        r = p.add_run(text)
        set_font(r, bold=False)
        p.add_run("\t")
        r2 = p.add_run(page)
        set_font(r2)
    for i in range(start + len(LINES), end):
        doc.paragraphs[i].clear()
        doc.paragraphs[i].paragraph_format.line_spacing = Pt(1)
        doc.paragraphs[i].paragraph_format.space_after = Pt(0)
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(TMP_OUT)


if __name__ == "__main__":
    main()
