from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path.cwd()
TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_加图片代码前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = ROOT / "docs" / "_tmp_毕业设计正文模板_加图片代码.docx"
FIG_DIR = ROOT / "docs" / "figures"
ARCH_IMG = FIG_DIR / "system_function_architecture.png"


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def wrapped_lines(draw, text, box_w, fnt):
    lines, cur = [], ""
    for ch in text:
        trial = cur + ch
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= box_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = ch
    if cur:
        lines.append(cur)
    return lines


def center_text(draw, box, text, fnt, fill=(0, 0, 0)):
    x1, y1, x2, y2 = box
    lines = wrapped_lines(draw, text, x2 - x1 - 30, fnt)
    lh = fnt.size + 8
    total = lh * len(lines)
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=fnt, fill=fill)
        y += lh


def draw_arrow(draw, start, end, color=(54, 65, 82), width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    direction = 1 if x2 >= x1 else -1
    draw.polygon([(x2, y2), (x2 - direction * 18, y2 - 9), (x2 - direction * 18, y2 + 9)], fill=color)


def create_architecture_image():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 1050
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_f = font(44, True)
    head_f = font(30, True)
    node_f = font(25, True)
    small_f = font(22)

    title = "系统功能架构图"
    tb = d.textbbox((0, 0), title, font=title_f)
    d.text(((w - tb[2] + tb[0]) / 2, 40), title, font=title_f, fill=(0, 0, 0))

    layers = [
        (110, 140, 1690, 245, "表现层", ["风场总览", "单机详情", "实时监测", "故障诊断", "健康报告", "运维问答"], (237, 246, 255)),
        (110, 320, 1690, 425, "接口层", ["REST接口", "SSE数据流", "报告接口", "问答接口", "权限接口"], (239, 248, 240)),
        (110, 500, 1690, 645, "业务服务层", ["数据采集与融合", "振动特征提取", "故障规则库", "健康评分/RUL", "记录与工单", "知识库检索"], (255, 248, 232)),
        (110, 730, 1690, 860, "数据层", ["SQLite业务库", "风机统计数据", "故障代码库", "运维知识库", "报告历史"], (247, 242, 255)),
    ]
    for x1, y1, x2, y2, name, items, fill in layers:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=(80, 110, 145), width=4)
        d.rounded_rectangle((x1 + 20, y1 + 18, x1 + 190, y2 - 18), radius=18, fill=(255, 255, 255), outline=(80, 110, 145), width=3)
        center_text(d, (x1 + 20, y1 + 18, x1 + 190, y2 - 18), name, head_f)
        gap = 18
        node_w = int((x2 - x1 - 250 - gap * (len(items) - 1)) / len(items))
        nx = x1 + 220
        for item in items:
            d.rounded_rectangle((nx, y1 + 24, nx + node_w, y2 - 24), radius=16, fill=(255, 255, 255), outline=(120, 145, 170), width=2)
            center_text(d, (nx, y1 + 24, nx + node_w, y2 - 24), item, node_f)
            nx += node_w + gap

    draw_arrow(d, (900, 245), (900, 320))
    draw_arrow(d, (900, 425), (900, 500))
    draw_arrow(d, (900, 645), (900, 730))
    d.rounded_rectangle((180, 915, 1620, 985), radius=18, fill=(248, 248, 248), outline=(180, 180, 180), width=2)
    note = "说明：系统以 Flask 后端为核心，通过接口层向前端提供实时数据、诊断结果和报告能力，底层数据由业务数据库、风机统计数据和知识库共同支撑。"
    center_text(d, (200, 920, 1600, 980), note, small_f, fill=(40, 40, 40))
    im.save(ARCH_IMG)


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def insert_before(anchor, text="", style=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def add_caption(anchor, text):
    p = insert_before(anchor, text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_run_font(r, size=10.5)
    return p


def shade_cell(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(anchor, title, code):
    title_p = insert_before(anchor, title)
    for r in title_p.runs:
        set_run_font(r, east_asia="黑体", ascii_font="Arial", size=10.5, bold=True)
    table = anchor._parent.add_table(rows=1, cols=1, width=Inches(6.0))
    # Move the new table before the anchor paragraph.
    anchor._p.addprevious(table._tbl)
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F4F4")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(11)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(32, 32, 32)
    note = insert_before(anchor, "")
    note.paragraph_format.space_after = Pt(6)
    return table


def add_architecture_figure(doc):
    if any("图4.2 系统功能架构图" in p.text for p in doc.paragraphs):
        return
    anchor_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "4.1 系统架构设计")
    after = doc.paragraphs[anchor_idx + 1]
    lead = insert_before(
        after,
        "系统功能架构如图4.2所示。系统按表现层、接口层、业务服务层和数据层组织，前端负责状态展示与交互操作，后端负责数据处理、诊断计算、报告生成和知识问答，数据层为诊断模型和页面展示提供基础支撑。",
    )
    for r in lead.runs:
        set_run_font(r)
    pic_p = insert_before(after, "")
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(ARCH_IMG), width=Inches(5.8))
    add_caption(after, "图4.2 系统功能架构图")


def add_code_snippets(doc):
    if any("代码清单5.1" in p.text for p in doc.paragraphs):
        return

    anchor_51_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "5.1 数据采集与融合模块实现")
    after_51 = doc.paragraphs[anchor_51_idx + 2]
    add_code_block(
        after_51,
        "代码清单5.1 振动信号特征提取核心代码",
        """
def process_signals(signal_data):
    arr = np.array(signal_data, dtype=float)
    rms = float(np.sqrt(np.mean(arr ** 2)))
    peak = float(np.max(np.abs(arr)))
    mean = float(np.mean(arr))
    std = float(np.std(arr)) or 1e-6
    centered = arr - mean
    kurtosis = float(np.mean(centered ** 4) / (std ** 4))
    crest_factor = float(peak / rms) if rms > 0 else 1.0
    impulse_factor = float(peak / (np.mean(np.abs(arr)) or 1e-6))
    envelope = get_envelope(arr.tolist())
    envelope_peak = float(max(envelope)) if envelope else 0.0
    return {
        "rms": round(rms, 4),
        "kurtosis": round(kurtosis, 4),
        "peak": round(peak, 4),
        "crest_factor": round(crest_factor, 4),
        "impulse_factor": round(impulse_factor, 4),
        "envelope_peak": round(envelope_peak, 4),
    }
""",
    )

    anchor_52_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "5.2 故障诊断算法模块实现")
    after_52 = doc.paragraphs[anchor_52_idx + 2]
    add_code_block(
        after_52,
        "代码清单5.2 故障诊断与健康评分调用逻辑",
        """
def run_diagnosis(unit_id="WTG-001", raw_signal=None, status=None):
    raw_signal = raw_signal or generate_demo_signal()
    features = process_signals(raw_signal)
    risk_score, risk_factors, snapshot = score_risk(features, status)
    fault = classify_fault(features, risk_score, risk_factors, status)
    health_score = calculate_health_score(risk_score, fault["severity"])
    rul = predict_rul(health_score, risk_score, fault["severity"])
    explanation = build_explanation(
        features, fault["severity"], rul, risk_score, risk_factors, fault
    )
    return {
        "unit_id": unit_id,
        "features": features,
        "fault_type": fault["type"],
        "health_score": health_score,
        "rul_days": rul,
        "explanation": explanation,
    }
""",
    )

    anchor_55_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() in {"5.5 AI运维问答模块实现", "5.5 AI 运维问答模块实现"})
    after_55 = doc.paragraphs[anchor_55_idx + 2]
    add_code_block(
        after_55,
        "代码清单5.3 前端实时状态刷新逻辑",
        """
async function refreshStatus() {
  const response = await fetch("/api/data/status");
  const status = await response.json();
  document.querySelector("#oilTemp").textContent = status.oil_temp + " ℃";
  document.querySelector("#vibrationRms").textContent =
    status.vibration_rms + " mm/s";
  healthChart.setOption({
    series: [{ data: [{ value: status.health_score, name: "健康评分" }] }]
  });
}
setInterval(refreshStatus, 5000);
refreshStatus();
""",
    )


def main():
    create_architecture_image()
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    add_architecture_figure(doc)
    add_code_snippets(doc)
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")
    print(f"image={ARCH_IMG}")


if __name__ == "__main__":
    main()
