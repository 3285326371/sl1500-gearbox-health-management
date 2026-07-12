from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_目录后分页.docx"


def main():
    doc = Document(TEMPLATE)
    toc_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() in {"目  录", "目□□录"})
    idx = next(
        i for i in range(toc_idx + 1, len(doc.paragraphs))
        if doc.paragraphs[i].text.strip().replace(" ", "") == "第1章绪论"
    )
    prev = doc.paragraphs[idx - 1]
    if 'w:type="page"' not in prev._element.xml:
        prev.add_run().add_break(WD_BREAK.PAGE)
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(TMP_OUT)


if __name__ == "__main__":
    main()
