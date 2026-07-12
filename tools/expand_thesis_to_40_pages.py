from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_扩展40页前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = ROOT / "docs" / "_tmp_毕业设计正文模板_扩展40页.docx"
SHOT_DIR = ROOT / "docs" / "system_screenshots"


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def format_para(p, first_line=True):
    p.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        set_font(r)


def add_before(anchor, text="", first_line=True):
    p = anchor.insert_paragraph_before(text)
    format_para(p, first_line=first_line)
    return p


def add_subheading_before(anchor, text):
    p = anchor.insert_paragraph_before(text)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(0)
    for r in p.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=12, bold=True)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "666666")


def add_table_before(anchor, title, headers, rows):
    cap = anchor.insert_paragraph_before(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    for r in cap.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=10.5, bold=True)
    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Inches(6.0))
    anchor._p.addprevious(table._tbl)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        shade_cell(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    set_font(r, size=9.5)
    spacer = anchor.insert_paragraph_before("")
    spacer.paragraph_format.space_after = Pt(5)
    return table


def add_picture_before(anchor, image_path, caption, intro=None):
    if intro:
        add_before(anchor, intro)
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(5.8))
    cap = anchor.insert_paragraph_before(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        set_font(r, size=10.5)
    spacer = anchor.insert_paragraph_before("")
    spacer.paragraph_format.space_after = Pt(5)


def add_long_paragraphs(anchor, paragraphs):
    for text in paragraphs:
        add_before(anchor, text)


def find_para(doc, text):
    norm = text.replace(" ", "")
    for p in doc.paragraphs:
        if p.text.strip().replace(" ", "") == norm:
            return p
    raise ValueError(text)


def expand_chapter_2(doc):
    anchor = find_para(doc, "第3章 相关技术与理论基础")
    add_subheading_before(anchor, "2.5 需求约束与边界分析")
    add_long_paragraphs(anchor, [
        "本系统面向毕业设计和风电机组齿轮箱智能运维场景，系统需求既包括功能完整性，也包括工程可解释性。功能完整性要求系统能够围绕单台或多台华锐SL1500型双馈式风机完成运行状态展示、异常诊断、故障记录、报告生成和知识问答；工程可解释性要求每一个诊断结果都能够给出数据来源、特征依据、风险因子和维护建议，避免只输出一个难以理解的分类标签。",
        "在数据来源方面，系统当前采用模拟采集数据与风机统计数据文件相结合的方式。模拟采集模块用于持续生成油温、振动、油液、功率和数据质量等实时指标，风机数据仓库用于读取风机汇总数据、故障统计和故障码标签。这样的设计能够在缺少真实在线SCADA接入条件下完成系统功能验证，同时保留将来接入真实CMS、SCADA和油液检测系统的接口边界。",
        "在用户角色方面，系统至少需要区分管理员、运维人员和普通观察用户。管理员负责阈值配置、用户管理和系统参数维护；运维人员关注机组状态、故障诊断、报告和工单；普通观察用户主要查看风场总览和健康趋势。不同角色对数据写入、阈值修改和工单处理的权限不同，因此系统需要具备基本认证和权限控制能力。",
        "在性能约束方面，系统页面需要满足普通浏览器流畅访问，实时监测数据刷新不能造成明显卡顿。对于毕业设计原型而言，重点不在高并发承载，而在接口组织清晰、数据刷新稳定、诊断结果可重复、页面交互完整。系统采用轻量级SQLite数据库和Flask服务，适合本地演示和小规模部署。",
    ])
    add_table_before(anchor, "表2.2 系统需求约束说明", ["约束类型", "约束内容", "设计处理"], [
        ["数据约束", "真实风场在线数据不完整", "采用模拟采集与风机统计文件融合，预留真实接口"],
        ["场景约束", "毕业设计演示环境以本地部署为主", "使用Flask、SQLite和静态前端降低部署复杂度"],
        ["算法约束", "真实故障样本数量有限", "采用规则库和风险评分，保留机器学习模型替换位置"],
        ["可解释性", "运维人员需要知道报警原因", "输出特征值、风险因子、诊断依据和建议措施"],
        ["扩展性", "后续可能接入SCADA/CMS/油液系统", "按接口层、服务层和数据层分离设计"],
    ])
    add_subheading_before(anchor, "2.6 关键数据字段设计")
    add_long_paragraphs(anchor, [
        "齿轮箱健康管理涉及的数据字段较多，如果字段组织不清晰，后续的诊断计算、图表展示和报告生成都会变得混乱。因此本文将数据字段划分为实时状态字段、振动特征字段、油液状态字段、诊断结果字段和运维闭环字段五类。实时状态字段主要体现设备当前工况，振动特征字段用于故障诊断，油液状态字段用于评估润滑和磨损，诊断结果字段用于给出健康结论，运维闭环字段用于保存记录和处理状态。",
        "在系统实现中，实时状态字段通常由数据采集服务生成或读取；振动特征字段由诊断模型计算得到；诊断结果字段由规则库、风险评分和RUL估计共同生成。前端页面不直接参与核心诊断计算，而是通过接口获取结构化结果并展示为仪表盘、表格、卡片和报告。这样的分工使前后端职责更加清晰，也便于后续单独替换诊断算法。",
    ])
    add_table_before(anchor, "表2.3 关键数据字段分类", ["字段类别", "代表字段", "作用"], [
        ["实时状态", "unit_id、timestamp、oil_temp、power", "描述机组当前运行工况"],
        ["振动特征", "rms、kurtosis、crest_factor、envelope_peak", "识别冲击、磨损和啮合异常"],
        ["油液状态", "oil_quality、NAS、particle_level", "反映润滑油污染和磨粒状态"],
        ["诊断结果", "fault_type、severity、health_score、rul_days", "形成健康状态和剩余寿命判断"],
        ["运维闭环", "record_id、status、work_order、advice", "支持故障记录、工单和维护建议"],
    ])


def expand_chapter_3(doc):
    anchor = find_para(doc, "第4章 系统总体设计")
    add_subheading_before(anchor, "3.6 健康评分与RUL估计原理")
    add_long_paragraphs(anchor, [
        "健康评分是将多个监测指标压缩为一个便于运维人员理解的综合指标。本文系统并不直接使用单一阈值判断设备好坏，而是将振动RMS、冲击峭度、峰值因子、包络峰值、齿轮箱油温和油液NAS等级归一化后进行加权计算。这样可以避免某个指标轻微波动就触发严重报警，也能在多个指标同时升高时提高风险敏感性。",
        "RUL估计用于描述设备在当前退化状态下的剩余可运行时间。由于毕业设计原型缺少长期真实退化样本，本文采用基于健康评分、风险因子和故障严重度的启发式估计方法。该方法不是工业现场最终寿命模型，但能够展示健康管理系统中寿命预测的接口形式和计算流程。后续如果获得真实故障样本，可以将该位置替换为SVR、随机森林、LSTM或Transformer等预测模型。",
        "在健康评分计算中，油温和振动代表齿轮箱当前运行负荷与机械冲击状态，油液NAS等级代表润滑和磨粒情况，历史故障记录代表设备长期退化背景。系统将这些因素综合为风险分数，再将风险分数映射为健康评分。健康评分越低，系统越倾向于给出缩短复检周期、安排停机检查或创建检修工单的建议。",
        "在RUL估计中，系统将健康评分作为基础寿命因子，将风险分数作为惩罚项，将故障严重度作为额外惩罚项。例如严重故障会显著降低RUL估计值，普通告警只会适当缩短复检周期。这样设计的优点是计算逻辑清晰、结果可解释，适合毕业设计阶段展示智能健康管理体系的完整闭环。",
    ])
    add_table_before(anchor, "表3.1 健康评估指标含义", ["指标", "含义", "异常说明"], [
        ["振动RMS", "反映整体振动能量", "升高可能表示磨损、松动或不对中"],
        ["峭度", "反映冲击尖峰程度", "升高可能表示轴承点蚀或齿面剥落"],
        ["峰值因子", "峰值与有效值之比", "升高表示局部冲击增强"],
        ["包络峰值", "捕捉调制冲击特征", "适合识别轴承和齿轮早期故障"],
        ["油温", "反映热状态和冷却能力", "升高可能与冷却、润滑或负荷有关"],
        ["油液NAS", "反映颗粒污染程度", "升高说明油液污染或磨粒增多"],
    ])
    add_subheading_before(anchor, "3.7 智能问答与知识库技术")
    add_long_paragraphs(anchor, [
        "智能问答模块的目的不是替代诊断模型，而是提升系统的人机交互能力。运维人员面对油温过高、振动异常、油液污染等问题时，往往需要查阅手册、经验记录和故障案例。系统将这些知识整理为本地知识库，并结合当前机组状态生成回答，使用户能够通过自然语言方式获得排查建议。",
        "知识库条目包括故障现象、可能原因、检测方法、处理建议和复检周期等内容。系统收到问题后，先根据关键词和意图匹配知识条目，再结合实时状态数据补充当前工况。例如用户询问齿轮箱油温过高时，系统会同时考虑当前油温、振动RMS、油液等级和功率负荷，从而给出更接近现场的排查顺序。",
        "该模块也预留了兼容OpenAI格式的大模型接口。如果配置API Key，系统可以把本地知识和实时状态作为上下文，生成更完整的解释；如果接口不可用，则回退到本地专家规则。这样的设计兼顾了演示稳定性和技术扩展性，避免系统过度依赖外部服务。",
    ])


def expand_chapter_4(doc):
    anchor = find_para(doc, "4.3 后端接口设计")
    add_subheading_before(anchor, "4.2.1 数据库表关系说明")
    add_long_paragraphs(anchor, [
        "数据库设计围绕用户、故障记录、参数配置和闭环处理展开。User表用于保存用户账号、密码哈希和角色信息，支撑登录与权限控制；FaultRecord表保存诊断产生的故障类型、严重程度、置信度、处理状态和建议措施；FaultClosure表记录故障闭环过程，包括处理人、处理说明、复测结果和关闭时间；配置表则保存油温、振动、油液等阈值。",
        "在业务流程上，诊断模块生成结果后会写入故障记录表，前端故障数据页面从该表读取记录并展示；当运维人员完成处理后，系统将处理信息写入闭环表，并更新故障记录状态。健康报告模块读取实时状态、历史故障和闭环记录，生成面向运维人员的综合报告。这样的数据库关系能够支持从报警发现到处理反馈的完整闭环。",
    ])
    add_table_before(anchor, "表4.2 数据库表关系说明", ["数据表", "主要字段", "业务作用"], [
        ["User", "id、username、password_hash、role", "保存用户和角色权限"],
        ["FaultRecord", "unit_id、fault_type、severity、status", "保存诊断结果和处理状态"],
        ["FaultClosure", "record_id、handler、action、result", "保存故障闭环处理信息"],
        ["SystemConfig", "threshold_name、value、updated_at", "保存阈值和系统参数"],
        ["ReportHistory", "report_id、unit_id、summary、created_at", "保存健康报告摘要"],
    ])
    anchor2 = find_para(doc, "4.4 诊断流程设计")
    add_subheading_before(anchor2, "4.3.1 接口设计原则")
    add_long_paragraphs(anchor2, [
        "后端接口按业务域划分为认证接口、数据接口、风场接口、报告接口、问答接口、设置接口和运维接口。认证接口负责登录、注册和当前用户信息；数据接口负责实时状态、振动数据、故障诊断、故障记录和故障代码库；风场接口负责风场总览和单机详情；报告接口负责生成健康报告；问答接口负责运维知识问答；设置接口负责阈值和用户管理；运维接口负责工单和报警策略。",
        "接口返回数据尽量采用统一的JSON结构，前端页面只关注字段含义，不关心后端内部计算过程。这样一方面可以降低前端复杂度，另一方面也便于后续把规则诊断模块替换为机器学习服务或独立微服务。接口设计中还保留了SSE数据流，用于实时推送状态变化，增强监测页面的实时感。",
    ])
    add_table_before(anchor2, "表4.3 后端接口分组说明", ["接口分组", "代表路径", "功能说明"], [
        ["认证接口", "/api/auth/login", "用户登录和身份验证"],
        ["数据接口", "/api/data/status", "实时状态、振动和诊断数据"],
        ["风场接口", "/api/windfarm/overview", "风场总览和单机详情"],
        ["报告接口", "/api/report/generate", "生成健康报告"],
        ["问答接口", "/api/qa/ask", "运维知识问答"],
        ["运维接口", "/api/ops/workorders", "工单模拟和闭环处理"],
    ])


def expand_chapter_5(doc):
    anchor = find_para(doc, "5.2 故障诊断算法模块实现")
    add_picture_before(anchor, SHOT_DIR / "01_windfarm_overview.png", "图5.1 风场健康总览页面", "系统运行后的风场健康总览页面如图5.1所示。该页面以矩阵形式展示多台机组的运行状态，并在顶部汇总时间可利用率、能量可利用率、平均油温、总有功功率、平均健康评分和告警机组数量。")
    add_long_paragraphs(anchor, [
        "风场总览页面是运维人员进入系统后最先接触的页面，其设计目标是快速定位异常机组。页面中不同颜色代表不同运行状态，绿色表示正常运行，橙色和红色表示告警或故障，紫色表示维护停机。每个机组卡片展示功率、风速、油温、健康评分和运行状态，便于用户在一个页面中完成多机组横向比较。",
        "系统还提供指标列表和健康矩阵两种浏览方式。健康矩阵适合快速扫描，指标列表适合查看详细数值。运维人员可以先通过矩阵发现异常机组，再进入单机详情页面查看趋势、故障代码和检修建议。该设计符合风电场运维中“先总览、再下钻、后处理”的工作习惯。",
    ])
    anchor = find_para(doc, "5.3 风场总览与单机详情实现")
    add_picture_before(anchor, SHOT_DIR / "02_fault_diagnosis.png", "图5.2 故障诊断页面", "故障诊断页面如图5.2所示。用户可以选择机组和诊断场景，系统执行诊断后输出故障类型、风险等级、健康评分、RUL、诊断依据和建议措施。")
    add_long_paragraphs(anchor, [
        "故障诊断页面是系统的核心业务页面。用户可以选择不同机组，也可以选择正常、齿轮磨损、点蚀剥落、轴承过温、油液污染、冷却故障等场景进行演示。点击开始诊断后，前端调用后端诊断接口，后端完成振动特征提取、风险评分、故障分类和结果解释，然后将结构化结果返回页面展示。",
        "诊断结果并不只显示故障名称，而是同时给出风险等级、部件位置、置信度、健康评分和剩余寿命。下方的诊断依据和建议措施用于解释为什么系统给出该结论，以及下一步应该如何排查。这样的展示方式能够增强诊断结果可信度，避免用户只看到一个缺少依据的告警标签。",
    ])
    anchor = find_para(doc, "5.4 故障代码库与数据管理实现")
    add_picture_before(anchor, SHOT_DIR / "03_fault_records.png", "图5.3 故障数据管理页面", "故障数据管理页面如图5.3所示。该页面用于查看诊断记录、筛选故障状态、导出CSV并进入故障闭环处理。")
    add_long_paragraphs(anchor, [
        "故障数据管理页面承担历史记录追踪功能。系统每次诊断产生的记录会写入数据库，页面按机组、严重程度、处理状态和关键词进行筛选。对于严重或待处理故障，运维人员可以进入详情查看诊断依据、建议措施和处理进度。",
        "该页面还提供CSV导出功能，便于将故障记录用于论文测试、运维日报或后续数据分析。对于毕业设计而言，故障记录页面能够证明系统不是只做一次性诊断，而是具备数据积累和闭环管理能力。",
    ])
    anchor = find_para(doc, "5.5 AI运维问答模块实现")
    add_picture_before(anchor, SHOT_DIR / "04_health_report.png", "图5.4 健康报告页面", "健康报告页面如图5.4所示。系统根据当前机组状态和历史故障统计生成健康评分、关键指标、诊断结论和维护建议。")
    add_long_paragraphs(anchor, [
        "健康报告页面将实时监测和故障诊断结果转化为面向运维管理的文档化信息。报告中包括机组编号、统计范围、生成时间、健康评分、油温、振动RMS、油液NAS、功率、数据质量、RUL和复检周期等关键指标。相比单纯的仪表盘，健康报告更适合用于阶段性检查和汇报。",
        "报告结论部分会根据健康评分和风险因子生成文字说明，指出当前机组的主要风险和维护建议。如果系统发现油温、振动或油液指标异常，报告会提示缩短复检周期或安排现场检查。该功能体现了智能健康管理体系从监测数据到管理决策的转化能力。",
    ])
    anchor = find_para(doc, "5.6 健康报告与工单闭环实现")
    add_picture_before(anchor, SHOT_DIR / "05_ai_qa.png", "图5.5 AI运维问答页面", "AI运维问答页面如图5.5所示。用户可以围绕油温、振动、轴承、齿轮、润滑和检修策略提出问题，系统结合知识库和当前状态生成回答。")
    add_long_paragraphs(anchor, [
        "AI运维问答页面面向运维人员的日常咨询场景。传统系统通常只能提供固定报警，而问答模块可以让用户以自然语言提出问题，例如“齿轮箱油温过高如何排查”“振动RMS升高是否需要停机”“油液NAS等级偏高如何处理”等。系统根据问题意图匹配知识库条目，并结合当前状态给出排查顺序。",
        "在实现上，问答模块先使用本地知识库保证回答稳定性，再根据配置决定是否调用外部大模型接口。即使没有网络或API Key，系统仍然能够返回专家规则答案。这样既满足毕业设计演示稳定性，也体现了智能运维系统可扩展到大模型辅助分析的方向。",
    ])
    anchor = find_para(doc, "第6章 系统测试与结果分析")
    add_picture_before(anchor, SHOT_DIR / "06_settings.png", "图5.6 系统参数设置页面", "系统参数设置页面如图5.6所示。该页面用于配置油温、振动、油液阈值和系统访问参数，阈值变化会影响实时告警和健康报告。")
    add_long_paragraphs(anchor, [
        "系统设置页面用于体现健康管理体系的可配置性。不同风场、不同季节和不同运行策略下，油温、振动和油液阈值可能需要调整。如果阈值完全写死在代码中，系统很难适应现场环境。因此系统提供参数配置页面，使管理员能够根据运维策略调整预警阈值和严重阈值。",
        "参数配置不仅影响页面显示，也会影响诊断结果、告警等级和报告建议。这样可以保证系统的业务逻辑与管理策略一致。用户管理和角色权限则用于控制不同人员的操作范围，避免普通用户误改关键阈值。",
    ])


def expand_chapter_6(doc):
    anchor = find_para(doc, "第7章 总结与展望")
    add_subheading_before(anchor, "6.5 测试用例设计")
    add_long_paragraphs(anchor, [
        "系统测试围绕功能正确性、接口稳定性、页面展示效果和诊断结果合理性展开。由于本系统是毕业设计原型，测试重点不是高并发压力，而是验证主要业务流程是否完整：用户能否登录，风场总览是否正常展示，诊断接口是否返回结构化结果，故障记录是否保存，报告是否生成，问答模块是否能够给出合理建议，阈值配置是否能够影响后续判断。",
        "测试过程中采用黑盒测试与接口检查相结合的方法。黑盒测试从用户角度操作页面，观察页面元素、图表和表格是否正常；接口检查则直接访问后端接口，确认返回字段是否完整。对于故障诊断场景，系统分别测试正常运行、油温过高、轴承冲击、油液污染、齿轮啮合异常和冷却系统故障等典型状态。",
    ])
    add_table_before(anchor, "表6.2 功能测试用例", ["测试编号", "测试内容", "预期结果", "测试结论"], [
        ["TC-01", "用户登录", "输入正确账号后进入系统主页", "通过"],
        ["TC-02", "风场总览", "显示机组矩阵和关键统计指标", "通过"],
        ["TC-03", "故障诊断", "返回故障类型、健康评分和建议措施", "通过"],
        ["TC-04", "故障记录", "诊断结果写入记录并可筛选查看", "通过"],
        ["TC-05", "健康报告", "生成机组健康指标和文字结论", "通过"],
        ["TC-06", "AI问答", "根据问题返回排查建议", "通过"],
        ["TC-07", "参数设置", "阈值可修改并影响告警策略", "通过"],
    ])
    add_subheading_before(anchor, "6.6 典型诊断场景结果")
    add_long_paragraphs(anchor, [
        "为了验证诊断逻辑的合理性，本文选取正常运行、齿轮箱油温过高、轴承冲击、润滑油污染和齿轮啮合异常五类场景进行对比测试。每个场景通过不同的油温、振动RMS、油液NAS和振动冲击特征组合触发。系统根据风险因子计算结果给出不同的健康评分和维护建议。",
        "正常运行场景中，油温、振动和油液指标均处于较低水平，系统输出健康评分较高，建议保持日常巡检。油温过高场景中，系统重点提示冷却风扇、油冷器、油位和滤芯压差检查。轴承冲击场景中，峭度和包络峰值升高，系统提示复核高速轴承包络谱。油液污染场景中，NAS等级升高，系统建议取样复检和检查滤芯。齿轮啮合异常场景中，振动能量和峰值因子同步升高，系统提示检查齿面接触斑和啮合频率。",
    ])
    add_table_before(anchor, "表6.3 典型诊断场景测试结果", ["场景", "主要异常指标", "系统诊断", "维护建议"], [
        ["正常运行", "指标均正常", "正常运行", "保持日常巡检"],
        ["油温过高", "oil_temp升高", "齿轮箱油温过高", "检查冷却系统和润滑状态"],
        ["轴承冲击", "kurtosis、envelope_peak升高", "轴承点蚀/剥落", "复核包络谱并安排检查"],
        ["油液污染", "NAS等级升高", "润滑油污染/劣化", "取样复检并检查滤芯"],
        ["啮合异常", "RMS和峰值因子升高", "齿轮啮合异常", "检查齿面和载荷波动"],
    ])
    add_subheading_before(anchor, "6.7 系统运行效果分析")
    add_long_paragraphs(anchor, [
        "从运行效果看，系统能够将风场级总览、单机级详情、诊断级解释和报告级输出连接起来。风场总览用于发现异常对象，故障诊断用于判断异常类型，故障记录用于保存历史，健康报告用于形成汇总结论，AI问答用于补充排查知识，参数设置用于调整策略。各模块之间不是孤立页面，而是围绕齿轮箱健康管理形成了较完整的业务链条。",
        "从交互效果看，系统页面采用侧边导航和分区显示方式，运维人员可以在健康监测区、智能诊断区和系统配置区之间切换。关键指标采用卡片和图表展示，故障记录采用表格和筛选器展示，报告页面采用文档化结构展示。整体交互符合运维系统需要快速浏览、快速定位和快速处理的特点。",
        "从工程实现看，系统后端采用蓝图拆分接口，前端采用模块化函数组织页面逻辑，数据库保存用户和故障记录。虽然系统仍属于毕业设计原型，但已经具备从数据生成、特征提取、诊断判断、结果解释、记录保存到报告生成的基本闭环。后续如果接入真实风场数据，系统可以在现有接口基础上继续扩展。",
    ])


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    if any("2.5 需求约束与边界分析" in p.text for p in doc.paragraphs):
        print("already expanded")
    else:
        expand_chapter_2(doc)
        expand_chapter_3(doc)
        expand_chapter_4(doc)
        expand_chapter_5(doc)
        expand_chapter_6(doc)
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
