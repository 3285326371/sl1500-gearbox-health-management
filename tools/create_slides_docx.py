import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "docs" / "答辩PPT" / "毕业答辩PPT大纲与演讲稿_学术绿.docx"

def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

def run_font(run, name="宋体", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(text)
    run_font(r, "黑体", 18, bold=True)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(text)
    run_font(r, "宋体", 12, italic=True, color="555555")
    return p

def add_slide_header(doc, title_zh, title_en):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    
    r_zh = p.add_run(title_zh + " ")
    run_font(r_zh, "黑体", 14, bold=True, color="1E3A2F")
    
    r_en = p.add_run(title_en)
    run_font(r_en, "Times New Roman", 12, bold=True, color="55665C")
    return p

def add_visual_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    
    r_prefix = p.add_run("💡 版面设计建议：")
    run_font(r_prefix, "宋体", 10, bold=True, color="2D4538")
    
    r_text = p.add_run(text)
    run_font(r_text, "宋体", 10, italic=True, color="55665C")
    return p

def add_bullet(doc, text, bold_lead_len=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    
    # Times New Roman vs 宋体 formatting
    if bold_lead_len > 0:
        lead = text[:bold_lead_len]
        body = text[bold_lead_len:]
        
        r_lead = p.add_run(lead)
        run_font(r_lead, "黑体", 10.5, bold=True, color="1E3A2F")
        
        # Split body to handle Times New Roman for English/Numbers
        r_body = p.add_run(body)
        run_font(r_body, "宋体", 10.5)
    else:
        r_text = p.add_run(text)
        run_font(r_text, "宋体", 10.5)
    return p

def add_speech(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.4
    r_prefix = p.add_run("口述稿备忘：\n")
    run_font(r_prefix, "黑体", 10, bold=True, color="2D4538")
    
    r_text = p.add_run(text)
    run_font(r_text, "楷体", 10.5, color="333333")
    return p

def add_section_title(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    run_font(r, "黑体", 16, bold=True, color="1E3A2F")
    return p

def add_category_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    run_font(r, "黑体", 12, bold=True, color="2D4538")
    return p

def add_qa_block(doc, q_num, question, intention, answer):
    # Add Question
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(10)
    p_q.paragraph_format.space_after = Pt(4)
    p_q.paragraph_format.keep_with_next = True
    r_prefix = p_q.add_run(f"问 {q_num}：")
    run_font(r_prefix, "黑体", 10.5, bold=True, color="1E3A2F")
    r_q = p_q.add_run(question)
    run_font(r_q, "黑体", 10.5, bold=True)
    
    # Add Intention
    p_i = doc.add_paragraph()
    p_i.paragraph_format.left_indent = Cm(0.5)
    p_i.paragraph_format.space_after = Pt(4)
    r_i_prefix = p_i.add_run("🎯 考察意图：")
    run_font(r_i_prefix, "宋体", 9.5, bold=True, color="55665C")
    r_i = p_i.add_run(intention)
    run_font(r_i, "宋体", 9.5, italic=True, color="55665C")
    
    # Add Answer
    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(1.0)
    p_a.paragraph_format.right_indent = Cm(1.0)
    p_a.paragraph_format.space_after = Pt(12)
    p_a.paragraph_format.line_spacing = 1.3
    r_a_prefix = p_a.add_run("💡 答辩参考回答：\n")
    run_font(r_a_prefix, "黑体", 9.5, bold=True, color="2D4538")
    r_a = p_a.add_run(answer)
    run_font(r_a, "楷体", 10.5, color="333333")

def main():
    doc = Document()
    set_doc_margins(doc)
    
    # Title Page
    add_title(doc, "《华锐SL1500型双馈式风机齿轮箱智能健康管理系统》")
    add_subtitle(doc, "毕业答辩 PPT 大纲、口述演讲稿与核心问题备忘\n（学术雅致·Times New Roman 规范版）")
    
    # Slide 1
    add_slide_header(doc, "Slide 1：封面页", "Title Slide")
    add_visual_note(doc, "米白色底色，四周使用细双线深墨绿色边框。标题大号中宋体，英文缩写（SL1500）使用 Times New Roman。")
    add_bullet(doc, "汇报题目：华锐 SL1500 型双馈式风机齿轮箱智能健康管理系统设计 with 实现", 5)
    add_bullet(doc, "汇报人：[你的名字]  |  指导教师：[导师姓名]", 4)
    add_bullet(doc, "专业班级：[你的专业班级]  |  答辩时间：2026年6月", 5)
    add_speech(doc, "“各位答辩委员会的老师，大家下午好！我是来自[专业班级]的[你的名字]。我的毕业设计论文题目是《华锐SL1500型双馈式风机齿轮箱智能健康管理系统设计与实现》，本课题是在[导师姓名]老师的悉心指导下完成的。在接下来的十分钟里，我将从课题的研究背景、系统的架构设计、核心技术突破以及系统的具体实现四个部分，向各位老师汇报我的研究成果，希望得到各位老师的批评与指正。”")
    
    # Slide 2
    add_slide_header(doc, "Slide 2：选题背景与研究意义", "Background & Significance")
    add_visual_note(doc, "左右分栏布局。左侧展示风力发电传动链中齿轮箱的受力特征与故障停机经济损失柱状图；右侧使用学术符号引导展示痛点。")
    add_bullet(doc, "服役规模巨大：华锐 SL1500 机型在我国分布广泛，是双馈式风力发电机的典型代表。", 7)
    add_bullet(doc, "工作负荷剧烈：长期承受复杂的交变载荷与极端外部温差，核心齿轮及轴承极易发生疲劳损伤。", 7)
    add_bullet(doc, "停机损失高昂：非计划停机维修周期长、起吊费用大，经济损失可达数十万元。", 7)
    add_bullet(doc, "核心研究目的：实现从“故障后被动维修”向“状态监测主动预测性维护”的转变，降低运维成本。", 7)
    add_speech(doc, "“首先介绍本课题的选题背景。华锐 SL1500 系列机组在我国服役基数非常庞大。而在风机运行中，齿轮箱作为传动链的核心枢纽，长期处于复杂的交变载荷与极端外部温差下，导致其成为整机故障率最高、维修成本最贵的部件之一。传统的故障后维修滞后性强，定期排查成本高昂。因此，设计开发一套在线智能健康管理系统，在微弱故障萌芽期进行精准预测，对降低停机损失、实现预防性维护，具有极高的实际工程应用价值。”")
    
    # Slide 3
    add_slide_header(doc, "Slide 3：系统总体设计与技术架构", "System Architecture")
    add_visual_note(doc, "层次清晰的技术架构图。采用墨绿色与灰色线条，清晰标注三层架构设计。英文技术词汇使用 Times New Roman。")
    add_bullet(doc, "表现层 (Web UI)：采用轻量级的原生 HTML5/CSS3 与 ES6 脚本，使用 Three.js 完成 3D 模型的高效渲染。", 13)
    add_bullet(doc, "业务层 (Flask Backend)：依托 Python Flask 构建，实现 HFRT 滤波算法运算、RUL 寿命预测以及 RAG 对话检索。", 17)
    add_bullet(doc, "数据中枢 (SQLite & VectorDB)：SQLite 存储机组遥测数据，向量数据库存储结构化的专业风电运维知识库。", 17)
    add_speech(doc, "“本系统基于实用与轻量化的原则，采用前后端分离的 B/S 架构。底层数据源通过仿真技术，提供持续的振动、温度、功率等多元遥测信号；后端服务采用 Flask 框架构建，负责精密信号处理算法（如希尔伯特变换）的计算、剩余寿命预测以及向量库 of 检索与大模型问答；前端则基于原生 Web 技术，使用 Three.js 在浏览器端渲染 3D 齿轮箱模型。整个系统遵循轻量化与安全性，保证了数据流的高效低延迟。”")
    
    # Slide 4
    add_slide_header(doc, "Slide 4：突破一：精密诊断——基于 HFRT 的早期故障特征提取", "Precision Diagnostics")
    add_visual_note(doc, "展示对比波形。上方为杂乱的时域原始波形 x(t)，下方为清晰的包络频谱图 E(f)，缺陷特征频率处有高亮谱线峰值。物理量、公式均渲染为 Times New Roman 格式。")
    add_bullet(doc, "故障特征提取难点：齿轮微裂纹、轴承剥落冲击能量微弱，完全淹没在齿轮箱强烈的宽带背景噪声中。", 9)
    add_bullet(doc, "希尔伯特包络解调 (HFRT)：构建解析信号 z(t) = x(t) + i*H[x(t)]，提取瞬时幅值包络，滤除高频无用载波。", 19)
    add_bullet(doc, "诊断效果与呈现：前端支持“原始波形 与 包络谱”一键切换，直观呈现明显的故障特征缺陷频率。", 8)
    add_speech(doc, "“下面汇报系统的第一大技术突破：精密诊断算法。早期故障产生的冲击能量非常微弱，直接进行普通频谱分析无法提取特征。本系统在后端编写了基于希尔伯特变换的包络解调算法，通过构建解析信号提取信号的瞬时幅值包络，从而有效滤除了背景宽带载波噪声。从解调后的包络谱可以看出，故障对应的特征频率处有极明显的谱线峰值突起，操作人员在前端通过一键切换，即可精准锁定故障源头。”")
    
    # Slide 5
    add_slide_header(doc, "Slide 5：突破二：数字孪生——3D 部件级物理与交互联动", "Digital Twin")
    add_visual_note(doc, "展示淡米色背景的 3D 齿轮箱数字孪生截图。高亮显示正在被检测的部件，并以箭头标注“物理变速旋转”和“健康画像浮窗”。")
    add_bullet(doc, "转速功率深度联动 (Power-Sync)：3D 行星轮、平行齿轮和散热风扇的旋转速度与实时功率 P 动态绑定变速，实现真实的物理镜像。", 18)
    add_bullet(doc, "三维交互与悬停高亮：使用 Raycaster 射线进行鼠标碰撞检测，悬停时组件发出亮蓝色（Emissive Glow）发光材质反馈。", 10)
    add_bullet(doc, "组件级健康画像：点击特定组件激活悬浮面板，实时呈现其温度变化、磨损率以及剩余有用寿命 (RUL) 预测值。", 8)
    add_speech(doc, "“系统的第二大突破在于 3D 数字孪生的物理联动。我们利用 Three.js 构建了齿轮箱各级传动副的三维几何模型。为了摆脱静态贴图的局限，我们通过运动方程让 3D 齿轮和风扇的自转速度与风机实时的功率遥测值挂钩，实现‘虚实物理镜像’。同时，我们编写了射线投射检测，当鼠标悬停于特定部件时其会高亮发光，点击部件即可在右上角调出该部件对应的磨损率、温度及寿命预测画像。”")
    
    # Slide 6
    add_slide_header(doc, "Slide 6：突破三：智慧决策——工况感知的 RAG 智能专家系统", "Smart LLM & RAG")
    add_visual_note(doc, "工况感知数据流图。展示当前遥测参数（温升、报警码）自动拼接并检索知识库手册，输入大模型后得到定制化策略。")
    add_bullet(doc, "打破知识孤岛：华锐 SL1500 技术规程、故障记录的向量化存储与检索。", 7)
    add_bullet(doc, "实时工况感知：大模型问答引擎能自动感知风机当前的遥测工况（如轴承温度 85℃ 超限）。", 7)
    add_bullet(doc, "定制决策输出：结合具体故障状态，大模型生成个性化的“限负荷运行建议”与“冷却水泵故障排查流程”。", 7)
    add_speech(doc, "“第三项突破是智慧决策系统的实现。传统的问答只提供死板的条文说明，不具备时效性。我们研发了上下文感知的检索增强生成技术。当机组出现高温或异常振动时，后台问答引擎会自动抓取当前的遥测异常参数，作为背景信息融入大模型的提示词中。此时大模型给出的将不再是泛泛的说明书，而是结合了当前实况、极具针对性的排查步骤和运维决策，大幅降低了现场对人工经验的依赖。”")
    
    # Slide 7
    add_slide_header(doc, "Slide 7：系统界面展示与核心功能协同", "System UI Demonstration")
    add_visual_note(doc, "整洁优雅的磨砂玻璃风格大屏主控界面截图。标注出实时指标卡片、中央 3D 数字孪生、包络诊断图表 and 右下角 AI 对话区。")
    add_bullet(doc, "主控数据卡片：采用 SSE 长连接推送技术，关键运行参数无刷新、毫秒级响应滚动展示。", 7)
    add_bullet(doc, "图表与趋势对标：集成 ECharts，提供劣化趋势线、健康度评分和同型集群标杆对标。", 8)
    add_bullet(doc, "闭环运维管理：包含登录控制、系统通知、参数设置、故障数据列表和健康报告导出。", 7)
    add_speech(doc, "“这是我们实际的系统运行大屏。整体界面设计遵循学术规范，排版大方。屏幕核心为 3D 数字孪生区，左侧展示由 SSE 单向长连接实时推送的风速、有功功率等关键工况，右下角融合了 RAG AI 专家对话终端，为现场值班人员提供诊断解释。系统还配备了权限认证、系统通知及参数设置等模块，全方位保障了系统闭环运行的安全可控。”")
    
    # Slide 8
    add_slide_header(doc, "Slide 8：总结与核心创新点", "Conclusion & Innovations")
    add_visual_note(doc, "严谨的学术三栏创新矩阵表格。将本系统的方案与传统只读监控进行对比，突出算法和 RAG 遥测感知的科学性。")
    add_bullet(doc, "精密物理与数字孪生深度融合：引入 HFRT 共振解调算法，实现变噪干扰下早期微弱故障提取。", 14)
    add_bullet(doc, "模型变速与实时遥测公式绑定：3D 孪生具备动力学方程，摆脱了单纯贴图，提升物理镜像价值。", 14)
    add_bullet(doc, "遥测数据与大模型知识感应：设计 Context-Aware 问答框架，解决工业数据与大模型知识割裂的难题。", 13)
    add_speech(doc, "“总结本系统的核心创新点：第一，实现了物理诊断与数字孪生的深度融合，利用包络解调有效提取了强噪下的微弱特征；第二，实现了三维渲染与物理方程的联动，使 3D 模型变速具备物理意义；第三，设计了工况感知的检索增强生成框架，解决了大模型与工业数据脱节的问题。本系统已经过完整测试，算法鲁棒，能够为风机齿轮箱健康管理提供闭环的应用成果。”")
    
    # Slide 9
    add_slide_header(doc, "Slide 9：未来展望与致谢", "Future Outlook & Q&A")
    add_visual_note(doc, "答辩致谢页。配以未来研究规划列表。英文使用 Times New Roman，谢谢大家使用中宋体。")
    add_bullet(doc, "寿命评估预测精细化：深入优化非平稳工况下小样本神经网络对剩余寿命 (RUL) 的拟合精度。", 10)
    add_bullet(doc, "边缘计算物理部署：探索基于边缘智能计算网关的算法硬件化移植，提升海上风场实际应用能力。", 10)
    add_bullet(doc, "谢辞：感谢指导老师的耐心教诲，感谢答辩委员会各位老师的批评指正！", 3)
    add_speech(doc, "“在未来的工作中，我们将进一步研究小样本条件下神经网络在多变工况下的剩余寿命拟合精度，并推进边缘网关的硬件集成部署。在课题完成之际，我向给予我悉心指导的导师表示最诚挚的感谢，同时也非常感谢在座的各位评委老师百忙之中聆听我的汇报。谢谢大家！请各位老师批评指正。”")
    
    # Q&A Section
    add_section_title(doc, "毕业答辩常见核心问题与标准回答备忘")
    
    add_category_title(doc, "一、 精密物理诊断与特征解调类")
    
    add_qa_block(doc, "1", 
                 "为什么要采用“包络解调（HFRT）”而不是普通的傅里叶变换（FFT）来提取故障？",
                 "考察你对机械故障诊断专业知识的深度理解，判断是否属于套用算法的“空壳系统”。",
                 "在齿轮箱运行早期，齿轮剥落或轴承点蚀引发的冲击信号极其微弱。由于齿轮箱内部背景噪声巨大，直接做傅里叶变换（FFT），故障谱线会被随机噪声完全淹没。而高频共振解调（HFRT）通过带通滤波和解析信号包络处理，剔除了背景无用高频载波，只保留冲击调制特征，从而直接在低频段显现故障特征频率。这是从‘频域淹没’到‘频域显现’的质的飞跃。")
                 
    add_qa_block(doc, "2",
                 "系统在后端具体是如何计算希尔伯特变换和解析信号包络的？",
                 "验证你对信号处理核心代码块的熟悉程度，防止学术不端或纯套库嫌疑。",
                 "我们在后端 ml_models.py 文件中，借助 scipy 库的 hilbert 函数计算。核心步骤是：将实测时域振动信号 x(t) 输入，计算其 Hilbert 变换正交分量 h(t)；构造复数形式的解析信号 z(t) = x(t) + i*h(t)；随后计算解析信号的幅值解析模 A(t) = |z(t)| = sqrt(x(t)^2 + h(t)^2)。这个 A(t) 就是提取的振动信号幅值包络，之后在前端做一键切换分析。")

    add_qa_block(doc, "3",
                 "系统中的振动数据源是如何得到的？如果是模拟生成的，数学公式是怎样的？",
                 "验证试验数据来源与数学建模能力，确保研究具备严谨的数据支持。",
                 "本系统的振动数据依托‘物理机制仿真’合成。为了模拟典型的齿面剥落等微弱故障特征，我们构建了多成份合成信号。包含：代表齿轮副啮合频的简谐分量 A_mesh*sin(2*pi*f_mesh*t)，模拟故障冲击特征的指数衰减简谐冲击分量 A_fault*e^(-alpha*t)*sin(2*pi*f_res*t)，以及一定强度的加性高斯白噪声。通过控制信噪比，我们验证了解调算法在复杂噪声中的有效性。")

    add_category_title(doc, "二、 数字孪生与 3D 可视化交互类")

    add_qa_block(doc, "4",
                 "你的 3D 数字孪生体如何实现“转速与风机实时输出功率绑定”的动态变速旋转？",
                 "考察前端与后端数据通道的实时性，判断是否是预渲染录像的“假孪生”。",
                 "我们打通了前后端的数据连接。在前端 main.js 的渲染循环 requestAnimationFrame 中，我们接收后端推送的发电机实时有功功率 Power 数据。我们建立转速物理联动公式：w = K * Power + w_min（其中 K 是传动系数，w_min 是零负载基准角速度）。在每一帧渲染计算中，系统动态改变 3D 行星架、平行级齿轮及冷却风扇的单帧旋转增量 delta_theta。当发电机有功功率增加时，齿轮组即物理加速旋转，实现真孪生联动。")

    add_qa_block(doc, "5",
                 "鼠标悬停 3D 组件高亮以及点击弹出组件健康画像，底层的核心原理是什么？",
                 "考察对 Three.js 核心三维空间拾取技术（Raycaster）的理解与应用能力。",
                 "底层利用了 Three.js 提供的射线投射器（Raycaster）原理。当用户操作鼠标滑过 3D 视口时，我们将屏幕的二维坐标归一化并传入 Raycaster。Raycaster 会从摄像机原点朝该二维坐标点发射一条虚拟的三维空间射线。系统计算该射线与齿轮箱中所有 3D 零部件网格（Mesh）的碰撞交点。一旦捕获到交点，证明部件被触碰，随即更新其材质自发光属性 emissive 实现高亮；点击时则提取网格属性中的部件名，动态请求后端接口刷新健康指标。")

    add_qa_block(doc, "6",
                 "为什么不直接使用商业大屏可视化工具，而是自行用原生 Three.js 编码开发？",
                 "考察你的工作量、技术含金量以及对系统底层自主可控性的把控。",
                 "商业可视化大屏虽然开发迅速，但多为封装好的“死图表”，无法深度嵌入复杂的物理计算公式，更无法支持如 3D 射线检测、基于物理功率变速的转速联动等深层实时数学运算。使用原生 Three.js 编码开发不仅展现了本科/硕士阶段充足的工程工作量，更为系统赋予了极高的数据自主权，能够完全打通‘传感器数据 -> 物理动力学计算 -> 3D动态渲染’的链路。")

    add_category_title(doc, "三、 AI 大模型与 RAG 决策应用类")

    add_qa_block(doc, "7",
                 "什么是检索增强生成（RAG）？它在你的健康管理系统中解决了什么问题？",
                 "考察对大模型前沿应用架构的理解，尤其是解决大模型工业领域常识盲区的问题。",
                 "RAG 即检索增强生成。普通的通用 LLM 缺乏风电领域的垂直机型说明书和现场维护记录，提问时容易出现常识性胡说（幻觉）。我们收集了华锐 SL1500 齿轮箱维护手册和案例库，结构化后向量化存入数据库。当用户提问时，系统先在本地查找最匹配的参考文档段落，与用户问题共同发给大模型。RAG 解决了大模型没有行业私有知识的弊端，使 AI 回答准确、专业、有据可查。")

    add_qa_block(doc, "8",
                 "系统是如何实现大模型对“实时工况（Context-Aware）”的感知的？",
                 "考察系统的智能化融合水平以及工况注入大模型提示词的具体实现方法。",
                 "在 rag_service.py 中，大模型运行前，系统会先请求当前的 SCADA 和数据传感器状态。若检测到特定的遥测指标异常（如油温 85.2℃ 高温警告），问答引擎会自动将该‘遥测工况上下文’拼接成一段格式化提示词段落。这样，AI 在检索手册的同时，能够完全感知当前风机发生了什么，输出的排查方案就是针对当前这台高温警告机组的“针对性急救流程”，而非死板的通用条款。")

    add_qa_block(doc, "9",
                 "如果向量数据库中的设备维护指南更新了，你的系统需要如何同步处理？",
                 "考察对向量化更新机制与知识库维护流程的理解。",
                 "更新知识库非常便捷。我们只需要将更新后的 PDF/Word 版维护守则进行文本切片，使用同样的 Embedding 编码模型将其编码为新的高维数学向量，然后执行数据库覆盖写操作或增量插入操作更新向量数据库。系统在运行时会自动使用新的索引，大模型在下一次查询中即可阅读并采纳最新的维护建议。")

    add_category_title(doc, "四、 工程实践与安全性考量类")

    add_qa_block(doc, "10",
                 "系统的实时数据推送为什么选择服务器发送事件（SSE）而不是 WebSocket？",
                 "考察网络通信协议选型的深度对比与工程实用性考量。",
                 "因为风电场中控大屏属于典型的单向数据流推送。数据由服务器源源不断地以高频状态向前端推送更新，前端并不需要向服务器频繁写入数据。WebSocket 虽然支持双向通信，但协议较重、维护心跳开销大，且较容易因网络网闸拦截导致掉线。而 SSE 基于标准的 HTTP 协议，轻量高效，具备原生断线重连机制，能完美满足单向高频流数据的高稳定推送。")

    add_qa_block(doc, "11",
                 "风场高频采集传感器（数十 kHz）数据量极大，在系统上线部署时如何缓解网络与存储瓶颈？",
                 "考察在真实高频海量工业互联网环境下的系统性能优化与工程落地经验。",
                 "实际工程中我们会使用‘分布式边缘计算 + 快照回传’架构。高频采集的振动加速度信号并不会直接全量传回中心服务器。信号解调、特征提取（如 RMS、峭度）等高计算量过程会在风机底部的边缘计算网关直接算完。平时仅向中心回传 1Hz 的低频特征指标。只有当特征指标超限触发报警时，边缘网关才会抓取一段几秒钟的高频时域波形快照回传，用于详细故障诊断分析。这既保证了诊断精度，又节省了 95% 以上的数据带宽存储。")

    # Save document
    doc.save(str(OUT_FILE))
    print("Docx file generated successfully at:", OUT_FILE)

if __name__ == "__main__":
    main()
