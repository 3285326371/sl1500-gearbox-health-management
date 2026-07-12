from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = sorted(DOCS.glob("*代码分析优化版.docx"), key=lambda p: p.stat().st_mtime, reverse=True)[0]
OUT = DOCS / "毕业设计正文_华锐SL1500齿轮箱智能健康管理系统_开题文献版.docx"

REFERENCES = [
    "[1] 杨俊峰.小样本下基于深度神经网络的行星齿轮箱故障诊断[D].成都:电子科技大学,2024.DOI:10.27005/d.cnki.gdzku.2024.005469.",
    "[2] 周昶清.基于统计模型与稀疏表示的齿轮箱故障检测与诊断方法研究[D].杭州:浙江大学,2024.DOI:10.27461/d.cnki.gzjdx.2024.000307.",
    "[3] 武甲.基于深度学习与降维技术的风电机组齿轮箱状态监测可视化和健康评估[D].呼和浩特:内蒙古工业大学,2024.DOI:10.27225/d.cnki.gnmgu.2024.000206.",
    "[4] 许晴.基于时频分析和机器学习的齿轮箱故障诊断[D].北京:华北电力大学(北京),2024.",
    "[5] 吴岚.基于先进信号分析方法的风电机组齿轮箱故障诊断[D].北京:华北电力大学(北京),2024.",
    "[6] 马健.基于振动信号的风电机组齿轮箱特征提取方法和故障诊断研究[D].上海:上海电机学院,2024.",
    "[7] 吴迪.基于自动特征学习的风机齿轮箱故障诊断研究[D].沈阳:沈阳工业大学,2024.DOI:10.27322/d.cnki.gsgyu.2024.000493.",
    "[8] 孔德强.风电齿轮箱温度预警及关键部件故障诊断方法[D].沈阳:沈阳工程学院,2024.DOI:10.27845/d.cnki.gsygc.2024.000003.",
    "[9] 李常见.风电机组齿轮箱故障识别及安全状态评估研究[D].阜新:辽宁工程技术大学,2024.",
    "[10] 郭建超.风电机组齿轮箱轴承微弱故障诊断方法研究及系统开发[D].上海:上海电机学院,2024.",
    "[11] 王广.风力发电机组齿轮箱故障检测和寿命预测方法的研究[D].兰州:兰州理工大学,2024.",
    "[12] 范亚飞.基于变分模态分解和随机配置网络的齿轮箱故障诊断研究[D].石家庄:石家庄铁道大学,2024.DOI:10.27334/d.cnki.gstdy.2024.000257.",
    "[13] 杨青松.基于改进ShuffleNet的齿轮箱故障诊断研究[D].石家庄:石家庄铁道大学,2024.DOI:10.27334/d.cnki.gstdy.2024.000304.",
    "[14] 王莘然.基于迁移学习的风电齿轮箱故障诊断方法研究[D].大连:大连海事大学,2024.",
    "[15] 王丹.基于深度迁移学习的齿轮箱故障诊断方法研究[D].济南:齐鲁工业大学,2024.DOI:10.27278/d.cnki.gsdqc.2024.000503.",
]


def set_run_font(run, size=10.5):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run.font.size = Pt(size)


def replace_references(doc):
    paras = doc.paragraphs
    start = None
    end = None
    for idx, para in enumerate(paras):
        text = para.text.strip()
        if text == "参考文献":
            start = idx
            continue
        if start is not None and text == "致谢":
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("未找到参考文献或致谢段落，无法替换。")

    # Blank existing reference paragraphs between “参考文献” and “致谢”.
    for idx in range(start + 1, end):
        paras[idx].clear()

    insert_anchor = paras[end]
    for offset, ref in enumerate(REFERENCES):
        idx = start + 1 + offset
        if idx < end:
            p = paras[idx]
        else:
            p = insert_anchor.insert_paragraph_before()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ref)
        set_run_font(r, 10.5)


def main():
    doc = Document(SOURCE)
    replace_references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
