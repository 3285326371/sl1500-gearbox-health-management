from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "风电场数据对接与项目阶段成果说明.docx"
OUT = ROOT / "docs" / "风电场数据对接与项目阶段成果说明_信息安全修订版.docx"

OLD = "接口资料：SCADA/PLC/CMS 数据访问方式、IP 与端口、账号权限、协议说明、点表或字段字典。"
NEW = (
    "数据接口与信息安全资料：包括 SCADA、PLC、CMS 等系统的数据访问方式、通信协议、"
    "点表信息或字段字典等；涉及 IP 地址、端口、账号权限、访问密钥等敏感信息时，"
    "应按风电场网络安全要求进行脱敏处理，或仅在现场受控环境下查看。系统原则上以只读方式接入运行数据，"
    "不直接修改 PLC 控制逻辑，不越权访问生产控制网络，并对数据访问过程进行权限控制和操作记录，"
    "确保数据接入安全、可控、可追溯。"
)


def replace_in_paragraph(paragraph):
    if OLD not in paragraph.text:
        return False
    paragraph.text = paragraph.text.replace(OLD, NEW)
    return True


def main():
    doc = Document(SOURCE)
    changed = 0
    for paragraph in doc.paragraphs:
        changed += int(replace_in_paragraph(paragraph))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed += int(replace_in_paragraph(paragraph))
    doc.save(OUT)
    print(OUT)
    print("changed", changed)


if __name__ == "__main__":
    main()
