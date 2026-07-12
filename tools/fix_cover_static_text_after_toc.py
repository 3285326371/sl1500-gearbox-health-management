from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"C:\Users\隐涅\Desktop\毕业设计正文模板_检测报告\毕业设计正文模板.docx")
TITLE_1 = "华锐 SL1500 型双馈式风机齿轮箱"
TITLE_2 = "智能健康管理体系"


def set_font(run, east_asia="黑体", latin="Arial", size=None, bold=None, underline=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.underline = underline


doc = Document(DOCX)
for paragraph in doc.paragraphs[:40]:
    text = paragraph.text.strip()
    if "华锐 SL1500" in text and "齿轮箱" in text:
        paragraph.clear()
        run = paragraph.add_run(TITLE_1)
        set_font(run, "黑体", "Arial", 20, False, False)
        paragraph.alignment = 1
    elif text == "智能健康管理体系":
        paragraph.clear()
        run = paragraph.add_run(TITLE_2)
        set_font(run, "黑体", "Arial", 20, False, False)
        paragraph.alignment = 1
    elif "2026" in text and ("月" in text or "日" in text):
        paragraph.clear()
        run = paragraph.add_run("2026  年  05  月")
        set_font(run, "宋体", "Times New Roman", 14, False, False)
        paragraph.alignment = 1

doc.save(DOCX)
print("cover title/date fixed")
