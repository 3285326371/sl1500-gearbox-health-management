from __future__ import annotations

import copy
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DESKTOP = Path.home() / "Desktop"
TARGET = next(DESKTOP.glob("*正文模板.docx"))
SOURCE = max(DESKTOP.glob("*格式修改前备份_*.docx"), key=lambda p: p.stat().st_mtime)

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ET.register_namespace("w", NS["w"])


def text_of(el):
    return "".join(t.text or "" for t in el.findall(".//w:t", NS))


def find_decl_idx(children):
    for i, child in enumerate(children):
        if "毕业设计原创性声明" in text_of(child):
            return i
    raise RuntimeError("Cannot locate declaration paragraph")


def restore_cover_xml():
    tmp = TARGET.with_name(f"{TARGET.stem}_cover_tmp.docx")
    shutil.copy2(TARGET, tmp)
    with zipfile.ZipFile(SOURCE, "r") as src_zip:
        src_xml = src_zip.read("word/document.xml")
    with zipfile.ZipFile(tmp, "r") as tgt_zip:
        tgt_files = {name: tgt_zip.read(name) for name in tgt_zip.namelist()}
        tgt_xml = tgt_files["word/document.xml"]

    src_root = ET.fromstring(src_xml)
    tgt_root = ET.fromstring(tgt_xml)
    src_body = src_root.find("w:body", NS)
    tgt_body = tgt_root.find("w:body", NS)
    src_children = list(src_body)
    tgt_children = list(tgt_body)
    src_cut = find_decl_idx(src_children)
    tgt_cut = find_decl_idx(tgt_children)

    for child in tgt_children[:tgt_cut]:
        tgt_body.remove(child)
    for offset, child in enumerate(src_children[:src_cut]):
        tgt_body.insert(offset, copy.deepcopy(child))

    tgt_files["word/document.xml"] = ET.tostring(tgt_root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(TARGET, "w", zipfile.ZIP_DEFLATED) as out_zip:
        for name, data in tgt_files.items():
            out_zip.writestr(name, data)

    tmp.unlink(missing_ok=True)


def remove_shading(element):
    for shd in list(element.findall(".//w:shd", NS)):
        parent = None
        for p in element.iter():
            if shd in list(p):
                parent = p
                break
        if parent is not None:
            parent.remove(shd)


def set_border(parent, name, val="single", size="8", color="000000"):
    node = parent.find(qn(f"w:{name}"))
    if node is None:
        node = OxmlElement(f"w:{name}")
        parent.append(node)
    node.set(qn("w:val"), val)
    if val != "nil":
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
    return node


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        set_border(borders, edge, val="nil")


def remove_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)
    for p in cell.paragraphs:
        p_pr = p._p.get_or_add_pPr()
        for shd in list(p_pr.findall(qn("w:shd"))):
            p_pr.remove(shd)


def set_three_line(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideH", "insideV"):
        set_border(borders, edge, val="nil")
    set_border(borders, "top", val="single", size="12")
    set_border(borders, "bottom", val="single", size="12")

    for row in table.rows:
        for cell in row.cells:
            remove_cell_shading(cell)
            clear_cell_borders(cell)

    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            set_border(borders, "bottom", val="single", size="8")


def format_tables():
    doc = Document(TARGET)
    for table in doc.tables:
        set_three_line(table)
    doc.save(TARGET)


def main():
    backup = TARGET.with_name(f"{TARGET.stem}_封面表格修改前备份_{datetime.now():%Y%m%d_%H%M%S}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)
    restore_cover_xml()
    format_tables()
    print(f"saved={TARGET}")
    print(f"backup={backup}")
    print(f"cover_source={SOURCE}")


if __name__ == "__main__":
    main()
