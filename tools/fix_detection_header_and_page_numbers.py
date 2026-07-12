from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCX = (
    Path.home()
    / "Desktop"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f_\u68c0\u6d4b\u62a5\u544a"
    / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
)


def set_font(run, size=9, name="宋体", color="808080"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(p):
    for child in list(p._p):
        p._p.remove(child)


def set_header_text(header, text):
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 9, "宋体", "808080")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, 9, "宋体", "808080")


def set_footer_page_number(footer):
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)


def main():
    backup = DOCX.with_name(
        f"{DOCX.stem}_页眉页码修改前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}"
    )
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    for section in doc.sections:
        section.footer_distance = Cm(1.75)

    # Cover keeps a blank first-page footer. The rest of the first section and
    # all body pages use Arabic page numbers so the detector can find real PAGE fields.
    if doc.sections:
        first = doc.sections[0]
        first.different_first_page_header_footer = True
        set_footer_page_number(first.footer)
        set_footer_page_number(first.even_page_footer)

    if len(doc.sections) > 1:
        body = doc.sections[1]
        body.header.is_linked_to_previous = False
        body.first_page_header.is_linked_to_previous = False
        body.even_page_header.is_linked_to_previous = False
        for header in (body.header, body.first_page_header, body.even_page_header):
            set_header_text(header, "河北建筑工程学院学士学位论文")

        body.footer.is_linked_to_previous = False
        body.first_page_footer.is_linked_to_previous = False
        body.even_page_footer.is_linked_to_previous = False
        for footer in (body.footer, body.first_page_footer, body.even_page_footer):
            set_footer_page_number(footer)

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")
    for i, section in enumerate(doc.sections, 1):
        print(i, round(section.footer_distance.cm, 2))


if __name__ == "__main__":
    main()
