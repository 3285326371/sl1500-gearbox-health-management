from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")
DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"
SHOT_DIR = ROOT / "docs" / "system_screenshots_updated"


def set_font(run, size=12, name="宋体", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def style_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        set_font(run, 12)


def style_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_font(run, 10.5)


def style_image(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True


def para_after(anchor, text=""):
    p = anchor._parent.add_paragraph(text)
    anchor._p.addnext(p._p)
    if text:
        style_body(p)
    return p


def image_after(anchor, image_name, width_cm=15.2):
    p = anchor._parent.add_paragraph()
    anchor._p.addnext(p._p)
    style_image(p)
    p.add_run().add_picture(str(SHOT_DIR / image_name), width=Cm(width_cm))
    return p


def main():
    backup = DOCX.with_name(
        f"{DOCX.stem}_第5章图片顺序修正前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}"
    )
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)

    replacements = {
        "单机齿轮箱详情页面如图5.7所示。": "单机齿轮箱详情页面如图5.4所示。",
        "图5.7 单机齿轮箱详情页面": "图5.4 单机齿轮箱详情页面",
        "齿轮箱HMI权限登录页面如图5.8所示。": "齿轮箱HMI权限登录页面如图5.5所示。",
        "图5.8 齿轮箱HMI权限登录页面": "图5.5 齿轮箱HMI权限登录页面",
        "运行总览页面，如图5.9所示。": "运行总览页面，如图5.6所示。",
        "图5.9 HMI登录后运行总览页面": "图5.6 HMI登录后运行总览页面",
        "健康报告页面如图5.4所示。": "健康报告页面如图5.8所示。",
        "图5.4 健康报告页面": "图5.8 健康报告页面",
        "AI运维问答页面如图5.5所示。": "AI运维问答页面如图5.9所示。",
        "图5.5 AI运维问答页面": "图5.9 AI运维问答页面",
        "系统参数设置页面如图5.6所示。": "系统参数设置页面如图5.10所示。",
        "图5.6 系统参数设置页面": "图5.10 系统参数设置页面",
    }

    inserted = False
    for p in list(doc.paragraphs):
        text = p.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            p.text = new_text
            if new_text.strip().startswith("图5."):
                style_caption(p)
            else:
                style_body(p)

        if p.text.strip().startswith("HMI 登录后页面将单机关键指标和控制入口集中到一个界面中") and not inserted:
            more = para_after(
                p,
                "HMI阈值设置页面如图5.7所示。该页面用于配置油温、轴承温度和振动阈值，"
                "使现场人员能够在登录后根据运行策略调整告警边界，并与后台参数接口保持一致。",
            )
            img = image_after(more, "07_hmi_threshold_page.png", 15.2)
            cap = para_after(img, "图5.7 HMI阈值设置页面")
            style_caption(cap)
            inserted = True

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")
    print(f"inserted_threshold={inserted}")


if __name__ == "__main__":
    main()
