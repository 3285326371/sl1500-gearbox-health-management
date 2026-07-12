from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def text(p):
    return p.text.strip().replace("\n", "")


def set_font(p, size=12, east="黑体", bold=False):
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
        r.font.size = Pt(size)
        r.bold = bold


def delete_para(p):
    p._element.getparent().remove(p._element)


doc = Document(DOCX)

# Rename body heading so it matches the replanned TOC.
for p in doc.paragraphs:
    if text(p).startswith("4.2.1 数据库表关系说明"):
        p.clear()
        p.add_run("4.2.1 数据表设计")
        p.style = "Heading 3"
        set_font(p)
        break

# Ensure 5.3.1 appears in the body, immediately after the real 5.3 heading.
body_started = False
body_531 = None
body_53 = None
for p in list(doc.paragraphs):
    t = text(p)
    if t.startswith("第1章") and "\t" not in p.text:
        body_started = True
    if not body_started:
        continue
    if t == "5.3.1 风场总览页面":
        body_531 = p
    if t.startswith("5.3 风场总览与单机详情实现"):
        body_53 = p
if body_531 is not None:
    delete_para(body_531)
if body_53 is not None:
    newp = body_53.insert_paragraph_before("5.3.1 风场总览页面")
    body_53._element.addnext(newp._element)
    newp.style = "Heading 3"
    set_font(newp)

doc.save(DOCX)
print(f"fixed={DOCX}")
