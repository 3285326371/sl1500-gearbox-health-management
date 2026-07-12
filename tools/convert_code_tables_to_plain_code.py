from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def set_intro_style(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def set_code_style(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.1)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10.5)


def insert_paragraph_before_table(doc: Document, table, text: str, code: bool = False):
    p = doc.add_paragraph(text)
    table._tbl.addprevious(p._p)
    if code:
        set_code_style(p)
    else:
        set_intro_style(p)
    return p


def is_code_table(table) -> bool:
    if len(table.rows) != 1 or len(table.columns) != 1:
        return False
    text = table.cell(0, 0).text.strip()
    return text.startswith(("def ", "async function", "function ")) or "\n    " in text and ("return" in text or "fetch(" in text)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_代码清单去表格前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)

    converted = 0
    for table in list(doc.tables):
        if not is_code_table(table):
            continue
        code_text = table.cell(0, 0).text.rstrip()
        insert_paragraph_before_table(doc, table, "核心代码如下：", code=False)
        for line in code_text.splitlines():
            # Word keeps indentation when the run text contains leading spaces.
            insert_paragraph_before_table(doc, table, line, code=True)
        table._element.getparent().remove(table._element)
        converted += 1

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")
    print(f"converted={converted}")


if __name__ == "__main__":
    main()
