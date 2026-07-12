from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


DESKTOP = Path.home() / "Desktop"
DOCX = next(p for p in DESKTOP.glob("*正文模板.docx") if not p.name.startswith("~$"))
WORK = Path.cwd() / "docs" / f"working_restructured_{datetime.now():%Y%m%d_%H%M%S}.docx"
ER_PNG = Path.cwd() / "docs" / "system_screenshots" / "gearbox_er_diagram.png"

TOC_LINES = [
    ("第1章 绪论", "", 0),
    ("1.1 研究背景与意义", "", 1),
    ("1.2 国内外研究现状", "", 1),
    ("1.2.1 传统故障诊断方法", "", 2),
    ("1.2.2 智能算法故障诊断方法", "", 2),
    ("1.2.3 现有技术的局限性", "", 2),
    ("1.3 本文研究内容", "", 1),
    ("1.4 技术路线", "", 1),
    ("第2章 系统需求分析", "", 0),
    ("2.1 业务需求分析", "", 1),
    ("2.2 功能需求分析", "", 1),
    ("2.2.1 健康监测需求", "", 2),
    ("2.2.2 智能诊断需求", "", 2),
    ("2.2.3 数据管理与报告需求", "", 2),
    ("2.3 非功能需求分析", "", 1),
    ("2.4 数据需求分析", "", 1),
    ("第3章 相关技术与理论基础", "", 0),
    ("3.1 齿轮箱典型故障机理", "", 1),
    ("3.2 多源状态监测技术", "", 1),
    ("3.3 振动特征提取方法", "", 1),
    ("3.3.1 时域特征", "", 2),
    ("3.3.2 包络特征", "", 2),
    ("3.4 风险评分与健康评估方法", "", 1),
    ("3.5 Web系统开发技术", "", 1),
    ("第4章 系统总体设计", "", 0),
    ("4.1 系统架构设计", "", 1),
    ("4.2 数据库设计", "", 1),
    ("4.2.1 数据表设计", "", 2),
    ("4.2.2 ER图设计", "", 2),
    ("4.3 后端接口设计", "", 1),
    ("4.4 诊断流程设计", "", 1),
    ("4.5 前端页面设计", "", 1),
    ("第5章 系统详细实现", "", 0),
    ("5.1 数据采集与融合模块实现", "", 1),
    ("5.2 故障诊断算法模块实现", "", 1),
    ("5.2.1 振动特征提取实现", "", 2),
    ("5.2.2 风险评分实现", "", 2),
    ("5.3 风场总览与单机详情实现", "", 1),
    ("5.3.1 风场总览页面", "", 2),
    ("5.3.2 单机详情页面", "", 2),
    ("5.4 故障代码库与数据管理实现", "", 1),
    ("5.5 AI运维问答模块实现", "", 1),
    ("5.6 健康报告与工单闭环实现", "", 1),
    ("第6章 系统测试与结果分析", "", 0),
    ("6.1 测试环境", "", 1),
    ("6.2 功能测试", "", 1),
    ("6.3 诊断场景测试分析", "", 1),
    ("6.4 结果评价", "", 1),
    ("第7章 总结与展望", "", 0),
    ("参考文献", "", 0),
    ("致谢", "", 0),
]


def set_run_font(run, size=12, east="宋体", west="Times New Roman", bold=None):
    run.font.name = west
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_para_font(p, size=12, east="宋体", west="Times New Roman", bold=None):
    for run in p.runs:
        set_run_font(run, size, east, west, bold)


def text(p):
    return p.text.strip().replace("\n", "")


def find_para(doc, starts):
    if isinstance(starts, str):
        starts = (starts,)
    # Most calls are meant to target the real正文, not the static TOC.
    body_start = 0
    for i, p in enumerate(doc.paragraphs):
        if text(p).startswith("第1章") and "\t" not in p.text:
            body_start = i
            break
    for p in doc.paragraphs:
        if para_index(doc, p) < body_start and not any(s.startswith("目") for s in starts):
            continue
        t = text(p)
        if any(t.startswith(s) for s in starts):
            return p
    return None


def para_index(doc, para):
    for i, p in enumerate(doc.paragraphs):
        if p._element is para._element:
            return i
    raise ValueError("paragraph not found")


def insert_before(anchor, lines):
    inserted = []
    for item in reversed(lines):
        p = anchor.insert_paragraph_before(item)
        inserted.append(p)
    return list(reversed(inserted))


def delete_para(p):
    el = p._element
    el.getparent().remove(el)


def delete_body_range(doc, start_para, end_para):
    body = doc._body._element
    children = list(body)
    start = children.index(start_para._element)
    end = children.index(end_para._element)
    for el in children[start:end]:
        body.remove(el)


def delete_from_para_to_end(doc, start_para):
    body = doc._body._element
    children = list(body)
    start = children.index(start_para._element)
    for el in children[start:]:
        if el.tag.endswith("sectPr"):
            continue
        body.remove(el)


def format_heading(p, level):
    if level == 1:
        p.style = "Heading 1"
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(21)
        set_para_font(p, 15, "黑体", "Times New Roman", False)
    elif level == 2:
        p.style = "Heading 2"
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        set_para_font(p, 14, "黑体", "Times New Roman", False)
    else:
        p.style = "Heading 3"
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)
        set_para_font(p, 12, "黑体", "Times New Roman", False)


def format_body_para(p):
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_para_font(p, 12, "宋体", "Times New Roman", None)


def create_er_image():
    ER_PNG.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 34)
        font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 24)
        small = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 20)
    except Exception:
        title_font = font = small = ImageFont.load_default()

    boxes = {
        "User": (80, 120, 370, 310, ["id", "username", "password_hash", "role"]),
        "TurbineAsset": (600, 80, 960, 310, ["unit_id", "model", "location", "status", "commission_date"]),
        "FaultRecord": (1090, 120, 1430, 350, ["id", "unit_id", "fault_type", "severity", "probability", "created_at"]),
        "SystemConfig": (80, 500, 400, 700, ["config_key", "config_value", "threshold", "description"]),
        "HealthReport": (610, 500, 970, 730, ["report_id", "unit_id", "health_score", "rul_days", "suggestion"]),
        "DiagnosisResult": (1090, 500, 1430, 730, ["result_id", "unit_id", "risk_score", "features", "advice"]),
    }
    draw.text((515, 24), "齿轮箱健康管理系统ER图", fill="black", font=title_font)
    for name, (x1, y1, x2, y2, fields) in boxes.items():
        draw.rectangle((x1, y1, x2, y2), outline="black", width=3)
        draw.rectangle((x1, y1, x2, y1 + 45), outline="black", fill="#F7F7F7", width=2)
        draw.text((x1 + 18, y1 + 10), name, fill="black", font=font)
        y = y1 + 58
        for f in fields:
            draw.text((x1 + 20, y), f, fill="black", font=small)
            y += 28

    def arrow(a, b, label):
        draw.line((a[0], a[1], b[0], b[1]), fill="black", width=3)
        draw.polygon([(b[0], b[1]), (b[0] - 12, b[1] - 8), (b[0] - 12, b[1] + 8)], fill="black")
        mx, my = (a[0] + b[0]) // 2, (a[1] + b[1]) // 2
        draw.text((mx - 35, my - 30), label, fill="black", font=small)

    arrow((960, 200), (1090, 200), "1:N")
    arrow((780, 310), (790, 500), "1:N")
    arrow((970, 610), (1090, 610), "1:1")
    arrow((400, 600), (610, 600), "参数")
    img.save(ER_PNG)


def remove_shading_and_set_three_line(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, val, size in [
        ("top", "single", "12"), ("bottom", "single", "12"),
        ("left", "nil", "0"), ("right", "nil", "0"),
        ("insideH", "nil", "0"), ("insideV", "nil", "0"),
    ]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        if val != "nil":
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), "000000")

    for ri, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant = tr_pr.find(qn("w:cantSplit"))
        if cant is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            for shd in list(tc_pr.findall(qn("w:shd"))):
                tc_pr.remove(shd)
            cb = tc_pr.find(qn("w:tcBorders"))
            if cb is None:
                cb = OxmlElement("w:tcBorders")
                tc_pr.append(cb)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                node = cb.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    cb.append(node)
                node.set(qn("w:val"), "nil")
            if ri == 0:
                bottom = cb.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:space"), "0")
                bottom.set(qn("w:color"), "000000")
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, 10.5, "宋体", "Times New Roman", True if ri == 0 else None)


def rebuild_toc(doc):
    title = None
    for p in doc.paragraphs:
        if "目" in text(p) and "录" in text(p) and len(text(p)) <= 6:
            title = p
            break
    first_chapter = None
    seen_toc = False
    for p in doc.paragraphs:
        if p._element is title._element:
            seen_toc = True
            continue
        if seen_toc and text(p).startswith("第1章") and "\t" not in p.text:
            first_chapter = p
            break
    if not title or not first_chapter:
        raise RuntimeError("Cannot locate TOC or first chapter")
    delete_body_range(doc, doc.paragraphs[para_index(doc, title) + 1], first_chapter)
    for item in reversed(TOC_LINES):
        label, page, level = item
        p = first_chapter.insert_paragraph_before("")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24 * level)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(label)
        set_run_font(r, 12 if level == 0 else 10.5, "宋体", "Times New Roman", True if level == 0 else False)
        p.add_run("\t")
        r2 = p.add_run(page)
        set_run_font(r2, 12 if level == 0 else 10.5, "宋体", "Times New Roman", True if level == 0 else False)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    set_para_font(title, 16, "黑体", "Times New Roman", True)


def insert_missing_headings(doc):
    # Chapter 1 tertiary headings.
    p = find_para(doc, "1.2 国内外研究现状")
    if p and not find_para(doc, "1.2.1"):
        first_after = doc.paragraphs[para_index(doc, p) + 1]
        first_after.insert_paragraph_before("1.2.1 传统故障诊断方法")
        scada = find_para(doc, "SCADA 数据分析")
        if scada:
            scada.insert_paragraph_before("1.2.2 智能算法故障诊断方法")
        oil = find_para(doc, "油液监测")
        if oil:
            oil.insert_paragraph_before("1.2.3 现有技术的局限性")
    # Demand tertiary headings.
    p = find_para(doc, "系统功能需求包括")
    if p and not find_para(doc, "2.2.1"):
        p.insert_paragraph_before("2.2.1 健康监测需求")
        diag = find_para(doc, "智能诊断模块包括")
        if diag:
            diag.insert_paragraph_before("2.2.2 智能诊断需求")
        data = find_para(doc, "数据管理模块展示")
        if data:
            data.insert_paragraph_before("2.2.3 数据管理与报告需求")
    # Vibration tertiary headings.
    p = find_para(doc, "3.3 振动特征提取方法")
    if p and not find_para(doc, "3.3.1"):
        nxt = doc.paragraphs[para_index(doc, p) + 1]
        nxt.insert_paragraph_before("3.3.1 时域特征")
        p34 = find_para(doc, "3.4 风险评分")
        if p34:
            p34.insert_paragraph_before("3.3.2 包络特征")
            p34.insert_paragraph_before("包络分析用于突出齿轮和轴承局部损伤产生的调制冲击。当齿面点蚀、轴承剥落或断齿出现时，包络峰值与冲击类指标会同步升高，可作为系统规则诊断的重要依据。")
    # Database ER heading and image.
    p43 = find_para(doc, "4.3 后端接口设计")
    if p43 and not find_para(doc, "4.2.2 ER图设计"):
        create_er_image()
        p43.insert_paragraph_before("4.2.2 ER图设计")
        explain = p43.insert_paragraph_before("系统数据库围绕用户、风机资产、故障记录、诊断结果、健康报告和系统参数展开。风机资产表是业务核心，故障记录、诊断结果和健康报告均通过机组编号与其关联，系统参数表为阈值配置和诊断规则提供支撑。")
        imgp = p43.insert_paragraph_before("")
        imgp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        imgp.add_run().add_picture(str(ER_PNG), width=Inches(5.8))
        cap = p43.insert_paragraph_before("图4.2 系统数据库ER图")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Implementation tertiary headings.
    p52 = find_para(doc, "5.2 故障诊断算法")
    if p52 and not find_para(doc, "5.2.1"):
        nxt = doc.paragraphs[para_index(doc, p52) + 1]
        nxt.insert_paragraph_before("5.2.1 振动特征提取实现")
        p53 = find_para(doc, "5.3 风场总览")
        if p53:
            p53.insert_paragraph_before("5.2.2 风险评分实现")
            p53.insert_paragraph_before("风险评分实现将振动、油温、油液等级和故障严重度映射为统一风险值，再计算健康评分和剩余寿命。该方法虽然属于规则融合模型，但输出过程清晰，适合毕业设计原型系统展示和运维解释。")
    p = find_para(doc, "图5.1")
    if p and not find_para(doc, "5.3.1"):
        p.insert_paragraph_before("5.3.1 风场总览页面")
    p = find_para(doc, "图5.7")
    if p and not find_para(doc, "5.3.2"):
        p.insert_paragraph_before("5.3.2 单机详情页面")


def trim_content(doc):
    # Remove appendices entirely to control pages and word count.
    appendix = find_para(doc, "附录A")
    if appendix:
        delete_from_para_to_end(doc, appendix)

    # Compress late testing chapter.
    start = find_para(doc, "6.5 测试用例设计")
    end = find_para(doc, "第7章 总结与展望")
    if start and end:
        delete_body_range(doc, start, end)
        insert_before(end, [
            "6.4 结果评价",
            "测试结果表明，系统能够完成风场总览、单机详情、实时监测、故障诊断、故障数据管理、健康报告和AI运维问答等主要功能。对于油温升高、振动增强、油液污染和轴承冲击等典型场景，系统能够给出风险等级、故障类型、健康评分、RUL估计和维护建议。",
            "从运行效果看，系统界面能够支持从风场总览到单机详情的下钻分析，故障诊断结果具有一定可解释性。系统不足在于当前数据以模拟数据和统计文件为主，后续仍需接入真实SCADA、CMS和油液监测数据，并利用更多真实故障样本对模型进行训练和校验。",
        ])

    # Remove sections no longer included in the replanned TOC.
    for start_label, end_label in [
        ("2.5 需求约束", "第3章"),
        ("3.6 健康评分", "第4章"),
    ]:
        start = find_para(doc, start_label)
        end = find_para(doc, end_label)
        if start and end:
            delete_body_range(doc, start, end)

    # Delete selected long/repetitive paragraphs.
    remove_starts = [
        "在智能运维平台方面",
        "系统根据这些数据形成机组快照",
        "在数据来源方面，系统当前采用",
        "扩展性方面，系统当前虽然使用SQLite",
        "从可用性角度看，系统页面采用卡片",
    ]
    for p in list(doc.paragraphs):
        t = text(p)
        if any(t.startswith(s) for s in remove_starts):
            delete_para(p)


def format_document(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    for p in doc.paragraphs:
        t = text(p)
        if not t:
            continue
        if t.startswith("第") and "章" in t:
            format_heading(p, 1)
        elif t[:3].replace(".", "").isdigit() and t.count(".") >= 2 and " " in t[:8]:
            format_heading(p, 3)
        elif t[:2].replace(".", "").isdigit() and t.count(".") >= 1 and " " in t[:6]:
            format_heading(p, 2)
        elif t.startswith(("图", "表", "代码清单")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            set_para_font(p, 10.5, "宋体", "Times New Roman", False)
        elif t in ("摘要", "摘  要", "ABSTRACT", "参考文献", "致谢"):
            format_heading(p, 1)
        elif t == "目  录":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            set_para_font(p, 16, "黑体", "Times New Roman", True)
        elif "{" in t[:2] or t.startswith(("def ", "async ", "function ")):
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                set_run_font(r, 9, "Consolas", "Consolas", False)
        elif p._element.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            format_body_para(p)

    for shape in doc.inline_shapes:
        if shape.width and shape.width > Inches(5.8):
            ratio = shape.height / shape.width
            shape.width = Inches(5.8)
            shape.height = int(shape.width * ratio)

    for table in doc.tables:
        remove_shading_and_set_three_line(table)


def main():
    print(f"source_docx={DOCX}")
    backup = DOCX.with_name(f"{DOCX.stem}_重构40页前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    shutil.copy2(DOCX, WORK)
    doc = Document(WORK)

    trim_content(doc)
    insert_missing_headings(doc)
    rebuild_toc(doc)
    format_document(doc)

    doc.save(WORK)
    shutil.copy2(WORK, DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")
    print(f"work={WORK}")


if __name__ == "__main__":
    main()
