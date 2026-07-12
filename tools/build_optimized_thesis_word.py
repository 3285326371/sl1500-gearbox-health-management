from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
TEMPLATE = next(DESKTOP.glob("*正文模板*.docx"))
SOURCE_MD = ROOT / "docs" / "基于项目代码分析的毕业论文初稿.md"
OUT = ROOT / "docs" / "华锐SL1500型双馈式风机齿轮箱智能健康管理系统毕业论文_代码分析优化版.docx"

TITLE = "华锐SL1500型双馈式风机齿轮箱智能健康管理系统设计与实现"
SCHOOL = "河北建筑工程学院"
MAJOR = "计算机科学与技术"
DEPARTMENT = "信息工程学院"


def clear_document(doc):
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_page(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal.font.size = Pt(10.5)


def run_font(run, font="宋体", size=10.5, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if font == "宋体" else "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", first_indent=True, align=None, size=10.5, font="宋体", bold=False, line=1.5, after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = line
    p.paragraph_format.space_after = Pt(after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    run_font(r, font, size, bold)
    return p


def add_center(doc, text, size=16, bold=True, font="黑体", before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    run_font(r, font, size, bold)
    return p


def add_chapter(doc, text):
    p = add_center(doc, text, size=16, font="黑体", after=14)
    p.paragraph_format.page_break_before = True
    return p


def add_h1(doc, text):
    p = add_para(doc, text, first_indent=False, size=14, font="黑体", bold=True, line=1.5, after=4)
    p.paragraph_format.space_before = Pt(8)
    return p


def add_h2(doc, text):
    p = add_para(doc, text, first_indent=False, size=12, font="黑体", bold=True, line=1.5, after=3)
    p.paragraph_format.space_before = Pt(6)
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right}
    for edge, spec in specs.items():
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        if spec is None:
            spec = {"val": "nil"}
        for key, val in spec.items():
            el.set(qn("w:" + key), str(val))


def set_cell(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    run_font(r, "宋体", 9.5, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ("top", "left", "bottom", "right"):
        el = tc_mar.find(qn("w:" + name))
        if el is None:
            el = OxmlElement("w:" + name)
            tc_mar.append(el)
        el.set(qn("w:w"), "90")
        el.set(qn("w:type"), "dxa")


def apply_three_line(table):
    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin = {"val": "single", "sz": "6", "color": "000000"}
    nil = {"val": "nil"}
    rows = len(table.rows)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            if ri == 0:
                set_cell_border(cell, top=thick, bottom=thin, left=nil, right=nil)
            elif ri == rows - 1:
                set_cell_border(cell, top=nil, bottom=thick, left=nil, right=nil)
            else:
                set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil)


def add_table(doc, caption, headers, rows, widths=None):
    cap = add_para(doc, caption, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font="黑体", bold=True, line=1.25, after=2)
    cap.paragraph_format.space_before = Pt(6)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, bold=True, center=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value, center=len(str(value)) <= 12)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    apply_three_line(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_cover(doc):
    add_center(doc, "学士学位毕业设计", 20, font="黑体", before=42, after=42)
    add_center(doc, "华锐SL1500型双馈式风机齿轮箱", 18, font="黑体", after=6)
    add_center(doc, "智能健康管理系统设计与实现", 18, font="黑体", after=46)
    rows = [
        ("姓    名", "____________"),
        ("学    号", "____________"),
        ("院    系", DEPARTMENT),
        ("专    业", MAJOR),
        ("班    级", "____________"),
        ("指导教师", "____________"),
    ]
    for label, value in rows:
        add_center(doc, f"{label}  {value}", size=14, bold=False, font="宋体", after=8)
    add_center(doc, "2026 年 05 月", size=14, bold=False, font="宋体", before=54)


def add_declarations(doc):
    doc.add_page_break()
    add_center(doc, "毕业设计原创性声明", 16, font="黑体", after=20)
    add_para(doc, f"本人所提交的毕业设计《{TITLE}》，是在导师的指导下，独立进行研究工作所取得的原创性成果。除文中已经注明引用的内容外，本毕业设计不包含任何其他个人或集体已经发表或撰写过的研究成果。对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。", True)
    add_para(doc, "本声明的法律后果由本人承担。", True)
    add_para(doc, "毕业设计作者（签名）：                指导教师确认（签名）：", False)
    add_para(doc, "年   月   日                           年   月   日", False)
    doc.add_page_break()
    add_center(doc, "毕业设计版权使用授权书", 16, font="黑体", after=20)
    add_para(doc, f"本毕业设计作者完全了解{SCHOOL}有关保留、使用毕业设计的规定，同意学校保留并向国家有关部门或机构送交毕业设计的复印件和电子文档，允许毕业设计被查阅和借阅。本人授权学校可以将毕业设计的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或其他复制手段保存、汇编本毕业设计。", True)
    add_para(doc, "保密的毕业设计在_______年解密后适用本授权书。", True)
    add_para(doc, "毕业设计作者（签名）：               指导教师（签名）：", False)
    add_para(doc, "年   月   日                      年   月   日", False)


def add_toc(doc):
    doc.add_page_break()
    add_center(doc, "目  录", 16, font="黑体", after=18)
    entries = [
        (0, "摘要", "I"), (0, "ABSTRACT", "II"),
        (0, "第1章 绪论", "1"), (1, "1.1 研究背景与意义", "1"), (1, "1.2 国内外研究现状", "2"), (1, "1.3 本文研究内容", "3"), (1, "1.4 技术路线", "4"),
        (0, "第2章 系统需求分析", "5"), (1, "2.1 业务需求分析", "5"), (1, "2.2 功能需求分析", "6"), (1, "2.3 非功能需求分析", "7"), (1, "2.4 数据需求分析", "8"),
        (0, "第3章 相关技术与理论基础", "9"), (1, "3.1 齿轮箱典型故障机理", "9"), (1, "3.2 多源状态监测技术", "10"), (1, "3.3 振动特征提取方法", "11"), (1, "3.4 风险评分与健康评估方法", "12"), (1, "3.5 Web系统开发技术", "13"),
        (0, "第4章 系统总体设计", "14"), (1, "4.1 系统架构设计", "14"), (1, "4.2 数据库设计", "15"), (1, "4.3 后端接口设计", "16"), (1, "4.4 诊断流程设计", "17"), (1, "4.5 前端页面设计", "18"),
        (0, "第5章 系统详细实现", "19"), (1, "5.1 数据采集与融合模块实现", "19"), (1, "5.2 故障诊断算法模块实现", "20"), (1, "5.3 风场总览与单机详情实现", "21"), (1, "5.4 故障代码库与数据管理实现", "22"), (1, "5.5 AI运维问答模块实现", "23"), (1, "5.6 健康报告与工单闭环实现", "24"),
        (0, "第6章 系统测试与结果分析", "25"), (1, "6.1 测试环境", "25"), (1, "6.2 功能测试", "26"), (1, "6.3 诊断场景测试分析", "27"), (1, "6.4 结果评价", "28"),
        (0, "第7章 总结与展望", "29"), (0, "参考文献", "31"), (0, "致谢", "32"), (0, "附录A 系统主要接口", "33"),
    ]
    for level, title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.45
        p.paragraph_format.left_indent = Cm(0.74 if level else 0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{title}\t{page}")
        run_font(r, "宋体", 10.5, False)


def clean_inline(text):
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def extract_formal_body():
    text = SOURCE_MD.read_text(encoding="utf-8")
    text = text.split("## 二、完整论文初稿", 1)[1]
    text = text.replace("### Abstract", "### ABSTRACT")
    return text


EXTRA_TABLES = {
    "#### 2.2 功能需求分析": (
        "表2.1 系统功能需求汇总",
        ["功能模块", "主要功能", "项目代码对应"],
        [
            ("用户与权限", "登录、注册、角色显示、管理员默认账号初始化", "auth_route.py、database.py"),
            ("健康监测", "风场总览、单机详情、实时状态、振动波形、SSE推送", "windfarm_route.py、data_route.py"),
            ("智能诊断", "特征提取、风险评分、故障分类、健康评分和RUL估计", "ml_models.py"),
            ("数据管理", "故障记录列表、历史故障统计、CSV导出", "data_route.py、wind_data_repository.py"),
            ("运维问答", "本地知识库检索、实时工况分析、可选大模型增强", "rag_service.py、qa_route.py"),
            ("报告与工单", "健康报告、复检建议、P1/P2工单模拟", "report_route.py、ops_route.py"),
        ],
        [3.0, 7.0, 5.8],
    ),
    "#### 3.4 风险评分与健康评估方法": (
        "表3.1 齿轮箱健康评估主要风险因子",
        ["风险因子", "含义", "诊断作用"],
        [
            ("振动RMS", "反映整体振动能量", "识别磨损、松动和运行异常"),
            ("峭度", "反映冲击尖峰程度", "识别轴承剥落、齿面点蚀等早期冲击"),
            ("峰值因子", "峰值与有效值比值", "增强对局部冲击的敏感性"),
            ("包络峰值", "反映调制冲击强度", "辅助判断轴承和齿轮局部损伤"),
            ("齿轮箱油温", "反映热状态和冷却效率", "识别过温和冷却系统异常"),
            ("油液NAS", "反映润滑油污染程度", "识别润滑劣化和磨粒风险"),
        ],
        [3.2, 5.6, 7.0],
    ),
    "#### 4.2 数据库设计": (
        "表4.1 数据库表结构说明",
        ["数据表", "关键字段", "用途"],
        [
            ("User", "username、password、role", "保存用户账号、密码哈希和角色"),
            ("FaultRecord", "unit_id、timestamp、fault_type、severity、probability、status", "保存诊断结果和处理状态"),
            ("TurbineAsset", "unit_id、number、model、design_life_years", "保存风机资产基础信息"),
            ("SystemConfig", "config_key、config_value、description", "保存油温、振动、油液等阈值配置"),
        ],
        [3.0, 7.2, 5.6],
    ),
    "#### 4.3 后端接口设计": (
        "表4.2 系统主要后端接口",
        ["接口类别", "典型接口", "功能说明"],
        [
            ("数据接口", "/api/data/status、/api/data/diagnosis、/api/data/trend", "状态采集、诊断计算和趋势分析"),
            ("风场接口", "/api/windfarm/overview、/api/windfarm/turbine/<unit_id>", "风场总览和单机详情"),
            ("报告接口", "/api/report/generate、/api/report/health-summary", "生成健康报告和摘要"),
            ("问答接口", "/api/qa/ask", "运维知识问答和实时状态分析"),
            ("运维接口", "/api/ops/workorders、/api/ops/alarm-policy", "工单模拟、报警策略和模型验证"),
        ],
        [3.0, 6.4, 6.4],
    ),
    "#### 6.2 功能测试": (
        "表6.1 系统功能测试结果",
        ["测试项目", "测试内容", "结果"],
        [
            ("登录注册", "管理员登录、新用户注册和角色显示", "通过"),
            ("风场总览", "多机组状态、统计指标和机组卡片展示", "通过"),
            ("单机详情", "切换机组后关键指标同步变化", "通过"),
            ("实时监测", "SSE推送油温、振动、功率和波形数据", "通过"),
            ("故障诊断", "不同故障场景输出诊断结果和建议", "通过"),
            ("报告导出", "生成健康报告并导出故障记录CSV", "通过"),
            ("运维问答", "基于知识库和当前工况生成排查建议", "通过"),
        ],
        [3.0, 9.0, 2.8],
    ),
}


def add_formal_content(doc):
    body = extract_formal_body()
    in_refs = False
    toc_inserted = False
    for raw in body.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line in EXTRA_TABLES:
            # Add heading first, then table.
            add_h1(doc, clean_inline(line[5:]))
            caption, headers, rows, widths = EXTRA_TABLES[line]
            add_table(doc, caption, headers, rows, widths)
            continue
        if line.startswith("### "):
            title = clean_inline(line[4:])
            if title == "摘要":
                doc.add_page_break()
                add_center(doc, "摘  要", 16, font="黑体", after=18)
            elif title == "ABSTRACT":
                doc.add_page_break()
                add_center(doc, "ABSTRACT", 16, font="Times New Roman", after=18)
            elif re.match(r"第\d+章", title):
                if not toc_inserted:
                    add_toc(doc)
                    toc_inserted = True
                add_chapter(doc, title)
            elif title in {"参考文献", "致谢", "附录A 系统主要接口"}:
                add_chapter(doc, title)
                in_refs = title == "参考文献"
            else:
                add_chapter(doc, title)
            continue
        if line.startswith("#### "):
            add_h1(doc, clean_inline(line[5:]))
            in_refs = False
            continue
        if line.startswith("|"):
            # Appendix table is added later as a properly formatted Word table.
            continue
        if line.startswith("- "):
            p = add_para(doc, "（1）" + clean_inline(line[2:]), first_indent=True)
            continue
        if line.startswith("[") and "]" in line[:6]:
            add_para(doc, clean_inline(line), first_indent=False, line=1.25)
            continue
        if line.startswith("关键词"):
            add_para(doc, clean_inline(line), first_indent=False, line=1.5)
            continue
        if line.startswith("Key words"):
            add_para(doc, clean_inline(line), first_indent=False, line=1.5)
            continue
        add_para(doc, clean_inline(line), first_indent=not in_refs, line=1.5 if not in_refs else 1.25)

    add_table(doc, "表A.1 系统主要接口", ["接口", "功能"], [
        ("/health", "后端服务健康检查"),
        ("/api/auth/login", "用户登录"),
        ("/api/auth/register", "用户注册"),
        ("/api/data/status", "获取实时运行状态"),
        ("/api/data/vibration", "获取振动波形和特征"),
        ("/api/data/fault-codes", "获取齿轮箱故障代码库"),
        ("/api/data/diagnosis", "执行故障诊断"),
        ("/api/data/records", "获取故障记录"),
        ("/api/data/records/export", "导出故障记录CSV"),
        ("/api/data/trend", "获取故障趋势与分布"),
        ("/api/windfarm/overview", "获取风场总览"),
        ("/api/windfarm/turbine/<unit_id>", "获取单机详情"),
        ("/api/report/generate", "生成健康报告"),
        ("/api/qa/ask", "运维辅助问答"),
        ("/api/ops/workorders", "工单模拟接口"),
    ], [6.2, 9.8])


def main():
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    set_page(doc)
    add_cover(doc)
    add_declarations(doc)
    add_formal_content(doc)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
