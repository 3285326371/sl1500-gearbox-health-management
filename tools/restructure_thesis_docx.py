from __future__ import annotations

import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板.docx")
OUT = SRC.with_name("毕业设计正文模板_图表精简改稿版.docx")
ASSET_DIR = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\图表精简改稿_assets")


ABSTRACT_CN = [
    "本文面向华锐 SL1500 型双馈式风机齿轮箱，设计并实现了一套用于毕业设计原型验证的智能健康管理系统。系统将模拟 SCADA/CMS 数据、风机数据摘要、油温、振动、油液 NAS 等信息接入同一平台，采用 Flask、SQLite、HTML/CSS/JavaScript、ECharts、Three.js 和 SSE 完成风场总览、单机详情、实时监测、故障诊断、健康报告和运维问答等功能。",
    "诊断部分提取振动 RMS、峰值、峭度、峰值因子、脉冲因子和包络峰值，并结合油温、油液等级和历史故障记录计算综合风险。系统可识别齿面磨损、点蚀剥落、断齿、轴承损伤、油温过高、油液污染、轴系不对中、基础松动和冷却异常等典型场景，输出健康评分、RUL、诊断依据和维护建议。测试结果表明，该原型能够形成“数据采集-风险诊断-记录保存-报告输出-工单闭环”的基本流程，为风机齿轮箱智能运维系统开发提供参考。",
]

ABSTRACT_EN = [
    "This thesis designs and implements an intelligent health management prototype for the gearbox of the Sinovel SL1500 doubly-fed wind turbine. The system integrates simulated SCADA/CMS data, wind turbine summary files, oil temperature, vibration and oil quality indicators, and provides wind farm overview, turbine detail monitoring, real-time display, fault diagnosis, health reports and operation assistance through Flask, SQLite, ECharts, Three.js and SSE.",
    "The diagnosis module extracts RMS, peak value, kurtosis, crest factor, impulse factor and envelope peak from vibration signals, and combines oil temperature, NAS oil grade and historical fault records to calculate risk factors. Typical gearbox faults such as gear wear, pitting, broken teeth, bearing damage, high oil temperature, oil contamination, misalignment, looseness and cooling failure can be identified. Test results show that the prototype supports a basic closed loop from data acquisition and risk diagnosis to record storage, report generation and maintenance suggestions.",
]


REPLACEMENTS = {
    "风力发电是新能源电力系统的重要组成部分。随着“双碳”目标推进和风电装机容量增长，风电场从单纯追求装机规模逐步转向追求设备可靠性、发电效率和运维经济性。风电机组通常安装在山地、戈壁、海岸或高寒地区，运行环境复杂，设备长期受到风速波动、温度变化、沙尘、湿度、盐雾和并网冲击等因素影响。对于大型风机而言，关键部件发生故障后维修周期长、吊装成本高、备件费用高，因此状态监测和预测性维护已经成为风电运维的重要方向。": "风电场运维正在从定期检修转向状态检修。齿轮箱一旦发生严重故障，通常会带来长时间停机、吊装作业和高额备件成本，因此需要把温度、振动、油液和历史故障信息集中到可分析、可追溯的平台中。",
    "齿轮箱是双馈式风机传动链中的核心部件，其主要作用是将风轮低速大扭矩运动转化为适合发电机运行的高速运动。齿轮箱内部通常包含多级齿轮、轴承、润滑油路、冷却系统和箱体结构，工作时承受强烈交变载荷和冲击载荷。华锐 SL1500 型双馈式风机在我国早期风电场中应用较多，随着服役时间增加，齿轮箱油温偏高、轴承振动升高、齿轮啮合异常、油液污染和冷却系统故障等问题逐渐成为影响运行可靠性的关键因素。": "华锐 SL1500 型双馈式风机服役时间较长，齿轮箱常见风险集中在油温、轴承振动、齿轮啮合、油液污染和冷却效率等方面。本文围绕这些可观测指标开展系统设计。",
    "传统齿轮箱运维主要依赖人工巡检、定期检修和单一阈值报警。人工巡检能够发现明显异常，但对早期微弱故障和趋势性退化识别不足；单一阈值报警能够快速提示超限状态，但难以区分负荷变化、环境温度变化和设备劣化之间的关系。例如齿轮箱油温升高可能由环境温度、负荷升高、冷却风扇异常、润滑油劣化或轴承摩擦加剧等多种因素引起，仅靠固定阈值难以给出准确维护建议。因此，有必要构建融合多源数据、诊断模型和知识库的智能健康管理系统。": "单一阈值报警只能提示超限，难以说明异常来源。本文系统把指标趋势、故障规则和维护建议关联起来，使告警结果能够进一步转化为排查路径。",
    "本课题围绕华锐 SL1500 型双馈式风机齿轮箱健康管理展开，目标是将风机实时状态、历史故障统计、振动特征、油温趋势、油液状态和专家知识整合到统一平台中，实现状态监测、故障诊断、健康评估、RUL 估计和运维建议生成。该系统能够帮助运维人员快速掌握多台机组运行状态，识别重点风险机组，形成检修建议和报告，对降低非计划停机、提高机组可利用率和优化检修资源具有一定参考价值。": "课题目标是形成一个可演示、可扩展的健康管理原型，支撑风场总览、单机下钻、故障诊断、健康报告和工单闭环。",
    "本文基于实际项目代码和风机数据文件，完成华锐SL1500型双馈式风机齿轮箱智能健康管理体系的设计与实现。主要研究内容如下：": "本文研究内容由故障机理、数据融合、诊断模型、Web 平台和系统测试五部分组成，见表1.1。",
    "本文技术路线可以概括为：需求分析与故障机理研究、数据接入与预处理、特征提取与风险评分、故障诊断与健康评估、后端接口与数据库设计、前端可视化实现、系统测试与结果分析。系统首先获取齿轮箱油温、振动、油液、功率、风速、转速、故障码等数据；其次进行滤波、异常值处理和字段归一化；然后提取振动特征并计算风险因子；最后输出故障诊断、健康评分、RUL、复检周期和工单建议，并通过前端页面展示给运维人员。": "本文技术路线见图1.1。整体思路是先确定齿轮箱故障机理和业务需求，再完成数据接入、特征计算、风险评分、诊断输出和前端展示。",
    "该技术路线的重点在于把机械故障机理与软件系统实现结合起来。一方面，论文需要说明齿轮箱故障产生的物理原因和监测指标意义；另一方面，系统需要把这些指标转化为可操作的页面、接口和数据记录。通过这种方式，本文不是单独讨论算法，也不是单独开发管理页面，而是围绕齿轮箱健康管理目标形成较完整的技术闭环。": "该路线强调工程闭环：监测指标不只用于展示，还要参与诊断、记录、报告和维护建议生成。",
    "系统功能需求包括用户管理、健康监测、智能诊断、数据管理、报告生成、参数配置和运维问答七个方面。": "系统功能需求按模块梳理如表2.1所示，后续章节围绕这些功能进行设计与实现。",
    "系统需要具备较好的可用性、可扩展性和可解释性。可用性方面，界面应直观展示关键指标，避免运维人员在多个页面之间频繁查找数据。可扩展性方面，系统应能够在后续接入真实 SCADA、CMS、油液监测和工单系统。可解释性方面，诊断结果不能只输出故障名称，还应说明风险因子、诊断依据、健康评分、RUL 和建议措施。": "非功能需求主要包括可用性、可扩展性、可解释性和轻量部署，见表2.2。",
    "项目数据来源包括两类。第一类是系统模拟采集数据，包含齿轮箱油温、高速轴承振动、油液颗粒度和有功功率。第二类是风机数据文件夹生成的数据摘要，包括自由报表、实时数据、故障统计和故障信息统计。自由报表包含风机名称、平均有功功率、平均风速、最大风速和环境温度；实时数据包含日期、发电机转速、有功功率、叶片角度、齿轮箱油温或齿轮箱加热、机舱位置等字段；故障统计包含故障天数、故障总数和峰值日期；故障信息统计包含故障码、出现次数、开始时间和结束时间。": "系统数据需求见表2.3。两类数据共同支撑风场总览、单机详情、诊断判断和健康报告。",
    "风机齿轮箱故障可按照部件和机理分为齿轮故障、轴承故障、润滑故障、冷却故障、轴系故障和结构故障。齿轮故障包括齿面磨损、点蚀剥落、断齿和啮合异常。齿面磨损通常由润滑不足、载荷波动、异物颗粒和长期疲劳引起，会导致啮合频率附近能量升高；点蚀剥落属于接触疲劳损伤，常伴随冲击成分增强；断齿属于严重故障，可能出现周期性强冲击并引起快速停机风险。": "齿轮箱故障类型、典型表现和监测指标见表3.1。本文将故障机理转化为规则库条件和页面提示信息。",
    "振动信号是机械故障诊断的重要依据。本文系统对振动信号提取以下特征：RMS 用于反映整体振动能量；峰值用于反映瞬时冲击强度；峭度用于识别脉冲冲击和局部损伤；峰值因子用于描述峰值相对于有效值的突出程度；脉冲因子用于描述峰值相对于平均绝对值的大小；包络峰值用于捕捉调制冲击特征。": "振动特征的含义和用途见表3.2。系统重点关注能量类、冲击类和包络类指标。",
    "系统将多个特征和工况指标归一化为风险因子，包括振动 RMS、冲击峭度、峰值因子、包络峰值、齿轮箱油温和油液 NAS。每个风险因子根据经验阈值映射到 0 到 1 的区间，再采用加权求和得到综合风险值。振动 RMS 和冲击峭度权重较高，油温和油液状态作为重要补充。": "健康评估过程如图3.1所示。各类指标先归一化为风险因子，再汇总为综合风险、健康评分和 RUL。",
    "系统总体采用 B/S 架构，用户通过浏览器访问前端页面，前端通过 HTTP 接口和 SSE 通道与 Flask 后端通信。后端由应用入口、路由层、服务层、数据模型层和数据仓库层组成。应用入口负责创建 Flask 应用、配置数据库和注册蓝图；路由层提供认证、数据、风场、报告、问答、设置和运维接口；服务层负责数据模拟、数据融合、诊断计算和知识问答；数据模型层负责数据库表定义；数据仓库层负责读取风机数据摘要文件。": "系统采用 B/S 架构，前端通过 HTTP 和 SSE 与 Flask 后端交互，后端再调用服务层、数据库模型和风机数据仓库。系统数据流见图4.1。",
    "系统诊断流程如图4.3所示。该流程以多源状态数据为输入，经过数据清洗、特征提取、风险因子计算和规则库诊断后，形成健康评分、RUL估计、诊断依据与运维建议，并通过健康报告、故障记录和工单模拟实现闭环管理。": "诊断流程以多源状态数据为输入，经清洗、特征提取、风险评分和规则判断后，输出故障类型、健康评分、RUL、诊断依据和维护建议。关键输入输出见表4.5。",
    "前端首页包含登录、注册和主应用界面。用户通过登录页进入系统，系统根据角色显示高级工程师或运维人员身份；注册页可选择运维人员或管理员角色。登录后默认进入健康总览，左侧导航按照健康监测区、智能诊断区和系统配置区分组，顶部显示当前模块说明、通知状态和用户身份。风场总览页面展示多台机组卡片，单机详情页面展示齿轮箱温度、寿命、部件状态、告警和集群对标。实时监测页面进一步改为多标签结构，包含运行概览、振动趋势、孪生与采集、故障统计等页面。": "前端页面按“总览、下钻、诊断、报告、配置”组织，页面职责见表4.7。",
    "故障诊断页面是系统的核心业务页面。用户可以选择不同机组，也可以选择正常、齿轮磨损、点蚀剥落、轴承过温、油液污染、冷却故障等场景进行演示。点击开始诊断后，前端调用后端诊断接口，后端完成振动特征提取、风险评分、故障分类和结果解释，然后将结构化结果返回页面展示。": "故障诊断页面负责把算法输出转化为运维人员可读的结论，包括风险等级、关联部件、置信度、健康评分、RUL、依据和建议。",
    "系统提供多种诊断场景。正常运行场景中，油温、振动和油液质量均处于较低风险范围，系统输出正常运行，健康评分较高，维护建议为保持日常巡检。齿轮齿面磨损场景中，振动能量和油液污染风险轻度升高，系统输出齿轮齿面磨损，并建议 72 小时内复测振动趋势和取油样复核。": "诊断场景测试结果见表6.2。系统通过不同输入变化验证故障类型、风险等级和维护建议是否一致。",
    "从功能完整性看，系统覆盖了毕业设计要求中的登录、系统管理、用户管理、故障数据、数据分析、参数设置、故障分析、AI 智能问答和知识库等内容。从工程流程看，系统能够形成从数据采集、特征提取、诊断判断、健康评分、RUL 估计、建议生成、记录保存到报告和工单的闭环。从展示效果看，系统页面较适合风电运维场景，具有多机组总览、单机下钻、图表分析和 HMI 展示能力。": "测试结果表明，系统已覆盖登录、监测、诊断、数据管理、报告、参数配置、问答和知识库等主要功能，并形成从指标采集到维护建议的闭环。",
    "本文系统仍存在以下不足：第一，诊断模型以规则和模拟数据为主，缺少真实故障样本训练和验证；第二，M-IALO-SVR、VMD-CNN 等算法尚未形成完整训练代码和评价结果，更多体现为系统预留和方法说明；第三，健康评分和 RUL 估计采用经验规则，尚未经过长期运行数据校准；第四，系统权限审计、数据安全、工单闭环和生产部署能力仍有待加强。": "系统不足和改进方向见表7.1。当前原型已完成核心演示流程，但真实数据验证、模型训练和生产级部署仍需继续完善。",
}


DELETE_EXACT = {
    "第一，分析齿轮箱典型故障机理和系统业务需求。结合齿轮箱结构特点，梳理齿轮、轴承、润滑、冷却、轴系和基础结构相关故障类型，明确系统需要支持的状态监测、故障诊断、故障码管理、健康评估和运维问答功能。",
    "第二，设计多源数据接入与融合方法。系统一方面通过 DataAcquisitionSystem 模拟 SCADA、CMS 和油液测点，另一方面通过 WindDataRepository 读取风机数据摘要文件，融合自由报表、实时数据、故障统计和故障信息统计，形成风场总览、单机详情和趋势分析数据。",
    "第三，设计并实现齿轮箱故障诊断流程。系统对振动信号提取 RMS、峰值、峭度、峰值因子、脉冲因子和包络峰值，结合油温、振动 RMS 和油液 NAS 等指标计算风险因子，再根据故障规则库输出故障类型、严重程度、置信度、健康评分、RUL、诊断依据和维护建议。",
    "第四，设计并实现 Web 健康管理平台。后端使用 Flask 构建接口，前端使用 HTML、CSS、JavaScript、ECharts 和 Three.js 展示系统功能。系统实现了用户登录注册、风场总览、单机详情、实时监测、故障诊断、故障数据管理、健康报告、参数设置、故障代码库、HMI 页面和 AI 运维问答等模块。",
    "第五，对系统功能进行测试与分析。通过不同故障场景、不同机组切换和不同接口调用，验证系统能否稳定展示运行状态、生成诊断结果、保存故障记录、导出数据和输出报告。",
    "轴承故障包括磨损、点蚀、剥落和高速轴承过温。轴承磨损初期可能表现为振动 RMS 缓慢升高，点蚀和剥落则会导致峭度、峰值因子和包络峰值升高。高速轴承过温可能与润滑不足、游隙异常、冷却效率下降或测点异常有关。",
    "润滑和冷却故障会对齿轮箱整体健康产生连锁影响。润滑油污染或劣化会降低油膜承载能力，加速齿轮和轴承磨损；冷却系统故障会使油温升高，导致油液黏度变化和热变形增加。轴系不对中和箱体基础松动则通常表现为低频振动增强、宽频振动升高和工况变化下振动波动。",
    "系统中振动数据由正弦基频、轴频、啮合频率和随机噪声叠加生成，并在部分场景中加入冲击信号，以模拟轴承或齿面局部损伤。数据采集模块还使用移动平均方式进行滤波，减少随机噪声对波形展示和特征计算的影响。",
    "包络分析用于突出齿轮和轴承局部损伤产生的调制冲击。当齿面点蚀、轴承剥落或断齿出现时，包络峰值与冲击类指标会同步升高，可作为系统规则诊断的重要依据。",
    "健康评分由综合风险值和故障严重程度共同决定。风险越高，健康评分越低；若故障等级为警告或严重，则进一步扣减评分。RUL 根据健康评分、风险值和故障严重程度估计，健康评分高、风险值低时 RUL 较长；严重故障会显著缩短 RUL。虽然该方法属于规则型估计，但能够在毕业设计阶段形成可解释的健康管理流程。",
    "从系统实现角度看，后端、数据库和前端并不是孤立模块，而是共同支撑齿轮箱健康管理流程。后端负责把原始状态数据转化为结构化诊断结果，数据库负责保存用户、故障记录、配置参数和报告历史，前端则将复杂指标转化为运维人员能够快速理解的页面信息。三者结合后，系统才能同时具备实时监测、历史追溯和辅助决策能力。",
    "在毕业设计实现中，技术选型更强调可实现性和可演示性。Flask 便于快速组织接口，SQLite 适合轻量化部署，ECharts 能够直观展示趋势和分布，SSE 可以降低实时刷新实现复杂度。该组合虽然不等同于生产级风电场平台，但能够覆盖原型系统的核心功能，并为后续接入真实 SCADA、CMS 和油液监测数据预留扩展空间。",
    "Web系统开发技术的作用还体现在系统集成能力上。齿轮箱健康管理涉及实时数据、故障规则、历史记录、报告输出和用户交互等多个环节，如果缺少统一的平台载体，诊断结果难以被运维人员持续使用。通过 Web 系统将各类功能组织在同一界面中，可以提高系统演示的完整性，也便于后续扩展移动端访问、权限控制和远程运维功能。",
    "综上，本文后续系统设计将围绕数据流、业务流和界面流展开。数据流解决齿轮箱状态如何采集、处理和保存的问题，业务流解决故障如何诊断、记录和闭环的问题，界面流解决运维人员如何查看、理解和处理结果的问题。三类流程相互配合，构成系统总体设计和详细实现的基础。",
    "随后系统提取振动特征并计算风险因子。综合风险值输入故障分类规则，得到故障类型、关联部件、严重程度和诊断依据。系统进一步计算健康评分和 RUL，生成维护动作清单。若诊断结果为警告或严重，系统模拟生成检修工单信息，包括工单编号、优先级、责任班组、建议处置窗口和闭环流程。最后，诊断记录写入数据库，供故障数据管理页面查看和导出。",
    "在交互流程上，系统避免让用户在多个页面之间反复寻找同一机组信息。风场总览用于发现异常，单机详情用于查看设备指标，故障诊断用于形成结论，健康报告用于沉淀结果，HMI 页面用于表达现场监控对象。这样的页面组织方式可以让运维人员按照实际工作路径逐步深入，而不是在分散功能中被动查找数据。",
    "在前端人机交互反馈方面，系统将报警状态、健康评分、RUL、故障类型和建议措施放在同一业务链路中呈现。用户既能看到实时指标，也能看到系统给出判断的依据和后续处理建议。对于齿轮箱这类高价值设备，仅显示单一报警数值并不足够，必须同时提供故障解释、处理优先级和闭环入口，才能体现智能健康管理系统的实际意义。",
    "因此，本文前端设计并非单纯追求页面美观，而是围绕运维人员的判断效率展开。总览页面强调快速发现问题，详情页面强调指标解释，诊断页面强调结论可信度，HMI 页面强调现场设备对象感，报告页面强调结果沉淀。不同页面承担不同任务，能够减少信息堆叠造成的阅读负担。",
    "第4章的总体设计为第5章实现提供了结构依据。后续实现章节将按照数据采集、故障诊断、风场与单机页面、数据管理、问答辅助和健康报告等模块展开，分别说明关键代码、系统截图和运行效果，使系统设计与实际开发结果能够相互对应。",
    "风场总览页面是运维人员进入系统后最先接触的页面，其设计目标是快速定位异常机组。页面中不同颜色代表不同运行状态，绿色表示正常运行，橙色和红色表示告警或故障，紫色表示维护停机。每个机组卡片展示功率、风速、油温、健康评分和运行状态，便于用户在一个页面中完成多机组横向比较。",
    "系统还提供指标列表和健康矩阵两种浏览方式。健康矩阵适合快速扫描，指标列表适合查看详细数值。运维人员可以先通过矩阵发现异常机组，再进入单机详情页面查看趋势、故障代码和检修建议。该设计符合风电场运维中“先总览、再下钻、后处理”的工作习惯。",
    "诊断结果并不只显示故障名称，而是同时给出风险等级、部件位置、置信度、健康评分和剩余寿命。下方的诊断依据和建议措施用于解释为什么系统给出该结论，以及下一步应该如何排查。这样的展示方式能够增强诊断结果可信度，避免用户只看到一个缺少依据的告警标签。",
    "从健康管理角度看，单机详情页面能够把不同来源的数据组织到同一个设备视图中。油温和轴承温度用于判断热状态，振动RMS用于判断机械冲击，油液NAS用于判断润滑污染，RUL和复检周期用于形成维护计划。这样的页面结构有助于运维人员快速理解设备状态。",
    "通过风场总览、单机详情和HMI页面的组合，系统形成了三层展示结构：风场总览用于宏观监控，单机详情用于设备分析，HMI页面用于现场交互展示。这种结构较符合风电场运维从场站到机组再到部件的业务层级。",
    "报告结论部分会根据健康评分和风险因子生成文字说明，指出当前机组的主要风险和维护建议。如果系统发现油温、振动或油液指标异常，报告会提示缩短复检周期或安排现场检查。该功能体现了智能健康管理体系从监测数据到管理决策的转化能力。",
    "在实现上，系统默认使用大模型主答、本地知识库作为证据补充；当大模型接口不可用时，后端会回退到本地专家规则，保证演示稳定性。回答渲染函数会把“结论、风险等级、关键依据、建议步骤、后续追问”等标题转换为卡片化结构，并提供查看实时趋势、故障记录和健康报告的跳转按钮，增强运维闭环能力。",
    "参数配置不仅影响页面显示，也会影响诊断结果、告警等级和报告建议。这样可以保证系统的业务逻辑与管理策略一致。用户管理和角色权限则用于控制不同人员的操作范围，避免普通用户误改关键阈值。",
    "从运行效果看，系统界面能够支持从风场总览到单机详情的下钻分析，故障诊断结果具有一定可解释性。系统不足在于当前数据以模拟数据和统计文件为主，后续仍需接入真实SCADA、CMS和油液监测数据，并利用更多真实故障样本对模型进行训练和校验。",
    "测试结果表明，系统能够完成风场总览、单机详情、实时监测、故障诊断、故障数据管理、健康报告和AI运维问答等主要功能。对于油温升高、振动增强、油液污染和轴承冲击等典型场景，系统能够给出风险等级、故障类型、健康评分、RUL估计和维护建议。",
}


TABLES = {
    "1.3 本文研究内容": (
        "表1.1 本文研究内容与系统产出",
        ["研究内容", "主要工作", "系统产出"],
        [
            ["故障机理分析", "梳理齿轮、轴承、润滑、冷却等异常", "故障类型与监测指标对应关系"],
            ["数据接入融合", "整合模拟采集、风机摘要和历史故障", "统一机组快照与趋势数据"],
            ["诊断规则设计", "计算振动特征、风险因子和健康评分", "诊断结果、RUL 和维护建议"],
            ["Web 平台实现", "完成后端接口、数据库和前端页面", "总览、详情、诊断、报告、问答模块"],
            ["测试与评价", "验证页面、接口和典型故障场景", "功能测试表和场景测试结论"],
        ],
    ),
    "2.3 非功能需求分析": (
        "表2.2 系统非功能需求说明",
        ["需求类别", "设计要求", "本文实现方式"],
        [
            ["可用性", "关键状态一屏可见，异常颜色明确", "风场总览、单机详情、HMI 页面"],
            ["可扩展性", "后续可接入真实数据源和模型", "数据仓库、服务层和接口层分离"],
            ["可解释性", "诊断结果需给出依据和建议", "风险因子、特征含义、建议措施同步输出"],
            ["轻量部署", "适合本地演示和答辩运行", "Flask + SQLite + 前端静态资源"],
        ],
    ),
    "2.4 数据需求分析": (
        "表2.3 系统数据来源与用途",
        ["数据来源", "核心字段", "用途"],
        [
            ["模拟采集数据", "油温、振动 RMS、油液 NAS、有功功率", "实时监测和诊断演示"],
            ["自由报表", "风机名称、平均功率、平均风速、环境温度", "风场总览统计"],
            ["实时数据", "转速、功率、叶片角度、齿轮箱温度", "单机详情和趋势展示"],
            ["故障统计", "故障天数、故障次数、峰值日期", "健康报告和风险背景"],
            ["故障信息统计", "故障码、出现次数、开始/结束时间", "故障代码库和记录列表"],
        ],
    ),
    "3.1 齿轮箱典型故障机理": (
        "表3.1 齿轮箱典型故障与监测指标",
        ["故障类别", "典型表现", "主要指标"],
        [
            ["齿轮磨损", "啮合能量升高、油液颗粒增加", "RMS、油液 NAS、啮合异常"],
            ["点蚀/剥落", "局部冲击增强，包络峰值上升", "峭度、峰值因子、包络峰值"],
            ["断齿风险", "周期性强冲击，严重时需停机", "峰值、包络峰值、严重告警"],
            ["轴承损伤", "振动能量和冲击类特征升高", "RMS、峭度、轴承温度"],
            ["润滑劣化", "油膜能力下降，磨损加剧", "油液 NAS、水分、油温"],
            ["冷却异常", "油温升高，热状态恶化", "油温、冷却状态、温度趋势"],
        ],
    ),
    "3.3.1 时域特征": (
        "表3.2 振动特征提取指标",
        ["特征", "含义", "诊断作用"],
        [
            ["RMS", "整体振动能量", "识别磨损、松动和运行异常"],
            ["峰值", "瞬时冲击强度", "提示断齿或强冲击"],
            ["峭度", "信号尖峰程度", "识别早期点蚀和剥落"],
            ["峰值因子", "峰值与 RMS 的比值", "增强局部冲击敏感性"],
            ["脉冲因子", "峰值与平均绝对值的比值", "反映冲击突出程度"],
            ["包络峰值", "调制冲击包络幅值", "辅助识别齿轮/轴承局部损伤"],
        ],
    ),
    "4.4 诊断流程设计": (
        "表4.5 诊断流程关键输入输出",
        ["阶段", "输入", "输出"],
        [
            ["数据读取", "机组编号、场景、实时状态", "机组状态快照"],
            ["特征提取", "振动波形", "RMS、峭度、峰值因子、包络峰值"],
            ["风险评分", "特征、油温、油液等级", "综合风险值和风险因子"],
            ["故障分类", "风险因子、规则库", "故障类型、严重程度、置信度"],
            ["结果输出", "诊断结论和历史记录", "健康评分、RUL、报告和工单建议"],
        ],
    ),
    "5.1 数据采集与融合模块实现": (
        "表5.1 数据采集模块输出字段",
        ["字段类别", "字段示例", "页面用途"],
        [
            ["运行状态", "功率、风速、转速、状态", "风场总览和单机详情"],
            ["齿轮箱状态", "油温、轴承温度、振动 RMS", "实时监测和异常判断"],
            ["油液状态", "NAS 等级、污染风险", "润滑风险评估"],
            ["数据质量", "采样率、延迟、丢包率", "判断采集链路可靠性"],
            ["健康结果", "健康评分、RUL、复检周期", "报告和维护建议"],
        ],
    ),
    "5.2.1 振动特征提取实现": (
        "表5.2 故障诊断模块实现要点",
        ["函数/模块", "处理内容", "输出结果"],
        [
            ["process_signals", "计算 RMS、峰值、峭度等特征", "振动特征字典"],
            ["score_risk", "归一化油温、振动、油液等指标", "综合风险和分项风险"],
            ["classify_fault", "根据规则库识别故障类型", "故障类别、部件、严重程度"],
            ["calculate_health_score", "按风险和严重度扣减评分", "健康评分"],
            ["predict_rul", "估计剩余寿命", "RUL 和复检建议"],
        ],
    ),
    "6.3 诊断场景测试分析": (
        "表6.2 典型诊断场景测试结果",
        ["测试场景", "输入变化", "系统输出"],
        [
            ["正常运行", "油温、振动和油液均处于低风险", "正常状态，建议日常巡检"],
            ["齿轮磨损", "振动能量和油液污染轻度升高", "齿面磨损，建议 72 小时内复测"],
            ["点蚀/剥落", "峭度、峰值因子、包络峰值升高", "严重风险，建议降载复核"],
            ["油温过高", "齿轮箱油温风险升高", "检查冷却风扇、散热器和油位"],
            ["油液污染", "NAS 等级升高", "取油样复测并检查滤芯和密封"],
            ["基础松动", "低频振动增强", "复核地脚螺栓和基础连接"],
        ],
    ),
    "7.2 存在不足": (
        "表7.1 系统不足与改进方向",
        ["不足", "影响", "改进方向"],
        [
            ["真实样本不足", "模型评价说服力有限", "接入 SCADA、CMS 和油液检测历史数据"],
            ["规则模型为主", "泛化能力需要验证", "训练并对比 SVR、XGBoost、LSTM 等模型"],
            ["RUL 经验估计", "寿命预测精度有限", "用长期运行数据校准参数"],
            ["生产部署能力不足", "安全审计和并发能力未验证", "补充权限审计、日志和压力测试"],
        ],
    ),
}


FIGURES = {
    "本文技术路线见图1.1": ("图1.1 本文技术路线图", "fig_1_1_tech_route.png"),
    "系统采用 B/S 架构": ("图4.1 系统数据流与模块关系图", "fig_4_1_data_flow.png"),
    "健康评估过程如图3.1所示": ("图3.1 健康评分与 RUL 估计示意图", "fig_3_1_risk_model.png"),
    "用户完成 HMI 权限登录后进入运行总览页面": ("图5.6 HMI运行总览页面结构示意图", "fig_5_6_hmi_overview.png"),
    "在运维闭环中，系统首先根据阈值和风险因子判断异常": ("图5.11 故障诊断与运维闭环图", "fig_5_11_closed_loop.png"),
}


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
        paragraph._p = paragraph._element = None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side in ("top", "left", "bottom", "right"):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def insert_table_after(paragraph, title, headers, rows):
    caption = OxmlElement("w:p")
    paragraph._p.addnext(caption)
    cap_p = paragraph._parent.add_paragraph()
    cap_p._p = cap_p._element = caption
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(title)
    cap_run.font.size = Pt(10)
    cap_run.font.name = "宋体"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    table = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.0))
    paragraph._p.addnext(table._tbl)
    caption.addnext(table._tbl)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            set_cell_text(cells[idx], text)
    style_table(table)

    spacer = paragraph._parent.add_paragraph()
    table._tbl.addnext(spacer._p)
    return table


def add_picture_after(paragraph, caption, img_path):
    pic_p = paragraph._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(img_path), width=Inches(5.9))
    paragraph._p.addnext(pic_p._p)

    cap_p = paragraph._parent.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.font.size = Pt(10)
    cap_run.font.name = "宋体"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pic_p._p.addnext(cap_p._p)


def find_font():
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    return next((p for p in candidates if Path(p).exists()), None)


def draw_diagram(path, title, columns, rows):
    from PIL import Image, ImageDraw, ImageFont

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = find_font()
    title_font = ImageFont.truetype(font_path, 36) if font_path else ImageFont.load_default()
    label_font = ImageFont.truetype(font_path, 24) if font_path else ImageFont.load_default()
    small_font = ImageFont.truetype(font_path, 20) if font_path else ImageFont.load_default()
    img = Image.new("RGB", (1600, 820), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1599, 819], outline="#CBD5E1", width=3)
    d.text((60, 34), title, fill="#0F172A", font=title_font)
    d.line([60, 92, 1540, 92], fill="#2563EB", width=4)
    x0, y0 = 70, 150
    box_w, box_h = 245, 126
    gap_x, gap_y = 55, 70
    palette = ["#E0F2FE", "#DCFCE7", "#FEF3C7", "#FCE7F3", "#EDE9FE"]
    for r, row in enumerate(rows):
        y = y0 + r * (box_h + gap_y)
        for c, text in enumerate(row):
            if not text:
                continue
            x = x0 + c * (box_w + gap_x)
            fill = palette[(r + c) % len(palette)]
            d.rounded_rectangle([x, y, x + box_w, y + box_h], radius=18, fill=fill, outline="#334155", width=2)
            lines = []
            current = ""
            for ch in text:
                current += ch
                if len(current) >= 9:
                    lines.append(current)
                    current = ""
            if current:
                lines.append(current)
            total_h = len(lines) * 28
            ty = y + (box_h - total_h) / 2
            for line in lines:
                bbox = d.textbbox((0, 0), line, font=label_font)
                d.text((x + (box_w - (bbox[2] - bbox[0])) / 2, ty), line, fill="#0F172A", font=label_font)
                ty += 28
            if c < len(row) - 1 and row[c + 1]:
                ax1, ay = x + box_w + 6, y + box_h // 2
                ax2 = x + box_w + gap_x - 8
                d.line([ax1, ay, ax2, ay], fill="#2563EB", width=4)
                d.polygon([(ax2, ay), (ax2 - 16, ay - 9), (ax2 - 16, ay + 9)], fill="#2563EB")
    for i, col in enumerate(columns):
        x = x0 + i * (box_w + gap_x) + box_w / 2
        bbox = d.textbbox((0, 0), col, font=small_font)
        d.text((x - (bbox[2] - bbox[0]) / 2, 710), col, fill="#475569", font=small_font)
    img.save(path)


def make_figures():
    draw_diagram(
        ASSET_DIR / "fig_1_1_tech_route.png",
        "华锐 SL1500 齿轮箱健康管理技术路线",
        ["需求", "数据", "算法", "平台", "验证"],
        [
            ["故障机理\n需求分析", "数据接入\n预处理", "特征提取\n风险评分", "Web 展示\n报告输出", "场景测试\n结果评价"],
            ["齿轮/轴承\n润滑/冷却", "SCADA/CMS\n油液/故障", "RMS/峭度\nNAS/油温", "总览/详情\n诊断/问答", "功能测试\n闭环检查"],
        ],
    )
    draw_diagram(
        ASSET_DIR / "fig_4_1_data_flow.png",
        "系统数据流与模块关系",
        ["数据源", "后端服务", "数据库", "前端页面", "业务输出"],
        [
            ["模拟采集\n风机摘要", "Flask 路由\n服务层", "SQLite\n记录配置", "ECharts\nThree.js", "诊断结论\n健康报告"],
            ["油温/振动\n油液/功率", "数据融合\n风险诊断", "用户/故障\n资产/参数", "总览/详情\nHMI/问答", "工单建议\nCSV 导出"],
        ],
    )
    draw_diagram(
        ASSET_DIR / "fig_3_1_risk_model.png",
        "健康评分与 RUL 估计示意",
        ["指标输入", "风险归一", "综合判断", "健康输出", "维护建议"],
        [
            ["RMS\n峭度\n油温", "0-1 风险\n权重融合", "规则库\n严重度", "健康评分\nRUL", "复检周期\n处理动作"],
            ["包络峰值\n油液 NAS", "分项风险\n主因排序", "故障类型\n置信度", "诊断依据\n报告摘要", "工单闭环\n复测验收"],
        ],
    )
    draw_diagram(
        ASSET_DIR / "fig_5_6_hmi_overview.png",
        "HMI 运行总览页面结构",
        ["入口", "状态", "模型", "配置", "告警"],
        [
            ["权限登录\n角色确认", "油温/振动\n健康评分", "透明模型\n测点信号", "阈值设置\n参数保存", "临界报警\n故障代码"],
            ["机组选择\n控制台", "RUL\n复检周期", "采集链路\n数据质量", "工程师\n管理员", "排查建议\n闭环入口"],
        ],
    )
    draw_diagram(
        ASSET_DIR / "fig_5_11_closed_loop.png",
        "故障诊断与运维闭环",
        ["发现", "判断", "处理", "记录", "复盘"],
        [
            ["实时监测\n阈值告警", "风险评分\n故障分类", "维护建议\n工单生成", "故障记录\n报告归档", "复测结果\n参数优化"],
            ["总览定位\n单机下钻", "特征依据\n置信度", "责任班组\n处理窗口", "状态更新\nCSV 导出", "知识库补充\n模型迭代"],
        ],
    )


def set_paragraph_text(paragraph, text):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def replace_abstracts(doc):
    paras = doc.paragraphs
    cn_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "摘  要")
    kw_idx = next(i for i, p in enumerate(paras) if p.text.strip().startswith("关键词"))
    body = [p for p in paras[cn_idx + 1 : kw_idx]]
    for p in body[len(ABSTRACT_CN) :]:
        remove_paragraph(p)
    for p, text in zip(body, ABSTRACT_CN):
        set_paragraph_text(p, text)

    en_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "ABSTRACT")
    key_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Key words"))
    body = [p for p in doc.paragraphs[en_idx + 1 : key_idx]]
    for p in body[len(ABSTRACT_EN) :]:
        remove_paragraph(p)
    for p, text in zip(body, ABSTRACT_EN):
        set_paragraph_text(p, text)


def delete_code_blocks(doc):
    deleting = False
    keep_starts = (
        "表",
        "图",
        "4.5.2",
        "5.2",
        "5.3",
        "5.4",
        "5.6",
        "6.",
    )
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text == "核心代码如下：":
            remove_paragraph(p)
            deleting = True
            continue
        if deleting and (text.startswith(keep_starts) or re.match(r"^\d+\.\d+", text)):
            deleting = False
        if deleting:
            remove_paragraph(p)


def apply_text_edits(doc):
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text in REPLACEMENTS:
            set_paragraph_text(p, REPLACEMENTS[text])
        elif text in DELETE_EXACT:
            remove_paragraph(p)


def insert_visuals(doc):
    inserted_tables = set()
    inserted_figures = set()
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text in TABLES and text not in inserted_tables:
            insert_table_after(p, *TABLES[text])
            inserted_tables.add(text)
        for marker, (caption, filename) in FIGURES.items():
            if marker in text and marker not in inserted_figures:
                add_picture_after(p, caption, ASSET_DIR / filename)
                inserted_figures.add(marker)


def fix_numbering(doc):
    exact = {
        "表3.1 齿轮箱健康评估主要风险因子": "表3.3 齿轮箱健康评估主要风险因子",
        "图4.1 系统数据流与模块关系图": "图4.2 系统数据流与模块关系图",
        "图4.2 系统数据库ER图": "图4.3 系统数据库ER图",
        "图4.3 齿轮箱智能健康管理流程图": "图4.4 齿轮箱智能健康管理流程图",
        "表4.5 前端页面与交互功能对应": "表4.6 前端页面与交互功能对应",
        "图5.11 故障诊断与运维闭环图": "图5.10 故障诊断与运维闭环图",
        "图5.10 系统参数设置页面": "图5.11 系统参数设置页面",
    }
    inline = {
        "系统数据流见图4.1": "系统数据流见图4.2",
        "系统参数设置页面如图5.10": "系统参数设置页面如图5.11",
        "闭环过程见图5.11": "闭环过程见图5.10",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in exact:
            set_paragraph_text(p, exact[text])
            continue
        next_text = text
        for old, new in inline.items():
            next_text = next_text.replace(old, new)
        if next_text != text:
            set_paragraph_text(p, next_text)


def request_field_update_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)
    make_figures()
    doc = Document(OUT)
    replace_abstracts(doc)
    delete_code_blocks(doc)
    apply_text_edits(doc)
    insert_visuals(doc)
    fix_numbering(doc)
    request_field_update_on_open(doc)
    doc.save(OUT)
    edited = Document(OUT)
    text_chars = sum(len(p.text.strip()) for p in edited.paragraphs if p.text.strip())
    print(f"saved={OUT}")
    print(f"paragraphs={sum(1 for p in edited.paragraphs if p.text.strip())}")
    print(f"tables={len(edited.tables)}")
    print(f"inline_shapes={len(edited.inline_shapes)}")
    print(f"chars={text_chars}")


if __name__ == "__main__":
    main()
