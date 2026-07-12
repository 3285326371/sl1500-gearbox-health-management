from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_公式美化版.docx")
OUT = SRC.with_name("毕业设计正文模板_图位置优化版.docx")


ORDER = [
    ("第5章 系统详细实现", [
        "图5.12 系统整体分层架构图",
        "图5.17 数据库实体关系图",
        "图5.18 后端接口与蓝图结构图",
        "图5.19 前端页面模块结构图",
    ]),
    ("数据采集模块位于 services/data_acquisition.py", [
        "图5.14 多源数据采集与融合流程图",
    ]),
    ("故障诊断模块位于 services/ml_models.py", [
        "图5.15 齿轮箱故障诊断流程图",
        "图5.2 故障诊断页面",
    ]),
    ("风险评分把振动、油温、油液等级和故障严重度换算成统一风险值", [
        "图5.16 风险评分与健康评估模型图",
    ]),
    ("风场总览接口位于 routes/windfarm_route.py", [
        "图5.13 风场总览与单机下钻流程图",
        "图5.1 风场健康总览页面",
    ]),
    ("单机齿轮箱详情页面如图5.4所示", [
        "图5.4 单机齿轮箱详情页面",
    ]),
    ("齿轮箱HMI权限登录页面如图5.5所示", [
        "图5.5 齿轮箱HMI权限登录页面",
    ]),
    ("用户完成 HMI 权限登录后进入运行总览页面", [
        "图5.6 HMI运行总览页面结构示意图",
    ]),
    ("HMI阈值设置页面如图5.7所示", [
        "图5.7 HMI阈值设置页面",
        "图5.20 HMI运维控制台结构图",
    ]),
    ("故障代码库位于 routes/data_route.py", [
        "图5.21 故障代码库与记录管理关系图",
    ]),
    ("故障数据管理通过 FaultRecord 表和风机文件故障记录共同实现", [
        "图5.3 故障数据管理页面",
    ]),
    ("智能运维诊断助手涉及 frontend/index.html", [
        "图5.23 AI运维问答处理流程图",
        "图5.9 智能运维诊断助手页面",
    ]),
    ("健康报告接口位于 routes/report_route.py", [
        "图5.22 健康报告生成流程图",
        "图5.8 健康报告页面",
    ]),
    ("系统参数设置页面如图5.11所示", [
        "图5.24 参数设置影响链路图",
        "图5.11 系统参数设置页面",
    ]),
    ("在运维闭环中，系统首先根据阈值和风险因子判断异常", [
        "图5.10 故障诊断与运维闭环图",
        "图5.25 运维工单闭环流程图",
    ]),
]


def p_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def set_p_text(el, text):
    texts = list(el.iter(qn("w:t")))
    if not texts:
        return
    texts[0].text = text
    for t in texts[1:]:
        t.text = ""


def find_para(doc, startswith: str):
    for p in doc.paragraphs:
        text = p.text.strip()
        if "\t" in text:
            continue
        if text.startswith(startswith):
            return p
    raise ValueError(f"Anchor not found: {startswith}")


def extract_pair(doc, caption: str):
    body = doc._body._element
    target = next((p for p in doc.paragraphs if p.text.strip() == caption), None)
    if target is None:
        raise ValueError(f"Caption not found: {caption}")
    children = list(body)
    pos = children.index(target._p)
    if pos <= 0 or not children[pos - 1].xpath(".//w:drawing"):
        raise ValueError(f"Image paragraph not found before: {caption}")
    img_el = copy.deepcopy(children[pos - 1])
    cap_el = copy.deepcopy(target._p)
    body.remove(children[pos - 1])
    body.remove(target._p)
    title = re.sub(r"^图5\.\d+\s*", "", caption)
    return {"img": img_el, "cap": cap_el, "title": title, "old": caption}


def insert_after(cursor_el, pair):
    cursor_el.addnext(pair["img"])
    pair["img"].addnext(pair["cap"])
    return pair["cap"]


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    all_captions = [cap for _, caps in ORDER for cap in caps]
    pairs = {cap: extract_pair(doc, cap) for cap in all_captions}

    new_labels = {}
    number = 1
    for _anchor, caps in ORDER:
        for cap in caps:
            pair = pairs[cap]
            new_caption = f"图5.{number} {pair['title']}"
            set_p_text(pair["cap"], new_caption)
            new_labels[pair["title"]] = f"图5.{number}"
            number += 1

    for anchor_text, caps in ORDER:
        anchor = find_para(doc, anchor_text)
        cursor = anchor._p
        for cap in caps:
            cursor = insert_after(cursor, pairs[cap])

    doc.save(OUT)

    # Reopen and repair prose references that mention old figure numbers.
    doc = Document(OUT)
    reference_patterns = {
        "风场总览页面": "风场健康总览页面",
        "故障诊断页面": "故障诊断页面",
        "故障数据管理页面": "故障数据管理页面",
        "单机齿轮箱详情页面": "单机齿轮箱详情页面",
        "齿轮箱HMI权限登录页面": "齿轮箱HMI权限登录页面",
        "运行总览页面": "HMI运行总览页面结构示意图",
        "HMI阈值设置页面": "HMI阈值设置页面",
        "健康报告页面": "健康报告页面",
        "系统参数设置页面": "系统参数设置页面",
    }
    changed_refs = 0
    for p in doc.paragraphs:
        text = p.text
        next_text = text
        for phrase, title in reference_patterns.items():
            if phrase in next_text and re.search(r"图5\.\d+", next_text):
                next_text = re.sub(r"图5\.\d+", new_labels[title], next_text, count=1)
        if next_text != text:
            p.text = next_text
            changed_refs += 1

    doc.save(OUT)
    checked = Document(OUT)
    captions = [p.text.strip() for p in checked.paragraphs if re.match(r"^图5\.\d+\s", p.text.strip())]
    print(f"saved={OUT}")
    print(f"moved_figures={len(all_captions)}")
    print(f"renumbered_figures={len(captions)}")
    print(f"changed_refs={changed_refs}")
    print(f"inline_shapes={len(checked.inline_shapes)}")


if __name__ == "__main__":
    main()
