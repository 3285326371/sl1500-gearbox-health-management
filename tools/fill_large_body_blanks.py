from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path.home() / "Desktop" / "\u6bd5\u4e1a\u8bbe\u8ba1\u6b63\u6587\u6a21\u677f.docx"


def set_body_style(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def paragraph_after(paragraph, text: str):
    p = paragraph._parent.add_paragraph(text)
    paragraph._p.addnext(p._p)
    set_body_style(p)
    return p


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"Paragraph not found: {text}")


def insert_once(doc: Document, anchor_text: str, marker: str, paragraphs: list[str]):
    if any(marker in p.text for p in doc.paragraphs):
        return
    anchor = find_para(doc, anchor_text)
    cur = anchor
    for text in paragraphs:
        cur = paragraph_after(cur, text)


def main():
    backup = DOCX.with_name(f"{DOCX.stem}_减少正文留白前备份_{datetime.now():%Y%m%d_%H%M%S}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)
    doc = Document(DOCX)

    # Correct a figure reference while touching the nearby paragraph.
    for p in doc.paragraphs:
        if "系统功能架构如图4.2所示" in p.text:
            p.text = p.text.replace("系统功能架构如图4.2所示", "系统功能架构如图4.1所示")
            set_body_style(p)

    insert_once(
        doc,
        "系统采用 Flask 框架构建后端服务。Flask 具有轻量、灵活、易扩展的特点，适合毕业设计原型开发。数据库使用 Flask-SQLAlchemy 管理，表结构包括用户表、故障记录表、机组资产表和系统配置表。前端使用 HTML、CSS 和 JavaScript 实现，ECharts 用于趋势图、饼图和柱状图展示，Three.js 用于齿轮箱三维结构示意。系统通过 Server-Sent Events 实现实时状态推送，使页面能够持续刷新油温、振动和波形数据。",
        "从系统实现角度看，后端、数据库和前端并不是孤立模块",
        [
            "从系统实现角度看，后端、数据库和前端并不是孤立模块，而是共同支撑齿轮箱健康管理流程。后端负责把原始状态数据转化为结构化诊断结果，数据库负责保存用户、故障记录、配置参数和报告历史，前端则将复杂指标转化为运维人员能够快速理解的页面信息。三者结合后，系统才能同时具备实时监测、历史追溯和辅助决策能力。",
            "在毕业设计实现中，技术选型更强调可实现性和可演示性。Flask 便于快速组织接口，SQLite 适合轻量化部署，ECharts 能够直观展示趋势和分布，SSE 可以降低实时刷新实现复杂度。该组合虽然不等同于生产级风电场平台，但能够覆盖原型系统的核心功能，并为后续接入真实 SCADA、CMS 和油液监测数据预留扩展空间。",
        ],
    )

    insert_once(
        doc,
        "HMI 页面更强调现场设备交互。该页面围绕单台齿轮箱展开，将齿轮箱油温、高速轴承温度、振动RMS、润滑油等级、健康评分和报警状态集中到同一视图，并预留透明模型、测点信号、阈值设置和临界报警入口。相比普通管理页面，HMI 界面能够体现设备对象、运行状态和操作反馈之间的对应关系，是系统面向现场监控应用的重要补充。",
        "在交互流程上，系统避免让用户在多个页面之间反复寻找同一机组信息",
        [
            "在交互流程上，系统避免让用户在多个页面之间反复寻找同一机组信息。风场总览用于发现异常，单机详情用于查看设备指标，故障诊断用于形成结论，健康报告用于沉淀结果，HMI 页面用于表达现场监控对象。这样的页面组织方式可以让运维人员按照实际工作路径逐步深入，而不是在分散功能中被动查找数据。",
            "在人机交互反馈方面，系统将报警状态、健康评分、RUL、故障类型和建议措施放在同一业务链路中呈现。用户既能看到实时指标，也能看到系统给出判断的依据和后续处理建议。对于齿轮箱这类高价值设备，仅显示单一报警数值并不足够，必须同时提供故障解释、处理优先级和闭环入口，才能体现智能健康管理系统的实际意义。",
        ],
    )

    insert_once(
        doc,
        "系统参数设置页面如图5.6所示。该页面用于配置油温、振动、油液阈值和系统访问参数，阈值变化会影响实时告警和健康报告。",
        "参数设置页面与健康报告、故障诊断和工单闭环存在直接关联",
        [
            "参数设置页面与健康报告、故障诊断和工单闭环存在直接关联。油温阈值、振动阈值和油液阈值改变后，系统在实时监测页面中的状态颜色、诊断模块中的风险因子、报告模块中的维护建议都会随之变化。因此，参数设置并不是简单的后台配置，而是健康管理体系中连接管理策略和诊断结果的关键环节。",
            "在运维闭环中，系统首先根据阈值和风险因子判断异常，再生成诊断结果和维护建议；当风险等级较高时，系统进一步模拟生成工单信息，并把处理状态写入故障记录。这样可以把“发现异常、解释原因、安排处理、复测确认”的过程串联起来，使论文中的系统实现不只停留在页面展示，而是具备完整业务流程。",
        ],
    )

    doc.save(DOCX)
    print(f"saved={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
