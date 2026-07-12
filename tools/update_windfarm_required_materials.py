from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = DOCS / "风电场数据对接与项目阶段成果说明_信息安全修订版.docx"
OUT = DOCS / "风电场数据对接与项目阶段成果说明_资料清单优化版.docx"

NEW_ITEMS = [
    "机组基础信息：包括机组编号、机组型号、额定容量、投运日期、所属线路或区域、当前运行状态等，用于建立系统中的机组资产档案和机组编号映射关系。",
    "数据接口与信息安全资料：包括 SCADA、PLC、CMS 等系统的数据访问方式、通信协议类型、点表信息或字段字典等。涉及 IP 地址、端口、账号权限、访问密钥等敏感信息时，应按风电场网络安全要求进行脱敏处理，或仅在现场受控环境下查看。系统原则上以只读方式接入运行数据，不直接修改 PLC 控制逻辑，不越权访问生产控制网络，并对数据访问过程进行权限控制和操作记录，确保数据接入安全、可控、可追溯。",
    "历史运行数据：建议提供至少 3 至 12 个月的历史运行数据，主要包括齿轮箱油温、环境温度、风速、有功功率、发电机转速、变桨角度、偏航角度、振动特征、报警记录和停机事件等，用于模型训练、趋势分析和健康评估结果校准。",
    "故障与检修记录：包括齿轮箱、轴承、冷却系统、润滑系统等相关故障案例、报警原因、停机原因、处理措施、备件更换情况和复测结果，用于建立故障样本标签，并验证系统诊断结论的准确性。",
    "现场运行与验收规则：包括风电场网络安全要求、数据脱敏要求、接口访问审批流程、控制命令权限边界、数据使用范围、现场联调方式和验收流程等，用于保证系统对接过程符合现场生产管理和信息安全要求。",
    "数据使用授权说明：明确数据使用范围、使用周期、保密要求和成果展示边界。论文、答辩或公开展示中不直接暴露风电场敏感信息，必要时对机组编号、场站名称、网络地址和账号信息进行脱敏处理。",
    "通过上述资料，可完成机组资产建档、真实数据接入、字段映射、模型训练验证、故障样本标注和健康评估结果校准，为系统由演示原型向工程应用转化提供数据基础。",
]


def main():
    doc = Document(SOURCE)
    paragraphs = doc.paragraphs
    start = None
    end = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "九、需要风电场配合提供的资料":
            start = i
        elif start is not None and p.text.strip().startswith("十、"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("未找到第九节或第十节边界")

    # Remove old paragraphs between title and next section.
    for p in paragraphs[start + 1:end]:
        p._element.getparent().remove(p._element)

    anchor = paragraphs[start]._element
    body = anchor.getparent()
    insert_at = list(body).index(anchor) + 1
    for item in NEW_ITEMS:
        p = doc.add_paragraph(item)
        body.remove(p._element)
        body.insert(insert_at, p._element)
        insert_at += 1

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
