from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path.cwd()
TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_加流程图前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
IMG = ROOT / "docs" / "figures" / "gearbox_health_management_flowchart.png"


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def draw_wrapped(draw, xy, text, box_w, fnt, fill):
    lines = []
    current = ""
    for ch in text:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= box_w:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    x, y = xy
    line_h = fnt.size + 8
    total_h = len(lines) * line_h
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x + (box_w - (bbox[2] - bbox[0])) / 2, y + i * line_h), line, font=fnt, fill=fill)
    return total_h


def arrow(draw, start, end, color=(40, 40, 40), width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 16 * direction, y2 - 8), (x2 - 16 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - 16 * direction), (x2 + 8, y2 - 16 * direction)]
    draw.polygon(pts, fill=color)


def create_flowchart():
    IMG.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1700, 980
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_f = font(42, True)
    node_f = font(28, True)
    small_f = font(22)
    line = (45, 45, 45)
    fill_main = (235, 244, 255)
    fill_warn = (255, 247, 230)
    fill_done = (235, 249, 240)
    outline = (20, 75, 130)

    title = "齿轮箱智能健康管理流程"
    tb = d.textbbox((0, 0), title, font=title_f)
    d.text(((w - (tb[2] - tb[0])) / 2, 40), title, font=title_f, fill=(0, 0, 0))

    nodes = [
        (90, 170, 360, 130, "多源数据接入", "SCADA、振动、油温、油液、历史故障", fill_main),
        (500, 170, 360, 130, "数据清洗与融合", "异常值处理、时间对齐、机组状态归一", fill_main),
        (910, 170, 360, 130, "特征提取", "RMS、峭度、峰值因子、脉冲因子、包络峰值", fill_main),
        (1310, 170, 300, 130, "风险因子计算", "温度、振动、油质、历史故障综合评估", fill_warn),
        (290, 430, 360, 130, "规则库故障诊断", "齿面磨损、点蚀剥落、轴承过温、油液劣化", fill_warn),
        (760, 430, 360, 130, "健康评分与RUL", "输出健康分、剩余寿命和风险等级", fill_warn),
        (1230, 430, 360, 130, "诊断结果解释", "生成诊断依据、报警原因和处置建议", fill_warn),
        (500, 700, 360, 130, "健康报告生成", "形成单机健康报告与检修建议", fill_done),
        (970, 700, 360, 130, "运维闭环", "故障记录、工单模拟、知识库问答", fill_done),
    ]

    for x, y, bw, bh, head, desc, fill in nodes:
        d.rounded_rectangle((x, y, x + bw, y + bh), radius=22, fill=fill, outline=outline, width=4)
        hb = d.textbbox((0, 0), head, font=node_f)
        d.text((x + (bw - (hb[2] - hb[0])) / 2, y + 22), head, font=node_f, fill=(0, 0, 0))
        draw_wrapped(d, (x + 24, y + 72), desc, bw - 48, small_f, (50, 50, 50))

    arrow(d, (450, 235), (500, 235), line)
    arrow(d, (860, 235), (910, 235), line)
    arrow(d, (1270, 235), (1310, 235), line)
    arrow(d, (1460, 300), (1410, 430), line)
    arrow(d, (1230, 495), (1120, 495), line)
    arrow(d, (760, 495), (650, 495), line)
    arrow(d, (470, 560), (620, 700), line)
    arrow(d, (1120, 560), (1150, 700), line)
    arrow(d, (860, 765), (970, 765), line)

    note = "流程说明：系统先完成状态数据接入与特征提取，再结合规则库和风险因子进行故障诊断，最后输出健康评分、RUL、报告和运维建议。"
    d.rounded_rectangle((150, 880, 1550, 940), radius=18, fill=(248, 248, 248), outline=(180, 180, 180), width=2)
    draw_wrapped(d, (190, 897), note, 1320, small_f, (35, 35, 35))
    im.save(IMG)


def add_to_doc():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)

    # Avoid duplicate insertion if rerun.
    for p in doc.paragraphs:
        if "图4.1 齿轮箱智能健康管理流程图" in p.text:
            doc.save(TEMPLATE)
            return

    idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "4.4 诊断流程设计")
    anchor = doc.paragraphs[idx + 1]
    run = anchor.insert_paragraph_before().add_run()
    run.add_picture(str(IMG), width=Inches(5.7))
    pic_p = anchor._element.getprevious()
    # The inserted paragraph object is not directly returned by python-docx in older builds.
    for p in doc.paragraphs:
        if p._element is pic_p:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    caption = anchor.insert_paragraph_before("图4.1 齿轮箱智能健康管理流程图")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in caption.runs:
        r.font.size = Pt(10.5)

    lead = anchor.insert_paragraph_before(
        "系统诊断流程如图4.1所示。该流程以多源状态数据为输入，经过数据清洗、特征提取、风险因子计算和规则库诊断后，形成健康评分、RUL估计、诊断依据与运维建议，并通过健康报告、故障记录和工单模拟实现闭环管理。"
    )
    for r in lead.runs:
        r.font.size = Pt(12)

    doc.save(TEMPLATE)


def main():
    create_flowchart()
    add_to_doc()
    print(f"docx={TEMPLATE}")
    print(f"backup={BACKUP}")
    print(f"image={IMG}")


if __name__ == "__main__":
    main()
