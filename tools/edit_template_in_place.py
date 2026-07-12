from copy import deepcopy
from datetime import datetime
from pathlib import Path
from shutil import copy2
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TITLE_LINE_1 = "华锐SL1500型双馈式风机齿轮箱"
TITLE_LINE_2 = "智能健康管理系统设计与实现"
FULL_TITLE = f"{TITLE_LINE_1}{TITLE_LINE_2}"


ROOT = Path.cwd()
TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
SOURCE = max((ROOT / "docs").glob("*开题文献版.docx"), key=lambda p: p.stat().st_mtime)
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_修改前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_cover_title(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.25
    r1 = paragraph.add_run(TITLE_LINE_1)
    set_font(r1, east_asia="黑体", ascii_font="Arial", size=22, bold=False)
    r1.add_break()
    r2 = paragraph.add_run(TITLE_LINE_2)
    set_font(r2, east_asia="黑体", ascii_font="Arial", size=22, bold=False)


def replace_text(paragraph, text, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def copy_source_body_into_template():
    copy2(TEMPLATE, BACKUP)

    template_doc = Document(TEMPLATE)
    source_doc = Document(SOURCE)

    # Fill the template cover and front matter in place.
    set_cover_title(template_doc.paragraphs[6])
    replace_text(template_doc.paragraphs[18], "2026 年 05 月", east_asia="宋体", ascii_font="Times New Roman", size=14)
    replace_text(
        template_doc.paragraphs[24],
        f"本人所提交的毕业设计《{FULL_TITLE}》，是在导师的指导下，独立进行研究工作所取得的原创性成果。除文中已经注明引用的内容外，本毕业设计不包含任何其他个人或集体已经发表或撰写过的研究成果。对本文的研究做出重要贡献的个人和集体，均已在文中标明。",
        east_asia="宋体",
        ascii_font="Times New Roman",
        size=12,
    )

    # Remove visible template-only notes on the cover.
    for idx in sorted([19, 20], reverse=True):
        delete_paragraph(template_doc.paragraphs[idx])

    target_start = next(i for i, p in enumerate(template_doc.paragraphs) if p.text.strip() == "摘□□要")
    source_start = next(i for i, p in enumerate(source_doc.paragraphs) if p.text.strip() == "摘  要")

    body = template_doc._body._element
    target_elements = list(body)
    target_start_el = template_doc.paragraphs[target_start]._element
    target_start_index = target_elements.index(target_start_el)

    # Keep the template front matter before the abstract; remove all example content after it.
    for el in target_elements[target_start_index:]:
        body.remove(el)

    # Append the already-written thesis body into the original template document.
    source_body = source_doc._body._element
    source_elements = list(source_body)
    source_start_el = source_doc.paragraphs[source_start]._element
    source_start_index = source_elements.index(source_start_el)

    for el in source_elements[source_start_index:]:
        if el.tag == qn("w:sectPr"):
            continue
        body.append(deepcopy(el))

    # Restore the final section properties expected by Word.
    sect = deepcopy(source_body.sectPr)
    if sect is not None:
        body.append(sect)

    template_doc.save(TEMPLATE)


def strip_instruction_shapes_keep_logo():
    # python-docx preserves drawings, but the template has many instruction callouts.
    # Remove drawings unless they are the school logo picture.
    with ZipFile(TEMPLATE, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    xml = parts["word/document.xml"].decode("utf-8")
    from lxml import etree

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    root = etree.fromstring(xml.encode("utf-8"))
    removed = 0
    for drawing in list(root.xpath(".//w:drawing", namespaces=ns)):
        doc_pr = drawing.xpath(".//wp:docPr", namespaces=ns)
        name = doc_pr[0].get("name", "") if doc_pr else ""
        descr = doc_pr[0].get("descr", "") if doc_pr else ""
        keep = name.startswith("图片") or "校徽" in descr
        if not keep:
            parent = drawing.getparent()
            parent.remove(drawing)
            removed += 1

    parts["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with ZipFile(TEMPLATE, "w", ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    return removed


def main():
    copy_source_body_into_template()
    removed = strip_instruction_shapes_keep_logo()
    print(f"template={TEMPLATE}")
    print(f"source={SOURCE}")
    print(f"backup={BACKUP}")
    print(f"removed_instruction_shapes={removed}")


if __name__ == "__main__":
    main()
