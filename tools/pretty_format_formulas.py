from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_公式增强_截图彩色版.docx")
OUT = SRC.with_name("毕业设计正文模板_公式美化版.docx")


FORMULA_REPLACEMENTS = {
    "e_T = T_gbx - T_hat_gbx                                                       （1-1）": "e_T = T_GBX − T̂_GBX    （1-1）",
    "x_RMS = sqrt((1/N) * sum_{i=1}^{N} x_i^2)                                      （3-1）": "x_RMS = √( 1/N · ∑ᵢ₌₁ᴺ xᵢ² )    （3-1）",
    "K = [(1/N) * sum_{i=1}^{N}(x_i - x_bar)^4] / sigma^4                           （3-2）": "K = [ 1/N · ∑ᵢ₌₁ᴺ (xᵢ − x̄)⁴ ] / σ⁴    （3-2）",
    "C_f = max(|x_i|) / x_RMS                                                       （3-3）": "C_f = max(|xᵢ|) / x_RMS    （3-3）",
    "I_f = max(|x_i|) / [(1/N) * sum_{i=1}^{N}|x_i|]                                 （3-4）": "I_f = max(|xᵢ|) / [ 1/N · ∑ᵢ₌₁ᴺ |xᵢ| ]    （3-4）",
    "r_j = min(1, max(0, (z_j - z_warn) / (z_crit - z_warn)))                       （3-5）": "rⱼ = min{1, max[0, (zⱼ − z_warn)/(z_crit − z_warn)]}    （3-5）",
    "R = sum_{j=1}^{m} w_j r_j,    sum_{j=1}^{m} w_j = 1                              （3-6）": "R = ∑ⱼ₌₁ᵐ wⱼrⱼ，且 ∑ⱼ₌₁ᵐ wⱼ = 1    （3-6）",
    "H = max(0, 100 - 100R - Delta_s)                                               （5-1）": "H = max(0, 100 − 100R − Δ_s)    （5-1）",
    "RUL = RUL_max * (H / 100)^gamma * (1 - R)                                      （5-2）": "RUL = RUL_max × (H/100)^γ × (1 − R)    （5-2）",
}


TEXT_FIXES = {
    "式中，x_i 为第 i 个振动采样点，N 为采样点数，x_bar 为均值，sigma 为标准差，K、C_f 和 I_f 分别表示峭度、峰值因子和脉冲因子。": "式中，xᵢ 为第 i 个振动采样点，N 为采样点数，x̄ 为均值，σ 为标准差，K、C_f 和 I_f 分别表示峭度、峰值因子和脉冲因子。",
    "式中，z_j 为第 j 个监测指标，z_warn 和 z_crit 分别为关注阈值和临界阈值，r_j 为分项风险因子，w_j 为对应权重，R 为综合风险值。": "式中，zⱼ 为第 j 个监测指标，z_warn 和 z_crit 分别为关注阈值和临界阈值，rⱼ 为分项风险因子，wⱼ 为对应权重，R 为综合风险值。",
    "式中，H 为健康评分，Delta_s 为故障严重度修正量，RUL_max 为最大参考寿命，gamma 为寿命衰减系数。该估计方法用于原型演示，后续仍需要结合真实运行数据进行校准。": "式中，H 为健康评分，Δ_s 为故障严重度修正量，RUL_max 为最大参考寿命，γ 为寿命衰减系数。该估计方法用于原型演示，后续仍需要结合真实运行数据进行校准。",
}


def set_paragraph(paragraph, text, formula=False):
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if formula else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman" if formula else "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if not formula else "Times New Roman")
    run.font.size = Pt(11 if formula else 10.5)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    formula_changed = 0
    text_changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in FORMULA_REPLACEMENTS:
            set_paragraph(paragraph, FORMULA_REPLACEMENTS[text], formula=True)
            formula_changed += 1
        elif text in TEXT_FIXES:
            set_paragraph(paragraph, TEXT_FIXES[text], formula=False)
            text_changed += 1
    doc.save(OUT)

    checked = Document(OUT)
    formulas = [p.text.strip() for p in checked.paragraphs if re.search(r"（[135]-\d+）$", p.text.strip())]
    bad = [f for f in formulas if "sum_" in f or "sqrt" in f or "Delta" in f or "gamma" in f or "*" in f]
    print(f"saved={OUT}")
    print(f"formula_changed={formula_changed}")
    print(f"text_changed={text_changed}")
    print(f"formula_count={len(formulas)}")
    print(f"bad_formula_left={len(bad)}")


if __name__ == "__main__":
    main()
