from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path.cwd()
TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_补足40页前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = ROOT / "docs" / "_tmp_毕业设计正文模板_补足40页.docx"
SHOT_DIR = ROOT / "docs" / "system_screenshots"


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def format_para(p, first_line=True):
    p.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        set_font(r)


def find_para(doc, text):
    norm = text.replace(" ", "")
    for p in doc.paragraphs:
        if p.text.strip().replace(" ", "") == norm:
            return p
    raise ValueError(text)


def add_before(anchor, text="", first_line=True):
    p = anchor.insert_paragraph_before(text)
    format_para(p, first_line)
    return p


def add_subheading(anchor, text):
    p = anchor.insert_paragraph_before(text)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=12, bold=True)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "666666")


def add_table(anchor, title, headers, rows):
    cap = anchor.insert_paragraph_before(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    for r in cap.runs:
        set_font(r, east_asia="黑体", ascii_font="Arial", size=10.5, bold=True)
    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Inches(6.0))
    anchor._p.addprevious(table._tbl)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        shade_cell(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9.5)


def add_picture(anchor, path, caption, intro):
    add_before(anchor, intro)
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.85))
    cap = anchor.insert_paragraph_before(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        set_font(r, size=10.5)


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    if any("6.8 单机详情与HMI验证" in p.text for p in doc.paragraphs):
        print("already added")
    else:
        anchor5 = find_para(doc, "第6章 系统测试与结果分析")
        add_picture(anchor5, SHOT_DIR / "07_turbine_detail.png", "图5.7 单机齿轮箱详情页面", "单机齿轮箱详情页面如图5.7所示。页面从风场总览下钻到单台机组，集中展示齿轮箱油温、高速轴承温度、振动RMS、油液NAS等级、健康评分、RUL、复检周期和历史故障记录。")
        for text in [
            "单机详情页面体现了系统从风场级监控到设备级诊断的下钻能力。运维人员在风场总览中发现异常机组后，可以进入该页面查看机组的关键指标和部件状态。该页面还提供HMI控制台、故障诊断、实时趋势和健康报告入口，使单台机组相关功能形成集中工作台。",
            "从健康管理角度看，单机详情页面能够把不同来源的数据组织到同一个设备视图中。油温和轴承温度用于判断热状态，振动RMS用于判断机械冲击，油液NAS用于判断润滑污染，RUL和复检周期用于形成维护计划。这样的页面结构有助于运维人员快速理解设备状态。",
        ]:
            add_before(anchor5, text)
        add_picture(anchor5, SHOT_DIR / "08_hmi_page.png", "图5.8 齿轮箱HMI人机交互页面", "齿轮箱HMI页面如图5.8所示。该页面面向现场监控场景，用更加接近工业界面的方式展示单机齿轮箱状态和关键指标。")
        for text in [
            "HMI页面主要用于展示齿轮箱关键运行信息和人机交互状态。与普通Web仪表盘相比，HMI页面更强调现场监控感和设备对象感，适合在答辩演示时说明系统具备工业监控界面的扩展能力。页面可根据机组编号加载不同设备状态，并与后端接口保持一致的数据来源。",
            "通过风场总览、单机详情和HMI页面的组合，系统形成了三层展示结构：风场总览用于宏观监控，单机详情用于设备分析，HMI页面用于现场交互展示。这种结构较符合风电场运维从场站到机组再到部件的业务层级。",
        ]:
            add_before(anchor5, text)

        anchor7 = find_para(doc, "第7章 总结与展望")
        add_subheading(anchor7, "6.8 单机详情与HMI验证")
        for text in [
            "单机详情与HMI页面测试主要验证系统的下钻分析能力。测试时从风场总览页面选择WTG-001机组，系统能够进入齿轮箱详情页面，并显示齿轮箱油温、高速轴承温度、振动RMS、润滑油NAS等级、健康评分、剩余寿命和历史故障记录等信息。页面按钮能够跳转到HMI、故障诊断、实时趋势和健康报告，说明页面之间的业务关联能够正常运行。",
            "HMI页面测试采用直接访问和详情页跳转两种方式。测试结果表明，HMI页面能够按照指定机组加载显示内容，页面布局稳定，关键指标显示清晰。虽然当前HMI页面仍属于演示型界面，但它能够说明系统具备面向现场监控终端扩展的可能性，为后续接入真实PLC、SCADA网关或工业组态界面提供了展示基础。",
        ]:
            add_before(anchor7, text)
        add_table(anchor7, "表6.4 页面截图验证结果", ["页面", "验证内容", "结果"], [
            ["风场总览", "机组矩阵、统计指标、状态颜色", "显示正常"],
            ["单机详情", "油温、振动、RUL、复检周期", "显示正常"],
            ["故障诊断", "场景选择、诊断结果、建议措施", "显示正常"],
            ["故障记录", "记录列表、筛选、导出入口", "显示正常"],
            ["健康报告", "关键指标、结论、报告结构", "显示正常"],
            ["AI问答", "问题输入、知识回答、上下文展示", "显示正常"],
            ["参数设置", "阈值配置、角色和用户管理", "显示正常"],
            ["HMI页面", "工业监控界面和机组状态", "显示正常"],
        ])
        add_subheading(anchor7, "6.9 非功能测试与可用性分析")
        for text in [
            "非功能测试主要关注系统的稳定性、可用性、可维护性和扩展性。稳定性方面，系统在本地Flask服务环境下能够连续刷新数据并完成多次诊断请求，页面没有出现明显卡死。可用性方面，侧边栏按健康监测区、智能诊断区和系统配置区分组，用户能够按照业务流程寻找功能。可维护性方面，后端使用蓝图拆分路由，服务层负责数据采集、诊断模型、知识库和风机数据仓库，模块边界较清晰。",
            "扩展性方面，系统当前虽然使用SQLite和模拟数据，但数据接口、诊断服务和前端页面之间保持松耦合。后续如果接入真实风机数据，只需要替换数据采集服务或风机数据仓库，不需要大规模修改前端页面。若引入真实机器学习模型，也可以在run_diagnosis流程中替换风险评分和分类逻辑。",
            "从可用性角度看，系统页面采用卡片、表格、图表和报告相结合的方式展示信息。对于运维人员来说，卡片适合快速查看状态，表格适合追踪记录，图表适合观察趋势，报告适合形成结论。多种展示方式共同提高了系统在不同工作场景下的可读性。",
        ]:
            add_before(anchor7, text)
        add_table(anchor7, "表6.5 非功能测试结果", ["测试项", "测试方法", "结果分析"], [
            ["稳定性", "连续刷新页面和多次诊断", "接口响应正常，页面无明显卡顿"],
            ["可用性", "按业务流程切换模块", "导航清晰，核心功能易定位"],
            ["可维护性", "检查后端路由和服务划分", "蓝图和服务层结构较清晰"],
            ["扩展性", "分析真实数据接入改造范围", "可通过替换采集服务扩展"],
            ["可解释性", "查看诊断依据和建议措施", "结果包含特征和风险说明"],
        ])
        add_subheading(anchor7, "6.10 存在问题与改进方向")
        for text in [
            "本系统已经实现齿轮箱健康管理的主要业务链条，但仍存在一些不足。首先，当前诊断数据主要来自模拟生成和统计文件，不能完全代表真实风场复杂工况。真实风机运行会受到风速波动、环境温度、功率限制、并网状态和检修状态等多种因素影响，后续需要接入真实SCADA、CMS和油液检测数据进行验证。",
            "其次，当前故障诊断主要采用规则库和风险评分方法，优点是可解释性强，缺点是对复杂非线性退化过程表达能力有限。后续可以基于真实故障样本训练机器学习模型，例如支持向量机、随机森林、XGBoost、LSTM或Transformer模型，并与规则库形成混合诊断机制。",
            "再次，当前RUL估计仍属于启发式计算，能够展示寿命预测流程，但精度需要真实退化数据支撑。后续可以建立正常工况模型和退化趋势模型，结合剩余寿命标签进行监督学习。同时，应引入模型评价指标，如MAE、RMSE、准确率、召回率和F1值，对模型性能进行量化评价。",
            "最后，系统安全性和部署能力还有提升空间。毕业设计阶段主要在本地环境运行，后续若部署到真实风场，应增加HTTPS、日志审计、权限细分、数据备份和异常恢复机制，并考虑服务容器化部署。这样才能从演示型系统进一步发展为可工程应用的智能健康管理平台。",
        ]:
            add_before(anchor7, text)

    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
