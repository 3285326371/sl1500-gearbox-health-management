from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_正文逐段润色版.docx")
OUT = SRC.with_name("毕业设计正文模板_公式增强版.docx")


INSERTS = {
    "振动特征部分重点选取能量类、冲击类和包络类指标。RMS 用来描述整体振动能量，峭度和峰值因子用于反映冲击特征，包络峰值则辅助判断局部损伤。": [
        "为便于后续诊断计算，本文采用式（3-1）至式（3-4）提取振动时域特征。",
        "x_RMS = sqrt((1/N) * sum_{i=1}^{N} x_i^2)                                      （3-1）",
        "K = [(1/N) * sum_{i=1}^{N}(x_i - x_bar)^4] / sigma^4                           （3-2）",
        "C_f = max(|x_i|) / x_RMS                                                       （3-3）",
        "I_f = max(|x_i|) / [(1/N) * sum_{i=1}^{N}|x_i|]                                 （3-4）",
        "式中，x_i 为第 i 个振动采样点，N 为采样点数，x_bar 为均值，sigma 为标准差，K、C_f 和 I_f 分别表示峭度、峰值因子和脉冲因子。",
    ],
    "健康评估过程如图3.1所示。各类指标先被换算成风险因子，再进一步汇总为综合风险、健康评分和 RUL。": [
        "不同量纲的状态指标先按式（3-5）归一化，再按式（3-6）得到综合风险值。",
        "r_j = min(1, max(0, (z_j - z_warn) / (z_crit - z_warn)))                       （3-5）",
        "R = sum_{j=1}^{m} w_j r_j,    sum_{j=1}^{m} w_j = 1                              （3-6）",
        "式中，z_j 为第 j 个监测指标，z_warn 和 z_crit 分别为关注阈值和临界阈值，r_j 为分项风险因子，w_j 为对应权重，R 为综合风险值。",
    ],
    "风险评分把振动、油温、油液等级和故障严重度换算成统一风险值，再用于计算健康评分和剩余寿命。它还不是基于大量样本训练出的预测模型，但计算过程清楚，适合在毕业设计阶段解释诊断结论的来源。": [
        "在原型系统中，健康评分和剩余寿命按式（5-1）和式（5-2）进行经验估计。",
        "H = max(0, 100 - 100R - Delta_s)                                               （5-1）",
        "RUL = RUL_max * (H / 100)^gamma * (1 - R)                                      （5-2）",
        "式中，H 为健康评分，Delta_s 为故障严重度修正量，RUL_max 为最大参考寿命，gamma 为寿命衰减系数。该估计方法用于原型演示，后续仍需要结合真实运行数据进行校准。",
    ],
    "相比高频振动系统，SCADA 数据在风场中更容易获得，部署成本也更低。它通常记录风速、有功功率、发电机转速、齿轮箱油温、轴承温度、变桨角、偏航角和报警状态等信息。由于采样频率有限，SCADA 很难捕捉齿轮或轴承的高频冲击，但适合观察工况变化和长周期趋势。实际应用中，常通过支持向量回归、随机森林、神经网络等方法建立正常工况模型，再利用实测值与预测值的残差判断异常。": [
        "若后续接入真实 SCADA 数据，可用式（1-1）描述温度残差，用于判断油温是否偏离正常工况。",
        "e_T = T_gbx - T_hat_gbx                                                       （1-1）",
        "式中，T_gbx 为实测齿轮箱油温，T_hat_gbx 为正常工况模型预测油温，e_T 为温度残差。当残差长期为正且超过设定范围时，可作为过温风险的辅助判断依据。",
    ],
}


def set_run_font(run, font_name="宋体", size=10.5):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)


def insert_after(paragraph, text, is_formula=False):
    new_p = paragraph._parent.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_formula else WD_ALIGN_PARAGRAPH.LEFT
    run = new_p.add_run(text)
    set_run_font(run, "Times New Roman" if is_formula else "宋体", 10.5)
    paragraph._p.addnext(new_p._p)
    return new_p


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    inserted = 0
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text not in INSERTS:
            continue
        cursor = paragraph
        for line in INSERTS[text]:
            is_formula = bool(re.search(r"（[135]-\d+）$", line))
            cursor = insert_after(cursor, line, is_formula=is_formula)
            inserted += 1

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUT)
    checked = Document(OUT)
    formula_count = sum(1 for p in checked.paragraphs if re.search(r"（[135]-\d+）", p.text))
    print(f"saved={OUT}")
    print(f"inserted_paragraphs={inserted}")
    print(f"formula_count={formula_count}")
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")


if __name__ == "__main__":
    main()
