from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))


def main():
    doc = Document(TEMPLATE)
    declaration = next(p for p in doc.paragraphs if p.text.strip() == "毕业设计原创性声明")
    previous = declaration._element.getprevious()
    if previous is not None:
        # Add a page break in the paragraph immediately before the declaration,
        # unless it already contains one.
        text = previous.xml
        if 'w:type="page"' not in text:
            p = doc.paragraphs[[p._element for p in doc.paragraphs].index(previous)]
            p.add_run().add_break(WD_BREAK.PAGE)
    doc.save(TEMPLATE)
    print(TEMPLATE)


if __name__ == "__main__":
    main()
