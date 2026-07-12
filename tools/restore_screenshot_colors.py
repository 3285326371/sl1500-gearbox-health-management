from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\bishe\Huaren SL1500 Type Doubly-Fed Wind Turbine Gearbox Intelligent Health Management System1.1")
SRC = Path(r"D:\迅雷下载\毕业设计正文模板_检测报告\毕业设计正文模板_公式增强版.docx")
OUT = SRC.with_name("毕业设计正文模板_公式增强_截图彩色版.docx")


COLOR_SCREENSHOTS = {
    "图5.1 风场健康总览页面": ROOT / "docs/system_screenshots/01_windfarm_overview.png",
    "图5.2 故障诊断页面": ROOT / "docs/system_screenshots/02_fault_diagnosis.png",
    "图5.3 故障数据管理页面": ROOT / "docs/system_screenshots/03_fault_records.png",
    "图5.4 单机齿轮箱详情页面": ROOT / "docs/system_screenshots/07_turbine_detail.png",
    "图5.5 齿轮箱HMI权限登录页面": ROOT / "docs/system_screenshots_updated/05_hmi_login.png",
    "图5.7 HMI阈值设置页面": ROOT / "docs/system_screenshots_updated/07_hmi_threshold_page.png",
    "图5.8 健康报告页面": ROOT / "docs/system_screenshots/04_health_report.png",
    "图5.9 智能运维诊断助手页面": ROOT / "docs/system_screenshots_updated/09_ai_qa_assistant_updated.png",
    "图5.11 系统参数设置页面": ROOT / "docs/system_screenshots/06_settings.png",
}


def replace_picture_before_caption(doc: Document, caption: str, image_path: Path) -> bool:
    if not image_path.exists():
        return False
    body = doc._body._element
    target = next((p for p in doc.paragraphs if p.text.strip() == caption), None)
    if target is None:
        return False
    children = list(body)
    pos = children.index(target._p)
    if pos > 0 and children[pos - 1].tag == qn("w:p") and children[pos - 1].xpath(".//w:drawing"):
        body.remove(children[pos - 1])
    pic_p = target._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Inches(5.85))
    target._p.addprevious(pic_p._p)
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in target.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10)
    return True


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    replaced = 0
    missing = []
    for caption, path in COLOR_SCREENSHOTS.items():
        if replace_picture_before_caption(doc, caption, path):
            replaced += 1
        else:
            missing.append(f"{caption} -> {path}")
    doc.save(OUT)
    checked = Document(OUT)
    captions = [p.text.strip() for p in checked.paragraphs if re.match(r"^[图表]\d+(\.\d+)+", p.text.strip())]
    print(f"saved={OUT}")
    print(f"replaced_color_screenshots={replaced}")
    print(f"missing={len(missing)}")
    for item in missing:
        print(item)
    print(f"tables={len(checked.tables)}")
    print(f"inline_shapes={len(checked.inline_shapes)}")
    print(f"captions={len(captions)}")


if __name__ == "__main__":
    main()
