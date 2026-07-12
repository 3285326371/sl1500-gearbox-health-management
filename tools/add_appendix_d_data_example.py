from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = next(Path.home().joinpath("Desktop").glob("毕业设计正文模板.docx"))
BACKUP = TEMPLATE.with_name(
    f"{TEMPLATE.stem}_附录D前备份_{datetime.now():%Y%m%d_%H%M%S}{TEMPLATE.suffix}"
)
TMP_OUT = Path.cwd() / "docs" / "_tmp_毕业设计正文模板_附录D.docx"


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def p(doc, text, first=True):
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Pt(24) if first else Pt(0)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    for r in para.runs:
        set_font(r)
    return para


def code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)
    for i, line in enumerate(text.strip().splitlines()):
        if i:
            para.add_run("\n")
        r = para.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
    return para


def main():
    copy2(TEMPLATE, BACKUP)
    doc = Document(TEMPLATE)
    if any("附录D 接口返回字段示例" in para.text for para in doc.paragraphs):
        print("already exists")
    else:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = doc.add_paragraph("附录D 接口返回字段示例")
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            set_font(r, east_asia="黑体", ascii_font="Arial", size=16, bold=True)
        p(doc, "为了便于说明前后端数据交互方式，本附录给出系统主要接口的返回字段示例。接口数据采用JSON格式，前端根据字段名称渲染卡片、表格、图表和报告内容。统一的字段结构有助于降低页面逻辑复杂度，也方便后续替换真实数据源。")
        code(doc, """
{
  "unit_id": "WTG-001",
  "timestamp": "2026-05-24 20:00:22",
  "oil_temp": 68.4,
  "vibration_rms": 3.64,
  "oil_quality": 7,
  "health_score": 90,
  "rul_days": 376,
  "status": "normal",
  "alarm_level": "正常"
}
""")
        p(doc, "实时状态接口主要用于风场总览、单机详情和健康报告页面。字段中的oil_temp表示齿轮箱油温，vibration_rms表示振动有效值，oil_quality表示油液NAS等级，health_score和rul_days用于展示健康评分和剩余寿命。")
        code(doc, """
{
  "fault_type": "轴承点蚀/剥落",
  "severity": "严重",
  "component": "高速轴承",
  "probability": 0.86,
  "features": {
    "rms": 4.91,
    "kurtosis": 6.72,
    "crest_factor": 4.33,
    "envelope_peak": 1.48
  },
  "advice": "复核包络谱并安排现场检查"
}
""")
        p(doc, "故障诊断接口主要服务于诊断页面和故障记录页面。系统不仅返回故障类型，还返回严重程度、部件位置、置信度、振动特征和维护建议，从而保证诊断结果具有可解释性。")
    TMP_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TMP_OUT)
    print(f"tmp_out={TMP_OUT}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
