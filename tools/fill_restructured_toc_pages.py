from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCX = next(p for p in (Path.home() / "Desktop").glob("*正文模板.docx") if not p.name.startswith("~$"))
PDF = max((Path.cwd() / "docs" / "qa_template_inplace").glob("*重构二检.pdf"), key=lambda p: p.stat().st_mtime)


def norm(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def set_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def main():
    docx = Document(DOCX)
    toc_title = next(p for p in docx.paragraphs if "目" in norm(p.text) and "录" in norm(p.text) and len(norm(p.text)) <= 4)
    title_idx = next(i for i, p in enumerate(docx.paragraphs) if p._element is toc_title._element)
    entries = []
    for p in docx.paragraphs[title_idx + 1:]:
        t = p.text.strip()
        if not t:
            continue
        if "\t" not in t and t.startswith("第1章"):
            break
        if "\t" in t:
            entries.append((p, t.split("\t", 1)[0].strip()))

    pdf = fitz.open(PDF)
    page_texts = [norm(page.get_text("text")) for page in pdf]
    chapter_phys = None
    for i, txt in enumerate(page_texts, start=1):
        if norm("第1章 绪论") in txt and norm("风力发电是新能源") in txt:
            chapter_phys = i
            break
    if not chapter_phys:
        raise RuntimeError("Cannot locate first chapter page")

    page_map = {}
    for _, title in entries:
        nt = norm(title)
        found = None
        for i in range(chapter_phys, len(page_texts) + 1):
            if nt in page_texts[i - 1]:
                found = i
                break
        page_map[title] = str(found - chapter_phys + 1) if found else ""

    for p, title in entries:
        level = 0
        if re.match(r"^\d+\.\d+\.\d+", title):
            level = 2
        elif re.match(r"^\d+\.\d+", title):
            level = 1
        p.clear()
        p.paragraph_format.left_indent = Pt(24 * level)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        set_font(r, 12 if level == 0 else 10.5, bold=(level == 0))
        p.add_run("\t")
        r2 = p.add_run(page_map.get(title, ""))
        set_font(r2, 12 if level == 0 else 10.5, bold=(level == 0))

    docx.save(DOCX)
    print(f"updated={DOCX}")
    print(f"chapter_phys={chapter_phys}")
    for _, title in entries:
        print(f"{title}\t{page_map.get(title, '')}")


if __name__ == "__main__":
    main()
